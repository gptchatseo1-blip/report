import io
import json
from datetime import date, timedelta
from unittest.mock import patch

import pytest
from django.contrib.auth import get_user_model
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import Client
from django.urls import reverse
from django.utils import timezone
from docx import Document
from openpyxl import load_workbook
from PIL import Image

from apps.metrics.models import KeywordPosition, RankingSnapshot
from apps.projects.models import Project
from apps.reports.exporting import generate_artifact
from apps.reports.forms import ReportCreateForm, parse_named_url_groups
from apps.reports.models import (
    GeneratedArtifact,
    ProjectReportSettings,
    Report,
    ReportDatasetSnapshot,
)
from apps.reports.services import create_report_version
from apps.topvisor.models import TopvisorProjectMapping

pytestmark = pytest.mark.django_db


def mapping(project, configurations):
    return TopvisorProjectMapping.objects.create(
        project=project, topvisor_project_id="1", selected_configurations=configurations
    )


def ranking(
    project,
    day,
    engine,
    config,
    *,
    url="https://example.test/page",
    region="Москва",
):
    snapshot = RankingSnapshot.objects.create(
        project=project,
        snapshot_date=day,
        search_engine=engine,
        region=region,
        ranking_depth=20,
        topvisor_configuration_id=config,
        response_checksum=f"sum-{config}-{day}",
    )
    KeywordPosition.objects.create(
        ranking_snapshot=snapshot,
        query="ключ",
        normalized_query="ключ",
        frequency=42,
        position_raw="5",
        position_value=5,
        position_status="ranked",
        group_name="Группа",
        target_url=url,
        normalized_target_url=url,
    )
    return snapshot


@pytest.fixture
def user():
    return get_user_model().objects.create_user("calendar-user", password="password")


@pytest.fixture
def project():
    return Project.objects.create(name="Calendar", domain="calendar.example")


def test_engines_have_independent_complete_date_sets(project):
    mapping(
        project,
        [
            {"id": "ya-a", "search_engine": "yandex"},
            {"id": "ya-b", "search_engine": "yandex"},
            {"id": "go", "search_engine": "google"},
        ],
    )
    for config in ("ya-a", "ya-b"):
        ranking(project, date(2026, 6, 1), "yandex", config)
    ranking(project, date(2026, 6, 2), "yandex", "ya-a")
    ranking(project, date(2026, 7, 1), "google", "go")
    ranking(project, date(2026, 7, 2), "google", "go")
    form = ReportCreateForm(project=project)
    assert [value for value, _ in form.fields["yandex_dates"].choices] == ["2026-06-01"]
    assert [value for value, _ in form.fields["google_dates"].choices] == [
        "2026-07-02",
        "2026-07-01",
    ]


def test_report_defaults_select_only_top_10_and_search_segment(project):
    form = ReportCreateForm(project=project)

    assert form["include_top_5"].value() is False
    assert form["include_top_10"].value() is True
    assert form["include_top_20"].value() is False
    assert form["include_top_11_30"].value() is False
    assert form["include_top_30"].value() is False
    assert form["metrika_search_segment"].value() is True
    assert form["include_metrika_sources_table"].value() is False


def test_project_report_settings_autosave_and_restore_are_isolated(client, user, project):
    other = Project.objects.create(name="Other", domain="other-settings.example")
    client.force_login(user)
    response = client.post(
        reverse("reports:report-settings-save", args=[project.id]),
        data=json.dumps(
            {
                "include_top_10": False,
                "include_top_20": True,
                "metrika_search_segment": False,
                "metrika_info_url_groups": ("Статьи | https://calendar.example/articles/*\n*smas*"),
                "metrika_bar_search_engines": ["google", "bing"],
            }
        ),
        content_type="application/json",
    )

    assert response.status_code == 200
    assert response.json()["ok"] is True
    assert ProjectReportSettings.objects.filter(project=project).exists()
    restored = ReportCreateForm(project=project)
    assert restored["include_top_10"].value() is False
    assert restored["include_top_20"].value() is True
    assert restored["metrika_search_segment"].value() is False
    assert restored["metrika_info_url_groups"].value().startswith("Статьи |")
    assert list(restored["metrika_bar_search_engines"].value()) == ["google", "bing"]
    untouched = ReportCreateForm(project=other)
    assert untouched["include_top_10"].value() is True
    assert untouched["metrika_search_segment"].value() is True


def test_named_url_groups_accept_masks_and_repeated_labels():
    groups = parse_named_url_groups(
        "Лечение | https://example.test/treatment/*\n"
        "Лечение | *smas*\n"
        "https://example.test/diagnostics/*"
    )

    assert groups == [
        {
            "name": "Лечение",
            "patterns": ["https://example.test/treatment/*", "*smas*"],
        },
        {
            "name": "Diagnostics",
            "patterns": ["https://example.test/diagnostics/*"],
        },
    ]


def test_uploaded_webmaster_query_screenshot_is_frozen_and_rendered(
    client, user, project, settings, tmp_path
):
    settings.MEDIA_ROOT = tmp_path
    image_data = io.BytesIO()
    Image.new("RGB", (240, 80), "white").save(image_data, format="PNG")
    screenshot = SimpleUploadedFile("queries.png", image_data.getvalue(), content_type="image/png")
    client.force_login(user)

    response = client.post(
        reverse("reports:report-create", args=[project.id]),
        {
            "include_webmaster": "on",
            "include_webmaster_popular_queries": "on",
            "webmaster_queries_screenshot": screenshot,
        },
    )

    assert response.status_code == 302
    version = Report.objects.get(project=project).versions.get()
    frozen = version.snapshot.payload["display_options"]["webmaster_queries_screenshot"]
    assert frozen["name"] == "queries.png"
    document = Document(
        io.BytesIO(
            generate_artifact(version=version, artifact_type="docx", is_draft=True).file.read()
        )
    )
    assert len(document.inline_shapes) == 1


def test_tampered_date_and_minimum_are_rejected_per_engine(project):
    mapping(
        project, [{"id": "ya", "search_engine": "yandex"}, {"id": "go", "search_engine": "google"}]
    )
    for engine, config in (("yandex", "ya"), ("google", "go")):
        ranking(project, date(2026, 7, 1), engine, config)
        ranking(project, date(2026, 7, 2), engine, config)
    form = ReportCreateForm(
        {"yandex_dates": ["2026-07-01", "2026-07-09"], "google_dates": ["2026-07-01"]},
        project=project,
    )
    assert not form.is_valid()
    assert "yandex_dates" in form.errors and "google_dates" in form.errors


@pytest.mark.parametrize(
    ("count", "message"), [(0, "нет доступных дат"), (1, "нужна ещё минимум одна дата")]
)
def test_connected_engine_with_too_few_dates_is_visible(client, user, project, count, message):
    mapping(project, [{"id": "go", "search_engine": "google"}])
    if count:
        ranking(project, date(2026, 7, 1), "google", "go")
    client.force_login(user)
    response = client.get(reverse("reports:report-list", args=[project.id]))
    html = response.content.decode().lower()
    assert "google" in html and message in html
    assert "disabled" in html


def test_bound_form_preserves_dates_and_show_urls_after_error(client, user, project):
    mapping(project, [{"id": "go", "search_engine": "google"}])
    for day in (date(2026, 7, 1), date(2026, 7, 2)):
        ranking(project, day, "google", "go")
    client.force_login(user)
    response = client.post(
        reverse("reports:report-create", args=[project.id]),
        {"google_dates": ["2026-07-01"], "show_urls": "on"},
    )
    assert response.status_code == 400
    assert response.context["form"]["show_urls"].value() is True
    assert response.context["form"]["google_dates"].value() == ["2026-07-01"]


@pytest.mark.parametrize(("checkbox", "expected"), [(None, False), ("on", True)])
def test_new_snapshot_stores_url_option(client, user, project, checkbox, expected):
    mapping(project, [{"id": "go", "search_engine": "google"}])
    for day in (date(2026, 7, 1), date(2026, 7, 2)):
        ranking(project, day, "google", "go")
    client.force_login(user)
    data = {"google_dates": ["2026-07-01", "2026-07-02"]}
    if checkbox:
        data["show_urls"] = checkbox
    assert client.post(reverse("reports:report-create", args=[project.id]), data).status_code == 302
    assert ReportDatasetSnapshot.objects.get().payload["display_options"]["show_urls"] is expected


def test_new_snapshot_stores_flexible_report_options(client, user, project):
    mapping(project, [{"id": "ya", "search_engine": "yandex"}])
    for day in (date(2026, 7, 1), date(2026, 7, 2)):
        ranking(project, day, "yandex", "ya")
    client.force_login(user)
    response = client.post(
        reverse("reports:report-create", args=[project.id]),
        {
            "yandex_dates": ["2026-07-01", "2026-07-02"],
            "include_top_tables": "on",
            "include_top_5": "on",
            "include_top_11_30": "on",
            "include_webmaster": "on",
            "include_metrika": "on",
            "include_metrika_geography": "on",
            "geography_moscow": "on",
            "include_topvisor_report_link": "on",
            "topvisor_report_url": "https://topvisor.example/report/42",
        },
    )
    assert response.status_code == 302
    options = ReportDatasetSnapshot.objects.get().payload["display_options"]
    assert options["configuration_version"] == 3
    assert options["include_visibility"] is False
    assert options["include_top_5"] is True
    assert options["include_top_10"] is False
    assert options["include_top_11_30"] is True
    assert options["geography_moscow"] is True
    assert options["geography_saint_petersburg"] is False
    assert options["topvisor_report_url"] == "https://topvisor.example/report/42"
    assert options["topvisor_report_urls"] == {"ya": "https://topvisor.example/report/42"}


def test_topvisor_report_link_requires_url(project):
    form = ReportCreateForm(
        {"include_topvisor_report_link": "on"},
        project=project,
    )
    assert not form.is_valid()
    assert "topvisor_report_url" in form.errors


def test_topvisor_report_links_are_required_and_saved_per_engine_region(client, user, project):
    mapping(
        project,
        [
            {
                "id": "ya-msk",
                "search_engine": "yandex",
                "region_name": "Москва",
            },
            {
                "id": "ya-spb",
                "search_engine": "yandex",
                "region_name": "Санкт-Петербург",
            },
        ],
    )
    for day in (date(2026, 7, 1), date(2026, 7, 31)):
        ranking(project, day, "yandex", "ya-msk", region="Москва")
        ranking(project, day, "yandex", "ya-spb", region="Санкт-Петербург")
    client.force_login(user)
    list_url = reverse("reports:report-list", args=[project.id])
    response = client.get(list_url)
    html = response.content.decode()
    assert "Яндекс · Москва" in html
    assert "Яндекс · Санкт-Петербург" in html
    assert 'name="topvisor_report_url_0"' in html
    assert 'name="topvisor_report_url_1"' in html

    invalid = ReportCreateForm(
        {
            "yandex_dates": ["2026-07-01", "2026-07-31"],
            "include_topvisor_report_link": "on",
            "topvisor_report_url_0": "https://topvisor.example/msk",
        },
        project=project,
    )
    assert not invalid.is_valid()
    assert "topvisor_report_url_1" in invalid.errors

    response = client.post(
        reverse("reports:report-create", args=[project.id]),
        {
            "yandex_dates": ["2026-07-01", "2026-07-31"],
            "include_topvisor_report_link": "on",
            "topvisor_report_url_0": "https://topvisor.example/msk",
            "topvisor_report_url_1": "https://topvisor.example/spb",
        },
    )
    assert response.status_code == 302
    options = ReportDatasetSnapshot.objects.get().payload["display_options"]
    assert options["topvisor_report_urls"] == {
        "ya-msk": "https://topvisor.example/msk",
        "ya-spb": "https://topvisor.example/spb",
    }


@pytest.mark.parametrize("show_urls", [False, True])
def test_xlsx_url_policy_covers_positions_and_work(project, tmp_path, settings, show_urls):
    settings.MEDIA_ROOT = tmp_path
    ranking(project, date(2026, 7, 31), "google", "")
    report = Report.objects.create(project=project, report_month=date(2026, 7, 1))
    version = create_report_version(
        report=report, selection={"display_options": {"show_urls": show_urls}}
    )
    payload = version.snapshot.payload
    payload["completed_work"] = [
        {
            "date": "2026-07-01",
            "category": "SEO",
            "title": "Статья",
            "status": "done",
            "page_or_material_name": "Полезное название",
            "url": "https://secret.test/work",
            "result_url": "https://secret.test/result",
            "character_count": 10,
            "responsible": "Иван",
            "comment": "готово",
        }
    ]
    ReportDatasetSnapshot.objects.filter(
        pk=version.snapshot_id if hasattr(version, "snapshot_id") else version.snapshot.pk
    ).update(payload=payload)
    version.snapshot.refresh_from_db()
    artifact = generate_artifact(version=version, artifact_type="xlsx", is_draft=True)
    workbook = load_workbook(io.BytesIO(artifact.file.read()), read_only=True)
    positions = list(workbook["Позиции"].values)
    work = list(workbook["Выполненные работы"].values)
    all_values = "\n".join(
        str(value) for sheet in workbook for row in sheet.values for value in row if value
    )
    assert ("Релевантный URL" in positions[0]) is show_urls
    assert ("URL" in work[0] and "Результат" in work[0]) is show_urls
    assert "Полезное название" in work[1]
    assert ("https://secret.test" in all_values) is show_urls


def test_legacy_segment_without_configuration_exports_from_frozen_snapshot(
    client, user, project, tmp_path, settings
):
    settings.MEDIA_ROOT = tmp_path
    ranking(project, date(2026, 7, 31), "google", "legacy-config")
    version = create_report_version(
        report=Report.objects.create(project=project, report_month=date(2026, 7, 1))
    )
    payload = version.snapshot.payload
    payload.pop("display_options", None)
    payload["source_selection"]["topvisor"] = {"selected_dates": ["2026-07-01", "2026-07-31"]}
    for segment in payload["calculated"]["positions"]["segments"]:
        segment.pop("configuration_id", None)
    ReportDatasetSnapshot.objects.filter(pk=version.snapshot.pk).update(payload=payload)
    version.snapshot.refresh_from_db()
    RankingSnapshot.objects.all().delete()
    client.force_login(user)
    preview = client.get(reverse("reports:version-detail", args=[version.id]))
    assert preview.status_code == 200 and preview.context["show_urls"] is True
    with (
        patch(
            "apps.topvisor.client.TopvisorClient._request", side_effect=AssertionError("API called")
        ),
        patch(
            "apps.metrics.models.RankingSnapshot.objects.filter",
            side_effect=AssertionError("live DB called"),
        ),
    ):
        docx = generate_artifact(version=version, artifact_type="docx", is_draft=True)
        xlsx = generate_artifact(version=version, artifact_type="xlsx", is_draft=True)
    document = Document(io.BytesIO(docx.file.read()))
    doc_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    workbook = load_workbook(io.BytesIO(xlsx.file.read()), read_only=True)
    assert "ключ" in doc_text and "https://example.test/page" not in doc_text
    assert list(workbook["Позиции"].values)[1][3] == "ключ"
    assert "Релевантный URL" in list(workbook["Позиции"].values)[0]


@pytest.mark.parametrize(
    ("kind", "target"), [("docx", "_docx"), ("pdf", "_pdf"), ("xlsx", "_xlsx")]
)
def test_export_errors_mark_artifact_failed(project, kind, target):
    version = create_report_version(
        report=Report.objects.create(project=project, report_month=date(2026, 7, 1))
    )
    with (
        patch(
            f"apps.reports.exporting.{target}", side_effect=RuntimeError("/secret/path token=bad")
        ),
        pytest.raises(RuntimeError),
    ):
        generate_artifact(version=version, artifact_type=kind, is_draft=True)
    artifact = GeneratedArtifact.objects.get()
    assert artifact.status == "failed" and "/secret/path" not in artifact.generation_log


def artifact(version, status, age=0, with_file=True):
    row = GeneratedArtifact.objects.create(
        report_version=version, artifact_type="docx", status=status
    )
    if with_file:
        row.file.save("safe.docx", ContentFile(b"data"))
    if age:
        GeneratedArtifact.objects.filter(pk=row.pk).update(
            created_at=timezone.now() - timedelta(seconds=age)
        )
        row.refresh_from_db()
    return row


@pytest.mark.parametrize("status", ["ready", "failed"])
def test_ready_and_failed_artifacts_delete(client, user, project, tmp_path, settings, status):
    settings.MEDIA_ROOT = tmp_path
    version = create_report_version(
        report=Report.objects.create(project=project, report_month=date(2026, 7, 1))
    )
    row = artifact(version, status)
    client.force_login(user)
    assert client.post(reverse("reports:artifact-delete", args=[row.id])).status_code == 302
    assert not GeneratedArtifact.objects.filter(pk=row.id).exists()
    assert ReportDatasetSnapshot.objects.filter(version=version).exists()


def test_stale_is_failed_but_fresh_generating_cannot_delete(client, user, project):
    version = create_report_version(
        report=Report.objects.create(project=project, report_month=date(2026, 7, 1))
    )
    stale = artifact(version, "generating", age=1000, with_file=False)
    fresh = artifact(version, "generating", with_file=False)
    client.force_login(user)
    client.get(reverse("reports:version-detail", args=[version.id]))
    stale.refresh_from_db()
    assert stale.status == "failed"
    client.post(reverse("reports:artifact-delete", args=[fresh.id]))
    assert GeneratedArtifact.objects.filter(pk=fresh.id).exists()


def test_version_with_active_export_cannot_delete(client, user, project):
    version = create_report_version(
        report=Report.objects.create(project=project, report_month=date(2026, 7, 1))
    )
    fresh = artifact(version, "generating", with_file=False)
    client.force_login(user)
    response = client.post(reverse("reports:version-delete", args=[version.id]), follow=True)
    assert response.status_code == 200
    assert "Нельзя удалить версию, пока для неё формируется файл" in response.content.decode()
    assert GeneratedArtifact.objects.filter(pk=fresh.id).exists()
    assert ReportDatasetSnapshot.objects.filter(version=version).exists()


def test_delete_is_post_login_csrf_and_missing_file_safe(user, project):
    version = create_report_version(
        report=Report.objects.create(project=project, report_month=date(2026, 7, 1))
    )
    row = artifact(version, "failed", with_file=False)
    url = reverse("reports:artifact-delete", args=[row.id])
    anonymous = Client()
    assert anonymous.post(url).status_code == 302
    client = Client(enforce_csrf_checks=True)
    client.force_login(user)
    assert client.get(url).status_code == 405
    assert client.post(url).status_code == 403
    client.get(reverse("reports:version-detail", args=[version.id]))
    token = client.cookies["csrftoken"].value
    assert client.post(url, HTTP_X_CSRFTOKEN=token).status_code == 302
