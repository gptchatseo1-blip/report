import io
import json
import shutil
import socket

import pytest
from django.core.exceptions import ValidationError
from django.core.management import call_command
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from apps.metrics.models import KeywordPosition, RankingSnapshot, SourceSnapshot
from apps.projects.models import Project
from apps.reports.demo import DEMO_DOMAIN, create_demo_project
from apps.reports.exporting import generate_artifact
from apps.reports.models import Report, ReportVersion

pytestmark = [
    pytest.mark.django_db,
    pytest.mark.skipif(
        shutil.which("libreoffice") is None and shutil.which("soffice") is None,
        reason="requires LibreOffice",
    ),
]


def _bytes(artifact):
    with artifact.file.open("rb") as stream:
        return stream.read()


def _docx_content(data):
    document = Document(io.BytesIO(data))
    headings = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
    text = "\n".join(p.text for p in document.paragraphs)
    tables = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert all(
        (round(section.page_width.cm, 1), round(section.page_height.cm, 1)) == (21.0, 29.7)
        for section in document.sections
    )
    return headings, text, tables


def test_reproducible_offline_demo_complete_e2e(monkeypatch, settings, tmp_path):
    settings.MEDIA_ROOT = tmp_path

    def deny_network(*args, **kwargs):
        raise AssertionError("demo attempted an external HTTP connection")

    monkeypatch.setattr(socket, "create_connection", deny_network)
    call_command("create_demo_project")
    project, report, version = create_demo_project()
    call_command("create_demo_project")
    project_again, report_again, version_again = create_demo_project()

    assert (project_again.pk, report_again.pk, version_again.pk) == (
        project.pk,
        report.pk,
        version.pk,
    )
    assert Project.objects.filter(normalized_domain=DEMO_DOMAIN).count() == 1
    assert Report.objects.filter(project=project).count() == 1
    assert ReportVersion.objects.filter(report=report).count() == 1
    assert RankingSnapshot.objects.filter(project=project).count() == 6
    assert SourceSnapshot.objects.filter(project=project).count() == 6
    assert KeywordPosition.objects.filter(ranking_snapshot__project=project).count() == 36
    assert not KeywordPosition.objects.filter(
        ranking_snapshot__project=project, frequency__isnull=True
    ).exists()

    payload = version.snapshot.payload
    serialized = json.dumps(payload, ensure_ascii=False).lower()
    assert "authorization" not in serialized and "oauth" not in serialized
    assert "@" not in serialized and "127.0.0.1" not in serialized
    assert payload["completed_work"]
    assert len(payload["project"]["brand_rules"]) == 2
    assert len(payload["project"]["url_groups"]) == 2
    segments = payload["calculated"]["positions"]["segments"]
    google = next(item for item in segments if item["search_engine"] == "google")
    yandex = next(item for item in segments if item["search_engine"] == "yandex")
    assert google["ranking_depth"] == 20
    assert set(google["distribution"]["ranges"]) == {"1-3", "4-10", "11-20"}
    assert google["distribution"]["top_30"] is None
    assert yandex["ranking_depth"] == 100
    assert set(yandex["distribution"]["ranges"]) == {
        "1-3",
        "4-10",
        "11-20",
        "21-30",
        "31-50",
        "51-100",
    }
    assert google["top_11_20"] and yandex["top_11_20"]
    assert (
        sum(
            "точная позиция не определена" in b.effective_text
            for b in version.narrative_blocks.all()
        )
        == 1
    )
    assert not version.validation_issues.filter(severity="error").exists()

    checksum = version.snapshot.checksum
    version.snapshot.payload = {}
    with pytest.raises(ValidationError):
        version.snapshot.save()
    version.snapshot.refresh_from_db()
    assert version.snapshot.checksum == checksum

    runs = []
    for _ in range(2):
        artifacts = {
            kind: _bytes(generate_artifact(version=version, artifact_type=kind, is_draft=False))
            for kind in ("docx", "pdf", "xlsx")
        }
        headings, text, tables = _docx_content(artifacts["docx"])
        workbook = load_workbook(io.BytesIO(artifacts["xlsx"]), read_only=True, data_only=True)
        workbook_signature = tuple(
            (sheet.title, tuple(tuple(row) for row in sheet.iter_rows(values_only=True)))
            for sheet in workbook.worksheets
        )
        workbook.close()
        pdf = PdfReader(io.BytesIO(artifacts["pdf"]))
        assert artifacts["pdf"].startswith(b"%PDF-") and pdf.pages
        assert all((page.extract_text() or "").strip() for page in pdf.pages)
        runs.append((headings, text, tables, workbook_signature))

    assert runs[0][0] == [
        "1) Видимость сайта в поисковых системах Яндекс и Google по основным ключевым словам",
        "2) Индексация сайта (Яндекс.Вебмастер)",
        "3) Сводная информация по переходам на сайт (Яндекс.Метрика)",
        "Выполненные работы",
    ]
    assert "Самые кликабельные запросы" in runs[0][1]
    assert "Популярные страницы входа" in runs[0][1]
    assert "Сводная информация по конверсии" in runs[0][1]
    assert "WS" in runs[0][2]
    assert "TOP-30" not in google["depth_comment"]
    assert runs[0] == runs[1]
