import base64
import io
import re
import shutil
import zipfile
from datetime import date
from decimal import Decimal
from pathlib import Path
from urllib.parse import urlsplit

import pytest
from django.contrib.auth import get_user_model
from django.utils import timezone
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from pypdf import PdfReader

from apps.metrics.models import KeywordPosition, RankingSnapshot
from apps.metrics.synthetic import sync_synthetic_metrics
from apps.projects.models import Project
from apps.reports.exporting import (
    _configure_document,
    _configured_groups_table,
    _goal_card,
    _landing_hierarchy_order,
    _landing_pages_table,
    _manual_topvisor_segment,
    _metrika_detail_table,
    _metrika_goal_icon,
    _monthly_topvisor_rows,
    _position_fill,
    _render_monthly_topvisor_table,
    _render_topvisor_comparison,
    _webmaster_popular_table,
    _webmaster_query_summary_from_changes,
    _webmaster_query_summary_table,
    generate_artifact,
)
from apps.reports.models import Report, ReportDatasetSnapshot
from apps.reports.services import create_report_version

pytestmark = pytest.mark.django_db


def test_manual_dynamics_merge_history_override_and_sort_by_month():
    segment = {
        "configuration_id": "google",
        "search_engine": "google",
        "region": "Россия",
        "ranking_depth": 30,
        "three_month_series": [
            {"month": "2026-07-31", "distribution": {"total": 100}},
            {"month": "2026-08-31", "distribution": {"total": 100}},
        ],
    }
    payload = {
        "display_options": {
            "topvisor_manual_rows": [
                {
                    "engine": "google",
                    "region": "Россия",
                    "month": "2026-06-01",
                    "top3": 20,
                    "top10": 50,
                    "top11_30": 30,
                    "top3_percent": 20,
                    "top10_percent": 50,
                    "top11_30_percent": 30,
                },
                {
                    "engine": "google",
                    "region": "Россия",
                    "month": "2026-05-01",
                    "top3": 10,
                    "top10": 40,
                    "top11_30": 20,
                    "top3_percent": 10,
                    "top10_percent": 40,
                    "top11_30_percent": 20,
                },
                {
                    "engine": "google",
                    "region": "Россия",
                    "month": "2026-07-01",
                    "top3": 27,
                    "top10": 59,
                    "top11_30": 32,
                    "top3_percent": 27,
                    "top10_percent": 59,
                    "top11_30_percent": 32,
                },
                {
                    "engine": "yandex",
                    "region": "Россия",
                    "month": "2026-04-01",
                    "top3": 99,
                    "top10": 99,
                    "top11_30": 0,
                    "top3_percent": 99,
                    "top10_percent": 99,
                    "top11_30_percent": 0,
                },
            ]
        }
    }

    merged = _manual_topvisor_segment(payload, segment)
    assert [str(row["month"])[:7] for row in merged["three_month_series"]] == [
        "2026-05",
        "2026-06",
        "2026-07",
        "2026-08",
    ]
    july = merged["three_month_series"][2]
    assert july["manual_override"] is True
    assert july["distribution"]["manual_buckets"]["1-10"] == {"count": 59, "share": 59.0}


def _ranking(project, month, depth=100, count=36, engine="google", region="Россия"):
    snapshot = RankingSnapshot.objects.create(
        project=project,
        snapshot_date=month,
        search_engine=engine,
        region=region,
        ranking_depth=depth,
        depth_raw=f"TOP-{depth}",
        visibility="18.75",
        tracked_keyword_count=count,
    )
    positions = []
    for index in range(count):
        position = index % depth + 1
        positions.append(
            KeywordPosition(
                ranking_snapshot=snapshot,
                query="=2+2" if index == 0 else f"демо запрос {index:04d}",
                normalized_query=f"демо запрос {index:04d}",
                frequency=index + 10,
                position_raw=str(position),
                position_value=position,
                position_status=KeywordPosition.Status.RANKED,
                group_name="@формула" if index == 0 else "Услуги",
                target_url=f"https://demo.example/services/{index}",
                normalized_target_url=f"https://demo.example/services/{index}",
            )
        )
    KeywordPosition.objects.bulk_create(positions)


@pytest.fixture
def rich_version():
    project = Project.objects.create(name="Демонстрационный проект", domain="demo.example")
    sync_synthetic_metrics(project=project, report_month=date(2026, 7, 1))
    for month in (date(2026, 5, 31), date(2026, 6, 30), date(2026, 7, 31)):
        _ranking(project, month)
        _ranking(project, month, depth=50, count=24, engine="yandex", region="Москва")
    report = Report.objects.create(project=project, report_month=date(2026, 7, 1))
    return create_report_version(report=report)


@pytest.fixture
def version():
    project = Project.objects.create(name="Демонстрационный проект", domain="plain.example")
    report = Report.objects.create(project=project, report_month=date(2026, 7, 1))
    return create_report_version(report=report)


def _artifact_bytes(artifact):
    with artifact.file.open("rb") as stream:
        return stream.read()


def test_clickable_queries_show_colored_absolute_difference():
    document = Document()
    _configure_document(document, "example.test", date(2026, 7, 1))
    table = _webmaster_popular_table(
        document,
        [{"query": "пример", "shows": 120, "clicks": 12, "ctr": 10, "average_position": 4}],
        [{"query": "пример", "shows": 100, "clicks": 10, "ctr": 9, "average_position": 5}],
    )

    for cell in table.rows[1].cells[1:]:
        assert len(cell.paragraphs) == 2
        assert str(cell.paragraphs[1].runs[0].font.color.rgb) == "26A95B"
    assert [cell.paragraphs[1].text for cell in table.rows[1].cells[1:]] == [
        "20",
        "2",
        "1",
        "1",
    ]


def test_webmaster_summary_matches_service_truncation_and_absolute_differences():
    document = Document()
    _configure_document(document, "example.test", date(2026, 8, 1))
    table = _webmaster_query_summary_table(
        document,
        {
            "shows": 732834,
            "clicks": 18012,
            "ctr": "2.457859",
            "average_position": "9.009",
        },
        {
            "shows": 692963,
            "clicks": 19356,
            "ctr": "2.787859",
            "average_position": "10.059",
        },
    )

    assert [cell.paragraphs[0].text for cell in table.rows[1].cells[1:]] == [
        "732834",
        "18012",
        "2,45",
        "9",
    ]
    assert [cell.paragraphs[1].text for cell in table.rows[1].cells[1:]] == [
        "39871",
        "1344",
        "0,33",
        "1,05",
    ]


@pytest.mark.parametrize(
    ("goal_type", "asset_name"),
    [
        ("Call", "call"),
        ("ConditionalCall", "call"),
        ("Action", "action"),
        ("Url", "url"),
        ("Step", "step"),
        ("Multi", "multi"),
    ],
)
def test_metrika_goal_types_use_supplied_exact_icon_assets(goal_type, asset_name):
    assert _metrika_goal_icon({"type": goal_type}) == asset_name
    asset = (
        Path("apps/reports")
        / "assets"
        / "metrika_goal_icons"
        / f"{asset_name}.png"
    )
    assert asset.exists() and asset.stat().st_size > 0


def test_landing_hierarchy_sorts_sections_and_children_by_current_visits():
    def values(visits):
        return {"visits": Decimal(visits)}

    hierarchy = {
        "https://site.test/": values(200),
        "https://site.test/blog/": values(70),
        "https://site.test/blog/b/": values(20),
        "https://site.test/blog/a/": values(40),
        "https://site.test/services/": values(100),
        "https://site.test/services/b/": values(30),
        "https://site.test/services/a/": values(60),
    }

    assert _landing_hierarchy_order(hierarchy) == [
        "https://site.test/",
        "https://site.test/services/",
        "https://site.test/services/a/",
        "https://site.test/services/b/",
        "https://site.test/blog/",
        "https://site.test/blog/a/",
        "https://site.test/blog/b/",
    ]


def test_landing_url_rows_include_centered_project_favicon():
    document = Document()
    _configure_document(document, "site.test", date(2026, 7, 1))
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQIHWP4z8DwHwAFgAI/"
        "6e0zNwAAAABJRU5ErkJggg=="
    )
    payload = {
        "project": {"favicon": {"mime_type": "image/png", "data": base64.b64encode(png).decode()}}
    }
    current = {
        "https://site.test/": {"visits": Decimal(10), "users": Decimal(8)},
        "https://site.test/blog/": {"visits": Decimal(7), "users": Decimal(6)},
    }

    table = _landing_pages_table(document, payload, current, {})

    assert len(document.inline_shapes) == 2
    assert all(
        row.cells[0].vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for row in table.rows[1:]
    )


def test_info_comparison_contains_only_aggregated_sections_sorted_by_visits():
    document = Document()
    _configure_document(document, "site.test", date(2026, 7, 1))

    def row(url, visits):
        return {
            "dimensions": [{"name": "Яндекс"}, {"name": url}],
            "visits": visits,
            "users": visits,
            "bounce_rate": 10,
        }

    current = [
        row("https://site.test/blog/a/", 30),
        row("https://site.test/blog/b/", 20),
        row("https://site.test/services/a/", 80),
    ]
    groups = [
        {"name": "Статьи", "patterns": ["https://site.test/blog/*"]},
        {"name": "Услуги", "patterns": ["https://site.test/services/*"]},
    ]

    table = _configured_groups_table(
        document,
        {"project": {"normalized_domain": "site.test"}},
        current,
        [],
        groups,
    )
    labels = [row.cells[0].text for row in table.rows[1:]]

    assert labels == [
        "Итого и среднее",
        "https://site.test/",
        "›  https://site.test/services/",
        "›  https://site.test/blog/",
    ]
    assert not any("/a/" in label or "/b/" in label for label in labels)


def test_each_metrika_goal_is_rendered_as_one_separate_image():
    document = Document()
    _configure_document(document, "example.test", date(2026, 7, 1))
    periods = [
        {"period_start": "2026-05-01", "conversion_rate": 1, "visits": 10, "reaches": 11},
        {"period_start": "2026-06-01", "conversion_rate": 2, "visits": 20, "reaches": 22},
        {"period_start": "2026-07-01", "conversion_rate": 3, "visits": 30, "reaches": 33},
    ]
    _goal_card(
        document,
        {"label": "Запись", "goal_id": "1", "conversion_rate": 3, "visits": 30, "reaches": 33},
        periods,
    )
    _goal_card(
        document,
        {"label": "Звонок", "goal_id": "2", "conversion_rate": 2, "visits": 20, "reaches": 22},
        periods,
    )

    assert len(document.inline_shapes) == 2
    assert len(document.tables) == 0


def test_manual_completed_work_keeps_paragraphs_lists_and_clickable_links(
    version, settings, tmp_path
):
    settings.MEDIA_ROOT = tmp_path
    payload = version.snapshot.payload
    payload.setdefault("display_options", {}).update(
        {
            "include_completed_work": True,
            "completed_work_text": (
                "<p><strong>Выполнен аудит</strong></p>"
                "<ul><li>Исправлены метатеги</li></ul>"
                '<ol><li><a href="https://example.test/result">Проверить результат</a></li></ol>'
            ),
        }
    )
    ReportDatasetSnapshot.objects.filter(pk=version.snapshot.pk).update(payload=payload)
    version.snapshot.refresh_from_db()

    data = _artifact_bytes(generate_artifact(version=version, artifact_type="docx", is_draft=True))
    document = Document(io.BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Выполнен аудит" in text
    assert "Исправлены метатеги" in text
    assert "Проверить результат" in text
    assert any(paragraph.style.name == "List Bullet" for paragraph in document.paragraphs)
    assert any(paragraph.style.name == "List Number" for paragraph in document.paragraphs)
    with zipfile.ZipFile(io.BytesIO(data)) as package:
        assert b"w:hyperlink" in package.read("word/document.xml")
        assert b"https://example.test/result" in package.read("word/_rels/document.xml.rels")


def test_legacy_webmaster_changes_use_provider_fidelity_summary_table_data():
    payload = {
        "calculated": {
            "sources": {
                "sources": {
                    "yandex_webmaster": {
                        "normalized_changes": {
                            "search_impressions": {"previous": "335656", "current": "291167"},
                            "search_clicks": {"previous": "16483", "current": "10264"},
                            "search_ctr": {"previous": "4.9107", "current": "3.5251"},
                            "average_position": {"previous": "7.6054", "current": "8.8972"},
                        }
                    }
                }
            }
        }
    }

    current, previous = _webmaster_query_summary_from_changes(payload)

    assert current == {
        "shows": "291167",
        "clicks": "10264",
        "ctr": "3.5251",
        "average_position": "8.8972",
    }
    assert previous == {
        "shows": "335656",
        "clicks": "16483",
        "ctr": "4.9107",
        "average_position": "7.6054",
    }


def test_position_fill_uses_distinct_light_green_bands():
    assert [_position_fill(position) for position in (1, 3, 4, 5, 6, 10, 11, 20, 21, 30)] == [
        "55C98A",
        "55C98A",
        "8FDDB0",
        "8FDDB0",
        "C9EFD7",
        "C9EFD7",
        "E2F5E9",
        "E2F5E9",
        "F1F9F4",
        "F1F9F4",
    ]


def test_monthly_topvisor_rows_use_one_boundary_point_per_calendar_month():
    segment = {
        "ranking_depth": 100,
        "three_month_series": [
            {"month": "2026-05-31", "visibility": "2", "distribution": {}},
            {"month": "2026-07-01", "visibility": "5", "distribution": {}},
            {"month": "2026-06-01", "visibility": "3", "distribution": {}},
            {"month": "2026-05-01", "visibility": "1", "distribution": {}},
            {"month": "2026-07-31", "visibility": "6", "distribution": {}},
            {"month": "2026-06-30", "visibility": "4", "distribution": {}},
        ],
    }

    rows = _monthly_topvisor_rows(segment)

    assert [row[0] for row in rows] == ["Май 2026", "Июнь 2026", "Июль 2026"]
    assert [row[1] for row in rows] == ["1%", "4%", "6%"]


def test_monthly_topvisor_rows_collapse_duplicate_month_start_dates():
    segment = {
        "ranking_depth": 100,
        "three_month_series": [
            {"month": "2026-05-01", "visibility": "1", "distribution": {}},
            {"month": "2026-05-01", "visibility": "2", "distribution": {}},
            {"month": "2026-06-01", "visibility": "3", "distribution": {}},
            {"month": "2026-06-01", "visibility": "4", "distribution": {}},
            {"month": "2026-07-01", "visibility": "5", "distribution": {}},
            {"month": "2026-07-01", "visibility": "6", "distribution": {}},
        ],
    }

    rows = _monthly_topvisor_rows(segment)

    assert [row[0] for row in rows] == ["Май 2026", "Июнь 2026", "Июль 2026"]
    assert [row[1] for row in rows] == ["1%", "4%", "6%"]


def test_monthly_topvisor_table_hides_visibility_column_when_disabled():
    document = Document()
    _configure_document(document, "site.test", date(2026, 7, 1))
    segment = {
        "ranking_depth": 30,
        "three_month_series": [
            {
                "month": "2026-07-01",
                "visibility": 12,
                "distribution": {
                    "total": 10,
                    "top_10": 4,
                    "ranges": {"1-3": 2, "4-10": 2, "11-20": 1, "21-30": 1},
                },
            }
        ],
    }
    _render_monthly_topvisor_table(document, segment, show_visibility=False)
    assert [cell.text for cell in document.tables[0].rows[0].cells] == [
        "Месяц",
        "в топ 3",
        "в топ 10",
        "в топ 11-30",
    ]


def test_topvisor_comparison_uses_percent_point_change_and_requested_labels():
    document = Document()
    _configure_document(document, "site.test", date(2026, 7, 1))
    segment = {
        "ranking_depth": 30,
        "three_month_series": [
            {
                "month": "2026-06-01",
                "distribution": {
                    "manual_buckets": {
                        "1-3": {"share": 34, "count": 340},
                        "1-10": {"share": 60, "count": 600},
                        "11-30": {"share": 40, "count": 400},
                    }
                },
            },
            {
                "month": "2026-07-01",
                "distribution": {
                    "manual_buckets": {
                        "1-3": {"share": 27, "count": 270},
                        "1-10": {"share": 59, "count": 590},
                        "11-30": {"share": 32, "count": 320},
                    }
                },
            },
        ],
    }

    _render_topvisor_comparison(document, segment, show_visibility=False)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Запросов в топ 3 — 27% (уменьшилось на 7,0%)." in text
    assert "Запросов в топ 10 — 59% (уменьшилось на 1,0%)." in text
    assert "Запросов в топ 11-30 — 32% (уменьшилось на 8,0%)." in text


def test_metrika_detail_table_adds_total_and_renders_missing_numbers_as_zero():
    document = Document()
    _configure_document(document, "site.test", date(2026, 7, 1))
    table = _metrika_detail_table(
        document,
        [("Москва", {"visits": None, "users": 2}, {"visits": 3, "users": None})],
        first_header="Регион",
        metrics=("visits", "users"),
    )
    assert table.rows[1].cells[0].text == "Итого и среднее"
    assert table.rows[2].cells[1].text == "0"
    assert table.rows[2].cells[2].text == "3"


def test_full_docx_matches_reference_report_structure_and_styles(rich_version, settings, tmp_path):
    settings.MEDIA_ROOT = tmp_path
    data = _artifact_bytes(
        generate_artifact(version=rich_version, artifact_type="docx", is_draft=True)
    )
    assert zipfile.is_zipfile(io.BytesIO(data))
    document = Document(io.BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "ЧЕРНОВИК" in text
    assert f"Версия {rich_version.number}" not in text
    assert "Источник данных" not in text
    assert "Сведения об источнике недоступны" not in text
    headings = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
    assert headings == [
        "1) Видимость сайта в поисковых системах Яндекс и Google по основным ключевым словам",
        "2) Индексация сайта (Яндекс.Вебмастер)",
        "3) Сводная информация по переходам на сайт (Яндекс.Метрика)",
        "4) Сводная информация по конверсии (Яндекс.Метрика).",
        "Выполненные работы",
    ]
    assert document.styles["Normal"].font.name == "Calibri"
    assert document.styles["Normal"].font.size.pt == 11
    assert document.core_properties.author == "SEO"
    assert document.core_properties.comments == "https://t.me/wmasterfl"
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]
    yandex_index = paragraph_text.index("Яндекс. Москва")
    google_index = paragraph_text.index("Google. Россия")
    assert yandex_index < paragraph_text.index("Запросы в TOP-10 по Яндекс.Москва") < google_index
    assert google_index < paragraph_text.index("Запросы в TOP-10 по Google.Россия")
    footer_text = document.sections[0].footer.paragraphs[0].text
    assert footer_text == ""
    for section in document.sections:
        dimensions = (round(section.page_width.cm, 1), round(section.page_height.cm, 1))
        assert dimensions == (21.0, 29.7)
    with zipfile.ZipFile(io.BytesIO(data)) as package:
        images = [name for name in package.namelist() if name.startswith("word/media/")]
        assert len(images) >= 8
        xml = package.read("word/document.xml")
        assert b"w:tblHeader" in xml
        assert b"w:cantSplit" in xml
        assert b"w:keepNext" in xml
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "Checksum / идентификатор" not in table_text
    assert "WS" in table_text
    assert "демо запрос 0019" in table_text
    yandex_table = next(
        table
        for table in document.tables
        if len(table.rows[0].cells) > 2 and table.rows[0].cells[2].text == "Yandex"
    )
    google_table = next(
        table
        for table in document.tables
        if len(table.rows[0].cells) > 2 and table.rows[0].cells[2].text == "Google"
    )
    assert [round(cell.width.cm, 2) for cell in yandex_table.rows[0].cells] == [
        7.74,
        1.5,
        1.56,
        8.25,
    ]
    assert [round(cell.width.cm, 2) for cell in google_table.rows[0].cells] == [
        8.63,
        1.17,
        1.52,
        7.42,
    ]
    monthly_tables = [
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells[:4]]
        == ["Месяц", "в топ 3", "в топ 10", "в топ 11-30"]
    ]
    assert len(monthly_tables) == 2
    assert all(
        re.fullmatch(r"\d+% \(\d+\)", row.cells[column].text)
        for table in monthly_tables
        for row in table.rows[1:]
        for column in range(1, 4)
    )
    assert all(
        row.cells[column].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        for table in monthly_tables
        for row in table.rows
        for column in range(1, 4)
    )
    assert all(
        cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    search_table = next(
        table for table in document.tables if table.rows[0].cells[0].text == "Поисковая система"
    )
    assert [cell.text for cell in search_table.rows[0].cells] == [
        "Поисковая система",
        "Визиты\nСегмент A",
        "Визиты\nСегмент B",
        "Посетители\nСегмент A",
        "Посетители\nСегмент B",
        "Отказы, %\nСегмент A",
        "Отказы, %\nСегмент B",
    ]
    assert search_table.rows[0].cells[0].paragraphs[0].runs[0].font.name == "Calibri"
    assert search_table.rows[0].cells[0].paragraphs[0].runs[0].font.size.pt == 11
    assert search_table.rows[0].cells[1].paragraphs[0].runs[0].font.size.pt == 9
    assert str(search_table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb) == "7A8796"
    assert search_table.rows[1].cells[0].paragraphs[0].runs[0].font.size.pt == 11
    assert search_table.rows[1].cells[1].paragraphs[0].runs[0].font.size.pt == 11
    search_data_row = next(
        row for row in search_table.rows[1:] if row.cells[0].text != "Итого и среднее"
    )
    assert search_data_row.cells[1].paragraphs[1].runs[0].font.size.pt == 8
    provider_tables = [
        table
        for table in document.tables
        if table.rows[0].cells[0].text in {"Группа запросов", "Запрос"}
    ]
    assert {table.rows[0].cells[0].text for table in provider_tables} == {
        "Группа запросов",
        "Запрос",
    }
    for table in provider_tables:
        assert table.rows[0].cells[0].paragraphs[0].runs[0].font.name == "Calibri"
        assert table.rows[0].cells[0].paragraphs[0].runs[0].font.size.pt == 11
        assert table.rows[1].cells[0].paragraphs[0].runs[0].font.size.pt == 11
        metric_paragraphs = table.rows[1].cells[1].paragraphs
        expected_sizes = [11, 8]
        assert [paragraph.runs[0].font.size.pt for paragraph in metric_paragraphs] == expected_sizes
        if table.rows[0].cells[0].text == "Группа запросов":
            assert str(metric_paragraphs[1].runs[0].font.color.rgb) in {"26A95B", "F04444"}
        else:
            assert str(metric_paragraphs[1].runs[0].font.color.rgb) in {"26A95B", "F04444"}
    landing_tables = [
        table for table in document.tables if table.rows[0].cells[0].text == "Страница входа"
    ]
    assert landing_tables
    assert all(
        "Ур. 1" not in [cell.text for cell in table.rows[0].cells] for table in landing_tables
    )
    assert not any(
        len(table.rows[0].cells) == 3 and "ID request" in table.rows[0].cells[1].text
        for table in document.tables
    )
    assert "4) Сводная информация по конверсии (Яндекс.Метрика)." in text
    assert "Ниже приведены диаграммы конверсии" in text
    assert "Динамика по поисковым системам за квартал:" in text
    assert "Период сравнения:" not in text
    assert "Трафик из региона «Не определено»" not in text
    assert "Трафик из региона «Область не определена»" not in text
    assert text.count("Выполненные работы отсутствуют.") == 1
    assert "Индекс качества сайта" in text
    warning_paragraphs = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Предупреждение:")
    ]
    assert len(warning_paragraphs) == len(set(warning_paragraphs))


def test_modern_geography_keeps_undefined_rows_out_of_chart_narrative(
    rich_version, settings, tmp_path
):
    settings.MEDIA_ROOT = tmp_path
    payload = rich_version.snapshot.payload
    payload["display_options"] = {
        "configuration_version": 3,
        "include_metrika": True,
        "include_metrika_geography": True,
        "geography_moscow": True,
        "geography_saint_petersburg": True,
    }
    ReportDatasetSnapshot.objects.filter(pk=rich_version.snapshot.pk).update(payload=payload)
    rich_version.snapshot.refresh_from_db()

    document = Document(
        io.BytesIO(
            _artifact_bytes(
                generate_artifact(version=rich_version, artifact_type="docx", is_draft=True)
            )
        )
    )
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    geography = next(table for table in document.tables if table.rows[0].cells[0].text == "Регион")
    labels = [row.cells[0].text for row in geography.rows[1:]]

    assert "Не определено" in labels
    assert "Область не определена" in labels
    assert "Трафик из региона «Не определено»" not in text
    assert "Трафик из региона «Область не определена»" not in text


def test_configured_url_segments_render_three_level_tables_and_separate_charts(
    rich_version, settings, tmp_path
):
    settings.MEDIA_ROOT = tmp_path
    payload = rich_version.snapshot.payload
    payload["display_options"] = {
        "configuration_version": 3,
        "include_metrika": True,
        "metrika_search_segment": True,
        "include_metrika_landing_page_comparison": True,
        "include_metrika_url_groups": True,
        "include_metrika_sections": True,
        "include_metrika_categories": True,
        "metrika_url_segments": {
            "information": [{"name": "Статьи", "patterns": ["https://demo.example/blog/*"]}],
            "commercial": [
                {"name": "Лечение", "patterns": ["https://demo.example/services/*"]},
                {
                    "name": "Диагностика",
                    "patterns": ["https://demo.example/catalog/diagnostics/*"],
                },
                {"name": "Реабилитация", "patterns": ["*rehabilitation*"]},
            ],
            "categories": [
                {
                    "name": "Приоритетные услуги",
                    "patterns": ["https://demo.example/services/priority/*"],
                }
            ],
        },
    }
    ReportDatasetSnapshot.objects.filter(pk=rich_version.snapshot.pk).update(payload=payload)
    rich_version.snapshot.refresh_from_db()

    document = Document(
        io.BytesIO(
            _artifact_bytes(
                generate_artifact(version=rich_version, artifact_type="docx", is_draft=True)
            )
        )
    )
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    landing_tables = [
        table for table in document.tables if table.rows[0].cells[0].text == "Страница входа"
    ]
    landing_urls = [row.cells[0].text for table in landing_tables for row in table.rows[1:]]

    assert "Данные по разделам" in text
    assert "Информационные разделы" in text
    assert "Коммерческие разделы" in text
    assert "Приоритетные услуги" in text
    assert "Коммерческий раздел «Лечение»" in text
    assert "Коммерческий раздел «Диагностика»" in text
    assert "Коммерческий раздел «Реабилитация»" in text
    assert "https://demo.example/services/priority/" in landing_urls
    assert all(
        len([part for part in urlsplit(url.lstrip("›⌄ ")).path.split("/") if part]) <= 2
        for url in landing_urls
    )


def test_modern_report_options_control_sections_and_top_tables(rich_version, settings, tmp_path):
    settings.MEDIA_ROOT = tmp_path
    version = create_report_version(
        report=rich_version.report,
        selection={
            "display_options": {
                "configuration_version": 2,
                "show_urls": False,
                "include_visibility": False,
                "include_monthly_dynamics": False,
                "include_top_tables": True,
                "include_top_5": True,
                "include_top_10": False,
                "include_top_20": False,
                "include_top_11_30": False,
                "include_top_30": False,
                "include_topvisor_report_link": True,
                "topvisor_report_url": "https://topvisor.example/report/42",
                "include_webmaster": False,
                "include_metrika": False,
                "include_metrika_geography": False,
            }
        },
    )
    data = _artifact_bytes(generate_artifact(version=version, artifact_type="docx", is_draft=True))
    document = Document(io.BytesIO(data))
    headings = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
    assert headings == [
        "1) Видимость сайта в поисковых системах Яндекс и Google по основным ключевым словам",
        "Выполненные работы",
    ]
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "https://topvisor.example/report/42" in text
    assert "Запросы в TOP-5 по Яндекс.Москва" in text
    assert "таблица запросов в TOP-5 сформирована" not in text
    with zipfile.ZipFile(io.BytesIO(data)) as package:
        assert b"55C98A" in package.read("word/document.xml")


def test_topvisor_links_follow_their_engine_and_region_data(rich_version, settings, tmp_path):
    settings.MEDIA_ROOT = tmp_path
    payload = rich_version.snapshot.payload
    config_by_engine = {"yandex": "ya-msk", "google": "go-ru"}
    for segment in payload["calculated"]["positions"]["segments"]:
        segment["configuration_id"] = config_by_engine[segment["search_engine"]]
    for source in payload["ranking_sources"]:
        source["configuration_id"] = config_by_engine[source["search_engine"]]
    payload.setdefault("display_options", {}).update(
        {
            "configuration_version": 3,
            "include_top_tables": True,
            "include_top_10": True,
            "include_topvisor_report_link": True,
            "topvisor_report_urls": {
                "ya-msk": "https://topvisor.example/yandex-moscow",
                "go-ru": "https://topvisor.example/google-russia",
            },
            "topvisor_report_url": "https://topvisor.example/yandex-moscow",
        }
    )
    ReportDatasetSnapshot.objects.filter(pk=rich_version.snapshot.pk).update(payload=payload)
    rich_version.snapshot.refresh_from_db()

    document = Document(
        io.BytesIO(
            _artifact_bytes(
                generate_artifact(version=rich_version, artifact_type="docx", is_draft=True)
            )
        )
    )
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    yandex_link = "Подробный отчёт Яндекс · Москва: https://topvisor.example/yandex-moscow"
    google_link = "Подробный отчёт Google · Россия: https://topvisor.example/google-russia"
    assert paragraphs.count(yandex_link) == 1
    assert paragraphs.count(google_link) == 1
    assert paragraphs.index("Запросы в TOP-10 по Яндекс.Москва") < paragraphs.index(yandex_link)
    assert paragraphs.index(yandex_link) < paragraphs.index("Google. Россия")
    assert paragraphs.index("Запросы в TOP-10 по Google.Россия") < paragraphs.index(google_link)


@pytest.mark.parametrize(
    ("depth", "expected_ranges"),
    [
        (10, {"1-3", "4-10"}),
        (20, {"1-3", "4-10", "11-20"}),
        (30, {"1-3", "4-10", "11-20", "21-30"}),
        (50, {"1-3", "4-10", "11-20", "21-30", "31-50"}),
        (100, {"1-3", "4-10", "11-20", "21-30", "31-50", "51-100"}),
    ],
)
def test_export_respects_each_confirmed_depth(depth, expected_ranges):
    project = Project.objects.create(name=f"Depth {depth}", domain=f"depth-{depth}.example")
    _ranking(project, date(2026, 7, 31), depth=depth, count=depth)
    report = Report.objects.create(project=project, report_month=date(2026, 7, 1))
    segment = create_report_version(report=report).snapshot.payload["calculated"]["positions"][
        "segments"
    ][0]
    assert set(segment["distribution"]["ranges"]) == expected_ranges
    assert (segment["distribution"]["top_30"] is None) is (depth < 30)


def test_google_depth_note_is_single_and_top_30_not_rendered_for_top_20(settings, tmp_path):
    settings.MEDIA_ROOT = tmp_path
    project = Project.objects.create(name="Depth", domain="depth.example")
    _ranking(project, date(2026, 7, 31), depth=20, count=20)
    version = create_report_version(
        report=Report.objects.create(project=project, report_month=date(2026, 7, 1))
    )
    document = Document(
        io.BytesIO(
            _artifact_bytes(generate_artifact(version=version, artifact_type="docx", is_draft=True))
        )
    )
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert text.count("Глубина проверки Google —") == 1
    assert "TOP-30" not in text


@pytest.mark.parametrize("depth", [30, 50, 100])
def test_top_30_is_rendered_only_when_confirmed(depth, settings, tmp_path):
    settings.MEDIA_ROOT = tmp_path
    project = Project.objects.create(name=f"Confirmed {depth}", domain=f"confirmed-{depth}.example")
    for month in (date(2026, 5, 31), date(2026, 6, 30), date(2026, 7, 31)):
        _ranking(project, month, depth=depth, count=depth)
    version = create_report_version(
        report=Report.objects.create(project=project, report_month=date(2026, 7, 1))
    )
    document = Document(
        io.BytesIO(
            _artifact_bytes(generate_artifact(version=version, artifact_type="docx", is_draft=True))
        )
    )
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "в топ 11-30" in text


@pytest.mark.parametrize(
    ("mode", "position", "present"),
    [("auto", 12, True), ("auto", 5, False), ("enabled", 5, True), ("disabled", 12, False)],
)
def test_top_11_20_project_modes(mode, position, present, settings, tmp_path):
    settings.MEDIA_ROOT = tmp_path
    project = Project.objects.create(
        name=f"Mode {mode}", domain=f"{mode}-{position}.example", top_11_20_mode=mode
    )
    snapshot = RankingSnapshot.objects.create(
        project=project,
        snapshot_date=date(2026, 7, 31),
        search_engine="google",
        region="Россия",
        ranking_depth=20,
        tracked_keyword_count=1,
    )
    KeywordPosition.objects.create(
        ranking_snapshot=snapshot,
        query="тест",
        normalized_query="тест",
        frequency=100,
        position_raw=str(position),
        position_value=position,
        position_status=KeywordPosition.Status.RANKED,
    )
    version = create_report_version(
        report=Report.objects.create(project=project, report_month=date(2026, 7, 1))
    )
    document = Document(
        io.BytesIO(
            _artifact_bytes(generate_artifact(version=version, artifact_type="docx", is_draft=True))
        )
    )
    text = "\n".join(p.text for p in document.paragraphs)
    assert ("Запросы в TOP-11–20" in text) is present


def test_xlsx_has_all_rows_native_types_and_formula_injection_protection(
    rich_version, settings, tmp_path
):
    settings.MEDIA_ROOT = tmp_path
    workbook = load_workbook(
        io.BytesIO(
            _artifact_bytes(
                generate_artifact(version=rich_version, artifact_type="xlsx", is_draft=True)
            )
        )
    )
    assert workbook.sheetnames == [
        "Метаданные",
        "Позиции",
        "История",
        "Метрика и Вебмастер",
        "Выполненные работы",
    ]
    assert "Provenance" not in [cell.value for cell in workbook["Позиции"][1]]
    assert "Provenance" not in [cell.value for cell in workbook["Метрика и Вебмастер"][1]]
    assert isinstance(workbook["Метаданные"]["B3"].value, date)
    exported_at = workbook["Метаданные"]["B6"].value
    expected_at = rich_version.snapshot.created_at.astimezone(
        timezone.get_current_timezone()
    ).replace(tzinfo=None)
    # XLSX stores datetimes with millisecond precision.
    assert abs(exported_at - expected_at).total_seconds() < 0.001
    positions = workbook["Позиции"]
    expected = sum(
        len(source["positions"]) for source in rich_version.snapshot.payload["ranking_sources"]
    )
    assert positions.max_row - 1 == expected
    assert positions["D2"].value == "'=2+2"
    assert positions["H2"].value == "'@формула"
    assert positions["E2"].data_type == "n"
    assert positions["C2"].is_date


def test_repeat_export_does_not_query_live_metric_sources(
    rich_version, settings, tmp_path, monkeypatch
):
    settings.MEDIA_ROOT = tmp_path

    def forbidden(*args, **kwargs):
        raise AssertionError("live metrics must not be queried during export")

    monkeypatch.setattr(RankingSnapshot.objects, "filter", forbidden)
    from apps.metrics.models import MetricPoint, SourceSnapshot

    monkeypatch.setattr(SourceSnapshot.objects, "filter", forbidden)
    monkeypatch.setattr(MetricPoint.objects, "filter", forbidden)
    first = generate_artifact(version=rich_version, artifact_type="xlsx", is_draft=True)
    second = generate_artifact(version=rich_version, artifact_type="xlsx", is_draft=True)
    assert first.status == second.status == "ready"
    first_book = load_workbook(io.BytesIO(_artifact_bytes(first)))
    second_book = load_workbook(io.BytesIO(_artifact_bytes(second)))
    assert list(first_book["Позиции"].values) == list(second_book["Позиции"].values)


def test_visibility_and_traffic_render_only_precalculated_series(rich_version, settings, tmp_path):
    settings.MEDIA_ROOT = tmp_path
    payload = rich_version.snapshot.payload
    segment = payload["calculated"]["positions"]["segments"][0]
    assert len(segment["three_month_series"]) == 3
    # Extra raw rows must not become chart/table points during export.
    payload["ranking_sources"].append(
        {**payload["ranking_sources"][-1], "id": "duplicate", "visibility": "999"}
    )
    traffic = payload["calculated"]["sources"]["sources"]["yandex_metrika"][
        "traffic_source_dynamics"
    ]
    expected_current = next(iter(traffic.values()))["change"]["current"]
    payload.setdefault("display_options", {})["include_metrika_sources_table"] = True
    for source in payload["source_snapshots"]:
        if source.get("source") == "yandex_metrika":
            source["metrics"] = []
    ReportDatasetSnapshot.objects.filter(pk=rich_version.snapshot.pk).update(payload=payload)
    document = Document(
        io.BytesIO(
            _artifact_bytes(
                generate_artifact(version=rich_version, artifact_type="docx", is_draft=True)
            )
        )
    )
    table_values = [
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    ]
    assert "999" not in table_values
    assert _number_text(expected_current) in table_values


def test_iks_narrative_uses_quality_index_data_instead_of_missing_fallback(rich_version):
    block = rich_version.narrative_blocks.get(section_code="iks")
    changes = rich_version.snapshot.payload["calculated"]["sources"]["sources"]["yandex_webmaster"][
        "normalized_changes"
    ]["quality_index"]
    assert block.generated_text != "Данные раздела отсутствуют."
    assert "ИКС" in block.generated_text
    assert _number_text(changes["current"]) in block.generated_text


def test_download_requires_login_and_generation_get_is_rejected(client, version):
    response = client.get(f"/versions/{version.id}/export/docx/")
    assert response.status_code == 302
    user = get_user_model().objects.create_user("exporter", password="secret-pass")
    client.force_login(user)
    response = client.get(f"/versions/{version.id}/export/docx/")
    assert response.status_code == 405


@pytest.mark.skipif(
    shutil.which("libreoffice") is None, reason="LibreOffice is tested in Docker CI"
)
def test_real_pdf_conversion_smoke(rich_version, settings, tmp_path):
    settings.MEDIA_ROOT = tmp_path
    artifact = generate_artifact(version=rich_version, artifact_type="pdf", is_draft=True)
    data = _artifact_bytes(artifact)
    assert data.startswith(b"%PDF-")
    pages = PdfReader(io.BytesIO(data)).pages
    assert len(pages) > 0
    allowed = ((595.28, 841.89), (841.89, 595.28))
    for page in pages:
        size = (float(page.mediabox.width), float(page.mediabox.height))
        assert any(
            abs(size[0] - width) <= 3 and abs(size[1] - height) <= 3 for width, height in allowed
        )
        assert len((page.extract_text() or "").strip()) >= 20 or list(page.images)
    texts = [page.extract_text() or "" for page in pages]
    assert f"Версия {rich_version.number}" not in "\n".join(texts)
    assert f"версия {rich_version.number}" not in "\n".join(texts).lower()
    assert "Источник данных" not in "\n".join(texts)
    assert "Checksum / идентификатор" not in "\n".join(texts)
    for heading in ("Яндекс.Вебмастер", "Яндекс.Метрика"):
        section_text = next(text for text in texts if heading in text)
        for unavailable in (
            "Данные раздела отсутствуют",
            "Данные недоступны",
            "Источник недоступен",
            "Сведения об источнике недоступны",
        ):
            assert unavailable not in section_text
    page_text = next(text for text in texts if "TOP-11–20" in text)
    assert "Запросы" in page_text and "WS" in page_text


def _number_text(value):
    return format(Decimal(str(value)).normalize(), "f").replace(".", ",")
