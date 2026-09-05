from pathlib import Path

import pytest
from django.conf import settings
from docx import Document

from apps.reports import exporting

pytestmark = pytest.mark.django_db


def _distribution():
    return {
        "total": 2435,
        "top_10": 754,
        "ranges": {
            "1-3": 253,
            "4-10": 501,
            "11-20": 826,
            "21-30": 0,
            "31-50": 337,
            "51-100": 159,
        },
    }


def test_distribution_cards_are_narrow_compact_and_use_11pt_text():
    document = Document()
    exporting._render_distribution_cards_table(document, _distribution(), 100, engine="yandex")

    outer = document.tables[0]
    assert len(outer.rows) == 3
    assert len(outer.columns) == 2
    assert outer.rows[0].cells[0].width.cm < 5

    nested_tables = [cell.tables[0] for row in outer.rows for cell in row.cells if cell.tables]
    assert len(nested_tables) == 6
    for nested in nested_tables:
        for cell in nested.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    assert round(run.font.size.pt) == 11

    spacer_sizes = []
    for row in outer.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text.strip() == "\u200b":
                    spacer_sizes.extend(run.font.size.pt for run in paragraph.runs if run.font.size)
    assert spacer_sizes
    assert max(spacer_sizes) <= 1


def test_topvisor_connection_has_visible_ajax_sync_status_assets():
    base = Path(settings.BASE_DIR)
    template = (base / "templates" / "topvisor" / "connection.html").read_text()
    javascript = (base / "static" / "topvisor" / "sync-status.js").read_text()
    css = (base / "static" / "topvisor" / "sync-status.css").read_text()

    assert "data-topvisor-sync-form" in template
    assert "data-topvisor-sync-status" in template
    assert "sync-status.js" in template
    assert "sync-status.css" in template
    assert "Синхронизация Topvisor идёт" in javascript
    assert "X-Requested-With" in javascript
    assert "aria-busy" in javascript
    assert ".topvisor-sync-spinner" in css
    assert "@keyframes topvisor-sync-spin" in css
