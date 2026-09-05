import json
from pathlib import Path

import pytest
from docx import Document

from apps.reports import exporting
from apps.reports.runtime_fixes_round2 import (
    _display_percent,
    _manual_topvisor_segment,
    validate_manual_rows,
)


def _distribution(total=100, top3=10, top10=30, top11=20):
    return {
        "total": total,
        "top_10": top10,
        "ranges": {
            "1-3": top3,
            "4-10": top10 - top3,
            "11-20": top11,
            "21-30": 0,
        },
    }


def _segment():
    return {
        "configuration_id": "yandex-msk",
        "search_engine": "yandex",
        "region": "Москва",
        "ranking_depth": 30,
        "three_month_series": [
            {
                "month": "2026-07-01",
                "visibility": 12.4,
                "distribution": _distribution(),
            },
            {
                "month": "2026-08-01",
                "visibility": 15.65,
                "distribution": _distribution(),
            },
        ],
    }


def _row(month, *, include=True, manual_override=False, visibility=None):
    return {
        "configuration_id": "yandex-msk",
        "engine": "yandex",
        "region": "Москва",
        "month": month,
        "include_in_report": include,
        "manual_override": manual_override,
        "visibility": visibility,
        "automatic_visibility": 10,
        "total": 100,
        "top3": 11,
        "top10": 35,
        "top11_30": 27,
        "top3_percent": 11,
        "top10_percent": 35,
        "top11_30_percent": 27,
    }


@pytest.mark.parametrize(
    ("source", "expected"),
    [(15.65, 16), (15.50, 16), (15.49, 15), (0, 0), (100, 100)],
)
def test_visibility_display_uses_half_up_rounding(source, expected):
    assert _display_percent(source) == expected


def test_validator_preserves_explicit_row_activation_and_legacy_default():
    explicit = validate_manual_rows([_row("2026-08-01", include=False)])[0]
    legacy_source = _row("2026-07-01")
    legacy_source.pop("include_in_report")
    legacy = validate_manual_rows([legacy_source])[0]

    assert explicit["include_in_report"] is False
    assert explicit["include_explicit"] is True
    assert legacy["include_in_report"] is True
    assert legacy["include_explicit"] is False


def test_four_activated_months_are_exposed_to_monthly_report_table():
    rows = validate_manual_rows(
        [
            _row("2026-05-01"),
            _row("2026-06-01"),
            _row("2026-07-01"),
            _row("2026-08-01"),
        ]
    )
    result = _manual_topvisor_segment(
        {"display_options": {"topvisor_manual_rows": rows}},
        _segment(),
    )

    assert [item["month"][:7] for item in result["monthly_table_series"]] == [
        "2026-05",
        "2026-06",
        "2026-07",
        "2026-08",
    ]


def test_inactive_month_stays_out_of_monthly_report_table():
    rows = validate_manual_rows(
        [
            _row("2026-07-01", include=False),
            _row("2026-08-01", include=True),
        ]
    )
    result = _manual_topvisor_segment(
        {"display_options": {"topvisor_manual_rows": rows}},
        _segment(),
    )
    assert [item["month"][:7] for item in result["monthly_table_series"]] == ["2026-08"]


def test_active_manual_visibility_overrides_automatic_value():
    rows = validate_manual_rows(
        [_row("2026-08-01", include=True, manual_override=True, visibility=13.5)]
    )
    result = _manual_topvisor_segment(
        {"display_options": {"topvisor_manual_rows": rows}},
        _segment(),
    )
    assert result["three_month_series"][-1]["visibility"] == 13.5
    assert result["monthly_table_series"][-1]["visibility"] == 13.5


def test_deleted_row_is_not_exposed_to_monthly_report_table():
    source = _row("2026-08-01", include=True)
    source["deleted"] = True
    rows = validate_manual_rows([source])
    result = _manual_topvisor_segment(
        {"display_options": {"topvisor_manual_rows": rows}},
        _segment(),
    )
    assert result["monthly_table_series"] == []


def test_legacy_manual_rows_do_not_drop_existing_months():
    source = _row("2026-08-01", manual_override=True, visibility=13)
    source.pop("include_in_report")
    rows = validate_manual_rows([source])
    result = _manual_topvisor_segment(
        {"display_options": {"topvisor_manual_rows": rows}},
        _segment(),
    )
    assert [item["month"][:7] for item in result["monthly_table_series"]] == [
        "2026-07",
        "2026-08",
    ]


def test_compact_distribution_table_uses_eleven_point_text_and_three_rows():
    doc = Document()
    table = exporting._render_distribution_cards_table(
        doc,
        {
            "total": 100,
            "ranges": {
                "1-3": 10,
                "4-10": 20,
                "11-20": 20,
                "21-30": 15,
                "31-50": 10,
                "51-100": 15,
            },
        },
        100,
        engine="yandex",
    )
    assert len(table.rows) == 3
    sizes = []
    for outer_row in table.rows:
        for outer_cell in outer_row.cells:
            for nested in outer_cell.tables:
                for cell in nested.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        sizes.extend(run.font.size.pt for run in paragraph.runs if run.font.size)
    assert sizes and all(size == 11 for size in sizes)


def test_round2_ui_contract_contains_dashed_trigger_and_row_actions():
    root = Path(__file__).resolve().parents[1]
    js = (root / "static/reports/report-polish-round2.js").read_text()
    css = (root / "static/reports/report-polish-round2.css").read_text()

    assert "trigger.textContent = 'Скорректировать'" in js
    assert "include_in_report" in js
    assert "manual-row-include" in js
    assert "manual-row-delete-round2" in js
    assert "border-bottom:1px dashed currentColor" in css
    assert "Скорректировать таблицы динамики" in js


def test_round2_serialized_schema_is_json_safe():
    rows = validate_manual_rows([_row("2026-08-01")])
    json.dumps(rows, ensure_ascii=False)
