import json
from datetime import date
from pathlib import Path

import pytest
from django.contrib.auth import get_user_model
from django.contrib.staticfiles.storage import staticfiles_storage
from django.core.exceptions import ValidationError
from django.urls import reverse
from docx import Document

from apps.metrics.models import RankingSnapshot, SourceSnapshot
from apps.projects.models import Project
from apps.reports import exporting
from apps.reports.forms import validate_topvisor_manual_rows
from apps.reports.models import ProjectReportSettings, Report, ReportDatasetSnapshot
from apps.topvisor.models import TopvisorProjectMapping

pytestmark = pytest.mark.django_db


@pytest.fixture
def user(db):
    return get_user_model().objects.create_user(
        username="report-polish",
        password="test-password",
    )


@pytest.fixture
def project(db):
    return Project.objects.create(
        name="Report polish",
        domain="report-polish.example",
    )


def _connect_google(project):
    TopvisorProjectMapping.objects.create(
        project=project,
        topvisor_project_id="42",
        selected_configurations=[
            {
                "id": "google-main",
                "search_engine": "google",
                "region_name": "Москва",
                "normalized_depth": 30,
            }
        ],
    )


def _ranking(project, day, visibility="16"):
    return RankingSnapshot.objects.create(
        project=project,
        snapshot_date=day,
        search_engine="google",
        region="Москва",
        topvisor_configuration_id="google-main",
        ranking_depth=30,
        visibility=visibility,
        response_checksum=f"google-{day}",
    )


def _distribution(total=100):
    return {
        "total": total,
        "top_10": 30,
        "ranges": {
            "1-3": 10,
            "4-10": 20,
            "11-20": 20,
            "21-30": 15,
            "31-50": 10,
            "51-100": 15,
        },
    }


def _segment():
    return {
        "configuration_id": "google-main",
        "search_engine": "google",
        "region": "Москва",
        "ranking_depth": 30,
        "three_month_series": [
            {
                "month": "2026-07-01",
                "visibility": 12,
                "distribution": _distribution(),
            },
            {
                "month": "2026-08-01",
                "visibility": 16,
                "distribution": _distribution(),
            },
        ],
    }


def _manual(visibility=13, month="2026-08-01", region="Москва"):
    return {
        "configuration_id": "google-main",
        "engine": "google",
        "region": region,
        "month": month,
        "visibility": visibility,
        "total": 100,
        "top3": 11,
        "top10": 35,
        "top11_30": 27,
        "top3_percent": 11,
        "top10_percent": 35,
        "top11_30_percent": 27,
    }


def test_report_builder_hides_manual_month_and_loads_modal_assets(
    client,
    user,
    project,
):
    client.force_login(user)
    html = client.get(reverse("reports:report-list", args=[project.id])).content.decode()

    assert "Отчётный месяц" not in html
    assert 'name="month"' in html and 'type="hidden"' in html
    assert staticfiles_storage.url("reports/report-polish.css") in html
    assert staticfiles_storage.url("reports/report-polish.js") in html


def test_august_position_dates_create_august_report_even_when_current_period_is_later(
    client,
    user,
    project,
    monkeypatch,
):
    _connect_google(project)
    _ranking(project, date(2026, 7, 14), "12")
    _ranking(project, date(2026, 8, 17), "16")
    monkeypatch.setattr(
        "apps.reports.forms.timezone.localdate",
        lambda: date(2026, 9, 5),
    )
    monkeypatch.setattr(
        "apps.reports.views.timezone.localdate",
        lambda: date(2026, 9, 5),
    )
    client.force_login(user)

    response = client.post(
        reverse("reports:report-create", args=[project.id]),
        {
            "google_dates": ["2026-07-14", "2026-08-17"],
            "topvisor_manual_rows": "[]",
        },
    )

    assert response.status_code == 302
    assert Report.objects.get().report_month == date(2026, 8, 1)


def test_without_position_dates_report_month_comes_from_latest_selected_source_snapshot(
    client,
    user,
    project,
):
    older = SourceSnapshot.objects.create(
        project=project,
        source=SourceSnapshot.Source.METRIKA,
        period_start=date(2026, 7, 1),
        period_end=date(2026, 7, 31),
        checksum="july",
    )
    august = SourceSnapshot.objects.create(
        project=project,
        source=SourceSnapshot.Source.METRIKA,
        period_start=date(2026, 8, 1),
        period_end=date(2026, 8, 31),
        checksum="august",
    )
    client.force_login(user)

    response = client.post(
        reverse("reports:report-create", args=[project.id]),
        {
            "include_metrika": "on",
            "metrika_snapshots": [str(older.id), str(august.id)],
            "topvisor_manual_rows": "[]",
        },
    )

    assert response.status_code == 302
    assert Report.objects.get().report_month == date(2026, 8, 1)


def test_without_selected_data_report_month_uses_previous_month_fallback(
    client,
    user,
    project,
    monkeypatch,
):
    monkeypatch.setattr(
        "apps.reports.forms.timezone.localdate",
        lambda: date(2026, 9, 5),
    )
    monkeypatch.setattr(
        "apps.reports.views.timezone.localdate",
        lambda: date(2026, 9, 5),
    )
    client.force_login(user)
    response = client.post(
        reverse("reports:report-create", args=[project.id]),
        {"topvisor_manual_rows": "[]"},
    )
    assert response.status_code == 302
    assert Report.objects.get().report_month == date(2026, 8, 1)


def test_manual_visibility_parser_accepts_percent_comma_zero_and_empty():
    comma = validate_topvisor_manual_rows([_manual("16,5%")])[0]
    zero = validate_topvisor_manual_rows([_manual(0)])[0]
    empty = validate_topvisor_manual_rows([_manual(None)])[0]
    assert comma["visibility"] == 16.5
    assert zero["visibility"] == 0
    assert empty["visibility"] is None


@pytest.mark.parametrize("value", ["abc", -1, 100.01, "101%"])
def test_manual_visibility_rejects_invalid_values(value):
    with pytest.raises(ValidationError):
        validate_topvisor_manual_rows([_manual(value)])


def test_manual_visibility_overrides_auto_value_without_mutating_source_snapshot(
    project,
):
    source = _ranking(project, date(2026, 8, 17), "16")
    payload = {"display_options": {"topvisor_manual_rows": [_manual(13)]}}
    effective = exporting._manual_topvisor_segment(payload, _segment())

    assert effective["three_month_series"][-1]["visibility"] == 13
    source.refresh_from_db()
    assert float(source.visibility) == 16


def test_missing_manual_visibility_returns_to_automatic_value():
    row = _manual(None)
    effective = exporting._manual_topvisor_segment(
        {"display_options": {"topvisor_manual_rows": [row]}},
        _segment(),
    )
    assert effective["three_month_series"][-1]["visibility"] == 16


def test_manual_zero_visibility_is_not_replaced_by_automatic_value():
    effective = exporting._manual_topvisor_segment(
        {"display_options": {"topvisor_manual_rows": [_manual(0)]}},
        _segment(),
    )
    assert effective["three_month_series"][-1]["visibility"] == 0


def test_manual_previous_month_is_added_to_effective_chart_series():
    effective = exporting._manual_topvisor_segment(
        {"display_options": {"topvisor_manual_rows": [_manual(9, "2026-06-01")]}},
        _segment(),
    )
    assert [row["month"][:7] for row in effective["chart_series"]] == [
        "2026-06",
        "2026-07",
        "2026-08",
    ]
    assert effective["chart_series"][0]["visibility"] == 9
    first_buckets = effective["chart_series"][0]["distribution"]["manual_buckets"]
    assert first_buckets["1-10"]["share"] == 35


def test_manual_visibility_isolated_by_search_engine_and_region():
    effective = exporting._manual_topvisor_segment(
        {"display_options": {"topvisor_manual_rows": [_manual(3, region="Санкт-Петербург")]}},
        _segment(),
    )
    assert effective["three_month_series"][-1]["visibility"] == 16


def test_saved_manual_visibility_survives_source_data_refresh(
    client,
    user,
    project,
):
    client.force_login(user)
    response = client.post(
        reverse("reports:report-settings-save", args=[project.id]),
        data=json.dumps({"topvisor_manual_rows": json.dumps([_manual(13)])}),
        content_type="application/json",
    )
    assert response.status_code == 200
    _ranking(project, date(2026, 8, 17), "18")

    settings_values = ProjectReportSettings.objects.get(project=project).values
    stored = json.loads(settings_values["topvisor_manual_rows"])
    assert stored[0]["visibility"] == 13


def test_created_report_version_freezes_manual_visibility(
    client,
    user,
    project,
):
    client.force_login(user)
    response = client.post(
        reverse("reports:report-create", args=[project.id]),
        {
            "month": "2026-08",
            "topvisor_manual_rows": json.dumps([_manual(13)]),
        },
    )
    assert response.status_code == 302
    snapshot = ReportDatasetSnapshot.objects.get()
    frozen = snapshot.payload["display_options"]["topvisor_manual_rows"]
    assert frozen[0]["visibility"] == 13

    client.post(
        reverse("reports:report-settings-save", args=[project.id]),
        data=json.dumps({"topvisor_manual_rows": json.dumps([_manual(18)])}),
        content_type="application/json",
    )
    snapshot.refresh_from_db()
    assert snapshot.payload["display_options"]["topvisor_manual_rows"][0]["visibility"] == 13


def test_visibility_chart_uses_fifty_percent_ceiling_for_low_series(monkeypatch):
    monkeypatch.setattr(exporting, "_save_figure", lambda figure: figure)
    figure = exporting._visibility_chart(
        [
            ("2026-06-17", 7),
            ("2026-07-14", 10),
            ("2026-08-17", 13),
        ]
    )
    try:
        assert figure.axes[0].get_ylim() == (0.0, 50.0)
        assert list(figure.axes[0].get_yticks()) == [0, 25, 50]
        assert len(figure.axes[1].patches) == 2
        assert all(getattr(wedge, "width", None) == 0.35 for wedge in figure.axes[1].patches)
    finally:
        exporting.plt.close(figure)


def test_distribution_legend_uses_topvisor_dot_markers_and_exact_palette(
    monkeypatch,
):
    monkeypatch.setattr(exporting, "_save_figure", lambda figure: figure)
    history = [
        {"month": "2026-07-01", "distribution": _distribution()},
        {"month": "2026-08-01", "distribution": _distribution()},
    ]
    figure = exporting._distribution_chart(history, 100)
    try:
        legend = figure.axes[0].get_legend()
        assert [text.get_text() for text in legend.get_texts()] == [
            "1-3",
            "1-10",
            "11-30",
            "31-50",
            "51-100",
            "101+",
        ]
        handles = getattr(
            legend,
            "legend_handles",
            getattr(legend, "legendHandles", []),
        )
        assert handles
        assert all(handle.get_marker() == "o" for handle in handles)
        expected_colors = {
            "1-3": "#3198DD",
            "1-10": "#21936C",
            "11-20": "#1ABC9C",
            "11-30": "#1ABC9C",
            "31-50": "#A2DF9F",
            "51-100": "#B0C7C7",
            "101+": "#FBC02D",
        }
        for name, color in expected_colors.items():
            assert exporting.TOPVISOR_COLORS[name] == color
    finally:
        exporting.plt.close(figure)


def test_visibility_sentence_agrees_with_feminine_noun():
    document = Document()
    segment = _segment()
    segment["three_month_series"][0]["visibility"] = 20
    segment["three_month_series"][1]["visibility"] = 29
    exporting._render_topvisor_comparison(document, segment)
    visibility_line = document.paragraphs[-1].text
    assert "Общая видимость сайта" in visibility_line
    assert "увеличилась" in visibility_line
    assert "увеличилось" not in visibility_line
    assert "увеличилась на 9%" in visibility_line


def test_distribution_summary_uses_equal_centered_label_cells():
    document = Document()
    exporting._render_distribution_cards_table(
        document,
        _distribution(),
        100,
        engine="yandex",
    )
    outer = document.tables[0]
    label_cells = []
    for row in outer.rows:
        for cell in row.cells:
            if cell.tables:
                label_cells.append(cell.tables[0].rows[0].cells[0])
    assert len(label_cells) == 6
    assert len({cell.width for cell in label_cells}) == 1
    assert all(cell.paragraphs[0].alignment is not None for cell in label_cells)


def test_polish_assets_define_modal_visibility_editor_and_fixed_layout(settings):
    static_root = Path(settings.BASE_DIR) / "static" / "reports"
    javascript = (static_root / "report-polish.js").read_text()
    css = (static_root / "report-polish.css").read_text()

    assert "Скорректировать таблицы динамики" in javascript
    assert "<th>Видимость</th>" in javascript
    assert "flushManualSave" in javascript
    assert "event.key === 'Escape'" in javascript
    assert "event.target === backdrop" in javascript
    assert "focusableSelector" in javascript
    assert "topvisor_manual_rows" in javascript
    assert "+ Добавить строку" in javascript
    assert "<form" not in javascript
    assert ".report-modal-backdrop" in css
    assert "details[data-topvisor-manual-editor]{display:none}" in css
    assert ".report-config-grid{grid-template-columns:repeat(3,minmax(0,1fr))}" in css
    assert ".position-distribution__range{width:58px;min-width:58px;max-width:58px" in css
