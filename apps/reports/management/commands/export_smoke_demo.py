"""Run the reproducible offline demo and verify the real export toolchain."""

import shutil
import subprocess
from pathlib import Path

from django.core.management.base import BaseCommand, CommandError
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from apps.reports.demo import create_demo_project
from apps.reports.exporting import generate_artifact


class Command(BaseCommand):
    help = "Create the MVP-1 demo, export twice from its snapshot, and smoke-test the files."

    def add_arguments(self, parser):
        parser.add_argument("--output", default="demo-artifacts")

    def handle(self, *args, **options):
        if not shutil.which("libreoffice") or not shutil.which("pdftoppm"):
            raise CommandError("LibreOffice and pdftoppm are required")
        output = Path(options["output"]).resolve()
        output.mkdir(parents=True, exist_ok=True)
        _, _, version = create_demo_project()
        checksum = version.snapshot.checksum
        baseline = None
        generated = {}
        for run in (1, 2):
            run_files = {}
            for kind in ("docx", "pdf", "xlsx"):
                artifact = generate_artifact(version=version, artifact_type=kind, is_draft=False)
                destination = output / f"seo-demo{'-repeat' if run == 2 else ''}.{kind}"
                with artifact.file.open("rb") as source, destination.open("wb") as target:
                    shutil.copyfileobj(source, target)
                run_files[kind] = destination
            content = self._content_signature(run_files)
            if baseline is None:
                baseline = content
                generated = run_files
            elif content != baseline:
                raise CommandError("Repeated exports from one snapshot have different content")
            version.snapshot.refresh_from_db()
            if version.snapshot.checksum != checksum:
                raise CommandError("Snapshot changed during export")

        reader = PdfReader(generated["pdf"])
        if not generated["pdf"].read_bytes().startswith(b"%PDF-") or not reader.pages:
            raise CommandError("PDF signature or page count is invalid")
        a4_sizes = ((595.28, 841.89), (841.89, 595.28))
        invalid_sizes = []
        for number, page in enumerate(reader.pages, 1):
            width, height = float(page.mediabox.width), float(page.mediabox.height)
            if not any(
                abs(width - expected_width) <= 3 and abs(height - expected_height) <= 3
                for expected_width, expected_height in a4_sizes
            ):
                invalid_sizes.append((number, round(width, 2), round(height, 2)))
        if invalid_sizes:
            raise CommandError(f"PDF contains non-A4 pages: {invalid_sizes}")
        sparse = [
            number
            for number, page in enumerate(reader.pages, 1)
            if number != 1
            and len((page.extract_text() or "").strip()) < 120
            and not list(page.images)
        ]
        if sparse:
            raise CommandError(f"PDF contains unexpectedly empty pages: {sparse}")
        subprocess.run(
            ["pdftoppm", "-png", str(generated["pdf"]), str(output / "seo-demo-page")],
            check=True,
            timeout=120,
        )
        previews = sorted(output.glob("seo-demo-page-*.png"))
        if len(previews) != len(reader.pages) or any(path.stat().st_size == 0 for path in previews):
            raise CommandError("PNG previews were not created for every PDF page")
        self.stdout.write(
            self.style.SUCCESS(
                f"MVP-1 smoke passed: snapshot={checksum}; {len(reader.pages)} pages; "
                f"almost-empty pages: none; artifacts={output}"
            )
        )

    @staticmethod
    def _content_signature(files):
        document = Document(files["docx"])
        if not document.inline_shapes:
            raise CommandError("DOCX has no charts")
        docx_text = tuple(p.text for p in document.paragraphs) + tuple(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        workbook = load_workbook(files["xlsx"], read_only=True, data_only=True)
        xlsx_values = tuple(
            (sheet.title, tuple(tuple(row) for row in sheet.iter_rows(values_only=True)))
            for sheet in workbook.worksheets
        )
        workbook.close()
        pdf_text = tuple(
            (page.extract_text() or "").strip() for page in PdfReader(files["pdf"]).pages
        )
        return docx_text, xlsx_values, pdf_text
