"""Offline renderers whose business data comes only from a frozen report snapshot."""

import base64
import hashlib
import io
import math
import re
import shutil
import subprocess
import tempfile
import textwrap
from datetime import date
from decimal import ROUND_DOWN, ROUND_HALF_UP, Decimal, InvalidOperation
from fnmatch import fnmatchcase
from html.parser import HTMLParser
from pathlib import Path
from urllib.parse import urlsplit

import matplotlib
from django.conf import settings
from django.core.files.base import ContentFile
from django.utils import timezone
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

matplotlib.use("Agg")
from matplotlib import dates as mdates  # noqa: E402
from matplotlib import pyplot as plt  # noqa: E402
from matplotlib.lines import Line2D  # noqa: E402
from matplotlib.patches import Circle, FancyBboxPatch, Patch, PathPatch, Rectangle  # noqa: E402
from matplotlib.path import Path as MplPath  # noqa: E402
from matplotlib.ticker import FuncFormatter, MaxNLocator  # noqa: E402

from .models import GeneratedArtifact, NarrativeBlock, ReportDatasetSnapshot, ValidationIssue
from .narratives import TOP_SECTION_RANGES, section_enabled
from .validation import get_publication_readiness

GENERATOR_VERSION = "mvp1.8-2026-09-04"
MIMES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
TITLES = {
    "visibility": "Видимость",
    "position_distribution": "Распределение позиций",
    "top_5": "Запросы в TOP-5",
    "top_10": "Запросы в TOP-10",
    "top_20": "Запросы в TOP-20",
    "top_11_30": "Запросы в TOP-11–30",
    "top_30": "Запросы в TOP-30",
    "top_11_20": "TOP-11–20",
    "position_dynamics": "Динамика позиций по месяцам",
    "traffic": "Трафик",
    "traffic_sources": "Источники трафика",
    "clicks_impressions": "Клики и показы",
    "ctr": "CTR",
    "indexing": "Индексация",
    "iks": "ИКС",
    "geography": "География посетителей",
    "completed_work": "Выполненные работы",
}
MONTHS = (
    "январь",
    "февраль",
    "март",
    "апрель",
    "май",
    "июнь",
    "июль",
    "август",
    "сентябрь",
    "октябрь",
    "ноябрь",
    "декабрь",
)
MONTHS_GENITIVE = (
    "января",
    "февраля",
    "марта",
    "апреля",
    "мая",
    "июня",
    "июля",
    "августа",
    "сентября",
    "октября",
    "ноября",
    "декабря",
)
ENGINE_LABELS = {"google": "Google", "yandex": "Яндекс"}
METRIC_LABELS = {
    "visits": "Визиты",
    "users": "Посетители",
    "new_users": "Новые посетители",
    "bounce_rate": "Показатель отказов",
    "page_depth": "Глубина просмотра",
    "avg_visit_duration_seconds": "Средняя длительность визита",
    "search_clicks": "Клики",
    "search_impressions": "Показы",
    "search_ctr": "CTR",
    "average_position": "Средняя позиция",
    "indexed_pages": "Индексируемые страницы",
    "iks": "ИКС",
    "quality_index": "ИКС",
    "geography_moscow_visits": "Москва",
    "geography_saint_petersburg_visits": "Санкт-Петербург",
    "geography_undefined_visits": "Не определено",
    "geography_area_undefined_visits": "Область не определена",
}
TOPVISOR_COLORS = {
    "visibility": "#66CC00",
    "1-3": "#2D9CDB",
    "1-10": "#15966F",
    "11-20": "#18BDA3",
    "11-30": "#18BDA3",
    "31-50": "#8CD993",
    "51-100": "#B0C8C8",
    "101+": "#FFB820",
}
WEBMASTER_COLORS = ("#FFD54A", "#8BCB55", "#FF6657", "#84BFE0")
INDEXING_GROUP_COLORS = (
    "#00B945",
    "#20B7A5",
    "#5B8FF9",
    "#8B6FD6",
    "#FF9D4D",
    "#D65DB1",
    "#8CBF26",
    "#38A3A5",
)
METRIKA_COLORS = ("#7A45E5", "#FF3399", "#0FBDA0", "#3388FF", "#FFB851")
CHART_FONT = "sans-serif"


class ExportBlocked(Exception):
    pass


def _clean(value):
    rendered = "" if value is None else str(value)
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", rendered)


def _excel_safe(value):
    """Prevent spreadsheet programs from interpreting untrusted strings as formulas."""
    value = _clean(value)
    return "'" + value if value.startswith(("=", "+", "-", "@")) else value


def _number(value, suffix="", *, decimal_places=None):
    if value is None:
        return "Данные недоступны"
    try:
        number = Decimal(str(value))
        if decimal_places is None:
            rendered = format(number.normalize(), "f")
        else:
            quantum = Decimal("1").scaleb(-decimal_places)
            rendered = format(
                number.quantize(quantum, rounding=ROUND_HALF_UP), f".{decimal_places}f"
            )
        rendered = rendered.replace(".", ",")
    except (InvalidOperation, ValueError):
        rendered = _clean(value)
    return rendered + suffix


def _month(value):
    parsed = date.fromisoformat(str(value)[:10])
    return f"{MONTHS[parsed.month - 1]} {parsed.year}"


def _month_short(value):
    parsed = date.fromisoformat(str(value)[:10])
    return MONTHS[parsed.month - 1]


def _set_cell_width(cell, width):
    cell.width = Cm(width)
    tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
    tc_width.set(qn("w:w"), str(int(width * 567)))
    tc_width.set(qn("w:type"), "dxa")


def _set_cell_margins(cell, *, top=45, bottom=45, left=70, right=70):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _repeat_header(row):
    props = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    props.append(element)


def _prevent_row_split(row):
    props = row._tr.get_or_add_trPr()
    props.append(OxmlElement("w:cantSplit"))


def _keep_small_table_together(table):
    """Keep compact tables on one page instead of orphaning their final row."""
    if len(table.rows) > 8:
        return
    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True


def _shade_cell(cell, color):
    if not color:
        return
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color.removeprefix("#"))


def _set_table_borders(table, color, *, size="4"):
    for row in table.rows:
        for cell in row.cells:
            properties = cell._tc.get_or_add_tcPr()
            borders = properties.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                properties.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                element = borders.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    borders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), size)
                element.set(qn("w:color"), color)


def _set_cell_bottom_border(cell, color, *, size="18"):
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:color"), color)


def _table(doc, headers, rows, widths=None, *, header_fill=None, cell_fills=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Report Table"
    table.autofit = False
    widths = widths or [16 / len(headers)] * len(headers)
    layout = table._tbl.tblPr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table._tbl.tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths, strict=True):
        grid_column.w = Cm(width)
    for cell, value, width in zip(table.rows[0].cells, headers, widths, strict=True):
        cell.text = _clean(value)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_width(cell, width)
        _set_cell_margins(cell)
        _shade_cell(cell, header_fill)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
                run.font.size = Pt(11)
                run.font.bold = True
    _repeat_header(table.rows[0])
    _prevent_row_split(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        _prevent_row_split(table.rows[-1])
        for column_index, (cell, value, width) in enumerate(
            zip(cells, values, widths, strict=True)
        ):
            cell.text = _clean(value if value is not None and value != "" else "—")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_width(cell, width)
            _set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
                    run.font.size = Pt(11)
                    run.font.bold = False
            if cell_fills and row_index < len(cell_fills):
                fills = cell_fills[row_index]
                if column_index < len(fills):
                    _shade_cell(cell, fills[column_index])
    _keep_small_table_together(table)
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(3)
    gap.paragraph_format.line_spacing = Pt(2)
    gap.add_run("\u200b").font.size = Pt(2)
    return table


def _save_figure(figure):
    output = io.BytesIO()
    figure.savefig(
        output,
        format="png",
        dpi=300,
        facecolor="white",
        metadata={"Software": GENERATOR_VERSION},
    )
    plt.close(figure)
    output.seek(0)
    return output


def _period_pills_image(items):
    """Render Metrika-like period controls as a sharp, layout-stable image."""
    widths = [0.9 if marker == "swap" else max(2.4, len(label) * 0.105) for label, marker in items]
    total = sum(widths) + 0.18 * (len(widths) - 1)
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 10}):
        figure, axis = plt.subplots(figsize=(total, 0.55), dpi=300, facecolor="white")
        axis.set_xlim(0, total)
        axis.set_ylim(0, 1)
        axis.axis("off")
        x = 0
        for (label, marker), width in zip(items, widths, strict=True):
            fill = "#F1F3F6" if marker != "swap" else "#FFFFFF"
            axis.add_patch(
                FancyBboxPatch(
                    (x, 0.16),
                    width,
                    0.68,
                    boxstyle="round,pad=0.02,rounding_size=0.18",
                    facecolor=fill,
                    edgecolor="#E1E5EA",
                    linewidth=0.8,
                )
            )
            if marker and marker != "swap":
                axis.add_patch(plt.Circle((x + 0.3, 0.5), 0.12, color="#FFFFFF"))
                axis.text(
                    x + 0.3, 0.5, marker, ha="center", va="center", fontsize=7, color="#526174"
                )
                text_x = x + 0.52
            else:
                text_x = x + width / 2
            axis.text(
                text_x,
                0.5,
                label,
                ha="center" if marker == "swap" else "left",
                va="center",
                fontsize=9,
                color="#27303F",
                fontweight="bold" if marker == "swap" else "normal",
            )
            x += width + 0.18
        figure.subplots_adjust(0, 0, 1, 1)
        return _save_figure(figure)


def _webmaster_period_image(period, *, detail=None):
    """Render the date controls used by Yandex Webmaster report pages."""
    date_width = max(3.7, len(period) * 0.088)
    detail_width = 2.55 if detail else 0
    gap = 0.18 if detail else 0
    total = date_width + detail_width + gap
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure, axis = plt.subplots(figsize=(total, 0.56), dpi=300, facecolor="white")
        axis.set_xlim(0, total)
        axis.set_ylim(0, 1)
        axis.axis("off")

        def field(x, width):
            axis.add_patch(
                FancyBboxPatch(
                    (x, 0.12),
                    width,
                    0.76,
                    boxstyle="round,pad=0.015,rounding_size=0.24",
                    facecolor="#FFFFFF",
                    edgecolor="#DDE2E8",
                    linewidth=0.9,
                )
            )

        field(0, date_width)
        axis.text(0.18, 0.5, period, ha="left", va="center", fontsize=8.6, color="#252B33")
        # Calendar icon from the same control family, reduced to its essential strokes.
        icon_x = date_width - 0.28
        axis.add_patch(
            FancyBboxPatch(
                (icon_x - 0.095, 0.34),
                0.16,
                0.27,
                boxstyle="round,pad=0.008,rounding_size=0.025",
                fill=False,
                edgecolor="#7A838D",
                linewidth=0.8,
            )
        )
        axis.plot((icon_x - 0.09, icon_x + 0.06), (0.52, 0.52), color="#7A838D", lw=0.7)
        axis.plot((icon_x - 0.05, icon_x - 0.05), (0.59, 0.65), color="#7A838D", lw=0.8)
        axis.plot((icon_x + 0.02, icon_x + 0.02), (0.59, 0.65), color="#7A838D", lw=0.8)
        if detail:
            x = date_width + gap
            field(x, detail_width)
            axis.text(
                x + 0.18,
                0.5,
                f"Детализация: {detail}",
                ha="left",
                va="center",
                fontsize=8.6,
                color="#252B33",
            )
            axis.plot(
                (x + detail_width - 0.32, x + detail_width - 0.24, x + detail_width - 0.16),
                (0.55, 0.43, 0.55),
                color="#7A838D",
                lw=0.8,
            )
        figure.subplots_adjust(0, 0, 1, 1)
        return _save_figure(figure)


def _metrika_period_image(period, *, detail):
    """Render compact shade buttons used by Yandex Metrika."""
    date_width = max(2.45, len(period) * 0.076)
    detail_label = f"По {detail}"
    detail_width = max(1.55, len(detail_label) * 0.075)
    gap = 0.12
    total = date_width + gap + detail_width
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure, axis = plt.subplots(figsize=(total, 0.48), dpi=300, facecolor="white")
        axis.set_xlim(0, total)
        axis.set_ylim(0, 1)
        axis.axis("off")
        for x, width in ((0, date_width), (date_width + gap, detail_width)):
            axis.add_patch(
                FancyBboxPatch(
                    (x, 0.14),
                    width,
                    0.72,
                    boxstyle="round,pad=0.015,rounding_size=0.18",
                    facecolor="#F1F3F5",
                    edgecolor="#F1F3F5",
                    linewidth=0.8,
                )
            )
        axis.text(0.17, 0.5, period, ha="left", va="center", fontsize=8.2, color="#252B33")
        axis.text(
            date_width + gap + 0.17,
            0.5,
            detail_label,
            ha="left",
            va="center",
            fontsize=8.2,
            color="#252B33",
        )
        for x in (date_width - 0.2, total - 0.2):
            axis.plot((x - 0.05, x, x + 0.05), (0.54, 0.44, 0.54), color="#66717D", lw=0.75)
        figure.subplots_adjust(0, 0, 1, 1)
        return _save_figure(figure)


def _style_axis(axis, *, grid_axis="both"):
    axis.set_axisbelow(True)
    axis.grid(axis=grid_axis, color="#E7EBEF", linewidth=0.75)
    axis.tick_params(colors="#8B98A7", labelsize=8, length=0)
    for spine in axis.spines.values():
        spine.set_visible(False)


def _plot_smooth_line(axis, x_values, y_values, *, color, linewidth=1.8, label=None):
    """Draw a visually smooth path through real points without creating a data series."""
    points = [(float(x), float(y)) for x, y in zip(x_values, y_values, strict=True)]
    if len(points) < 3:
        return axis.plot(
            x_values,
            y_values,
            color=color,
            linewidth=linewidth,
            solid_capstyle="round",
            solid_joinstyle="round",
            label=label,
        )[0]
    vertices = [points[0]]
    codes = [MplPath.MOVETO]
    for index in range(len(points) - 1):
        current = points[index]
        following = points[index + 1]
        before = points[index - 1] if index else current
        after = points[index + 2] if index + 2 < len(points) else following
        control_1 = (
            current[0] + (following[0] - before[0]) / 6,
            current[1] + (following[1] - before[1]) / 6,
        )
        control_2 = (
            following[0] - (after[0] - current[0]) / 6,
            following[1] - (after[1] - current[1]) / 6,
        )
        lower = min(current[1], following[1])
        upper = max(current[1], following[1])
        control_1 = (control_1[0], min(upper, max(lower, control_1[1])))
        control_2 = (control_2[0], min(upper, max(lower, control_2[1])))
        vertices.extend((control_1, control_2, following))
        codes.extend((MplPath.CURVE4, MplPath.CURVE4, MplPath.CURVE4))
    patch = PathPatch(
        MplPath(vertices, codes),
        fill=False,
        edgecolor=color,
        linewidth=linewidth,
        capstyle="round",
        joinstyle="round",
        label=label,
    )
    axis.add_patch(patch)
    return patch


def _date_ticks(labels, maximum=7):
    if len(labels) <= maximum:
        return list(range(len(labels))), labels
    indexes = sorted({round(index * (len(labels) - 1) / (maximum - 1)) for index in range(maximum)})
    return indexes, [labels[index] for index in indexes]


def _visibility_chart(points, title=None):
    useful = [(month, value) for month, value in points if value is not None]
    if not useful:
        return None
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure = plt.figure(figsize=(7.2, 3.0), dpi=150, facecolor="white")
        grid = figure.add_gridspec(1, 2, width_ratios=(3.5, 1.2), wspace=0.22)
        axis = figure.add_subplot(grid[0, 0])
        labels = [_date_label(month) for month, _value in useful]
        values = [float(value) for _month, value in useful]
        x_values = list(range(len(values)))
        green = TOPVISOR_COLORS["visibility"]
        _plot_smooth_line(axis, x_values, values, color=green, linewidth=1.8)
        axis.scatter(x_values, values, color=green, s=13, zorder=3)
        axis.fill_between(x_values, values, color=green, alpha=0.07)
        ticks, tick_labels = _date_ticks(labels)
        axis.set_xticks(ticks, tick_labels)
        axis.set_ylim(0, 100)
        axis.set_yticks((0, 50, 100), labels=("0%", "50%", "100%"))
        _style_axis(axis)
        donut = figure.add_subplot(grid[0, 1])
        current = max(0, min(100, values[-1]))
        donut.pie(
            [current, 100 - current],
            startangle=78,
            counterclock=False,
            colors=(green, "#F0F0F0"),
            wedgeprops={"width": 0.18, "edgecolor": "white"},
        )
        donut.text(
            0,
            0,
            f"{current:.0f}%",
            ha="center",
            va="center",
            fontsize=12,
            color=green,
        )
        donut.set_aspect("equal")
        figure.subplots_adjust(left=0.08, right=0.99, top=0.98, bottom=0.16)
        return _save_figure(figure)


def _topvisor_buckets(distribution, depth):
    manual = distribution.get("manual_buckets") or {}
    if manual:
        labels = ["1-3", "1-10", "11-30" if depth >= 30 else "11-20"]
        return [
            {
                "label": label,
                "count": (manual.get(label) or {}).get("count", 0),
                "share": (manual.get(label) or {}).get("share"),
            }
            for label in labels
        ]
    ranges = distribution.get("ranges") or {}
    total = distribution.get("total") or 0
    buckets = [
        ("1-3", ranges.get("1-3", 0)),
        ("1-10", distribution.get("top_10", 0)),
    ]
    if depth >= 30:
        buckets.append(("11-30", (ranges.get("11-20") or 0) + (ranges.get("21-30") or 0)))
    elif depth >= 20:
        buckets.append(("11-20", ranges.get("11-20", 0)))
    if depth >= 50:
        buckets.append(("31-50", ranges.get("31-50", 0)))
    if depth >= 100:
        buckets.append(("51-100", ranges.get("51-100", 0)))
        buckets.append(("101+", max(total - sum(ranges.values()), 0)))
    return [
        {
            "label": label,
            "count": count or 0,
            "share": Decimal(str(count or 0)) * 100 / Decimal(str(total)) if total else None,
        }
        for label, count in buckets
    ]


def _distribution_chart(history, depth):
    useful_rows = [row for row in history if (row.get("distribution") or {}).get("total")]
    if not useful_rows:
        return None
    bucket_rows = [_topvisor_buckets(row.get("distribution") or {}, depth) for row in useful_rows]
    labels = [_date_label(row.get("month")) for row in useful_rows]
    x_values = list(range(len(labels)))
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure, axis = plt.subplots(figsize=(7.2, 3.35), dpi=150, facecolor="white")
        bucket_names = [bucket["label"] for bucket in bucket_rows[0]]
        rows_by_name = [{bucket["label"]: bucket for bucket in buckets} for buckets in bucket_rows]
        for name in bucket_names:
            values = [float(row.get(name, {}).get("share") or 0) for row in rows_by_name]
            color = TOPVISOR_COLORS[name]
            _plot_smooth_line(axis, x_values, values, color=color, linewidth=1.7, label=name)
            axis.scatter(x_values, values, color=color, s=12, zorder=3)
            axis.fill_between(x_values, values, color=color, alpha=0.055)
        ticks, tick_labels = _date_ticks(labels)
        axis.set_xticks(ticks, tick_labels)
        top = max(float(bucket["share"] or 0) for row in bucket_rows for bucket in row)
        axis.set_ylim(0, max(10, math.ceil(top / 10) * 10 + 2))
        axis.yaxis.set_major_formatter(FuncFormatter(lambda value, _pos: f"{value:.0f}%"))
        _style_axis(axis)
        axis.legend(
            loc="upper center",
            bbox_to_anchor=(0.5, -0.14),
            ncol=min(6, len(bucket_names)),
            frameon=False,
            fontsize=8,
        )
        figure.subplots_adjust(left=0.14, right=0.98, top=0.98, bottom=0.25)
        return _save_figure(figure)


def _distribution_cards(distribution, depth, *, engine="yandex"):
    buckets = _topvisor_buckets(distribution, 20 if engine == "google" else depth)
    if engine == "google":
        buckets = [bucket for bucket in buckets if bucket["label"] in {"1-3", "1-10", "11-20"}]
    if not buckets:
        return None
    columns = 1 if engine == "google" else 2
    row_count = math.ceil(len(buckets) / columns)
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 12}):
        figure, axis = plt.subplots(figsize=(5.4, 0.9 * row_count), dpi=240, facecolor="white")
        axis.set_xlim(0, columns)
        axis.set_ylim(0, row_count)
        axis.axis("off")
        for index, bucket in enumerate(buckets):
            column = index // row_count
            row = index % row_count
            x = column + 0.03
            y = row_count - row - 0.9
            axis.add_patch(
                FancyBboxPatch(
                    (x, y),
                    0.91,
                    0.78,
                    boxstyle="round,pad=0.02,rounding_size=0.08",
                    linewidth=0,
                    facecolor="#F0F2F5",
                )
            )
            axis.add_patch(
                FancyBboxPatch(
                    (x + 0.04, y + 0.12),
                    0.28,
                    0.52,
                    boxstyle="round,pad=0.02,rounding_size=0.05",
                    linewidth=0,
                    facecolor=TOPVISOR_COLORS[bucket["label"]],
                )
            )
            axis.text(
                x + 0.18,
                y + 0.38,
                bucket["label"],
                ha="center",
                va="center",
                color="white",
                fontsize=12,
                fontweight="bold",
            )
            axis.text(
                x + 0.43,
                y + 0.38,
                _number(bucket["share"], "%", decimal_places=0),
                ha="left",
                va="center",
                color="#91A0AF",
                fontsize=12,
            )
            axis.text(
                x + 0.82,
                y + 0.38,
                _number(bucket["count"]),
                ha="right",
                va="center",
                color="#5D6875",
                fontsize=12,
            )
        figure.subplots_adjust(left=0, right=1, top=1, bottom=0)
        return _save_figure(figure)


def _shade_run(run, color):
    properties = run._element.get_or_add_rPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _render_distribution_cards_table(doc, distribution, depth, *, engine="yandex"):
    """Render the range summary as native Word content so PDF text stays sharp."""
    buckets = _topvisor_buckets(distribution, 20 if engine == "google" else depth)
    if engine == "google":
        buckets = [bucket for bucket in buckets if bucket["label"] in {"1-3", "1-10", "11-20"}]
    if not buckets:
        return None
    columns = 1 if engine == "google" else 2
    row_count = math.ceil(len(buckets) / columns)
    table = doc.add_table(rows=row_count, cols=columns)
    table.autofit = False
    for row in table.rows:
        _prevent_row_split(row)
        for cell in row.cells:
            _set_cell_width(cell, 9.25 if columns == 2 else 9.6)
            _set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            _shade_cell(cell, "F0F2F5")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
    for index, bucket in enumerate(buckets):
        column = index // row_count
        row_index = index % row_count
        paragraph = table.rows[row_index].cells[column].paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        label = paragraph.add_run(f"  {bucket['label']}  ")
        _style_run(label, size=12, color="FFFFFF", bold=True)
        _shade_run(label, TOPVISOR_COLORS[bucket["label"]].removeprefix("#"))
        share = paragraph.add_run(f"    {_number(bucket.get('share'), '%', decimal_places=0)}    ")
        _style_run(share, size=12, color="8491A5")
        count = paragraph.add_run(_number(bucket.get("count"), decimal_places=0))
        _style_run(count, size=12, color="3D4655")
    _set_table_borders(table, "FFFFFF", size="12")
    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(2)
    return table


def _metric_source(payload, source_name):
    return payload.get("calculated", {}).get("sources", {}).get("sources", {}).get(source_name, {})


def _current_position_source(payload, segment):
    configuration_id = segment.get("configuration_id")
    candidates = [
        s
        for s in payload.get("ranking_sources", [])
        if s.get("search_engine") == segment.get("search_engine")
        and s.get("region") == segment.get("region")
        and (not configuration_id or s.get("configuration_id") == configuration_id)
    ]
    return max(
        candidates, key=lambda item: (item.get("date") or "", item.get("id") or ""), default=None
    )


def _position_rows(source, start=None, end=None, *, show_urls=True):
    if not source:
        return []
    rows = []
    depth = source.get("ranking_depth") or 0
    for item in source.get("positions", []):
        position = item.get("position")
        if position is not None and position > depth:
            position = None
        if start is not None and (position is None or position < start or position > end):
            continue
        values = [
            item.get("query"),
            item.get("frequency"),
            position,
            item.get("group"),
            item.get("target_url"),
        ]
        rows.append(tuple(values if show_urls else values[:-1]))
    return rows


def _add_report_picture(doc, picture, *, width=18.2):
    if picture is None:
        doc.add_paragraph("Данные недоступны для построения графика.", style="Data Missing")
        return False
    doc.add_picture(picture, width=Cm(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def _position_fill(position):
    if position is None:
        return None
    if position <= 3:
        return "55C98A"
    if position <= 5:
        return "8FDDB0"
    if position <= 10:
        return "C9EFD7"
    if position <= 20:
        return "E2F5E9"
    return "F1F9F4"


def _render_position_table(doc, payload, segment, start, end):
    source = _current_position_source(payload, segment)
    rows = _position_rows(source, start, end, show_urls=False)
    engine = segment.get("search_engine")
    engine_header = "Yandex" if engine == "yandex" else "Google"
    if not rows:
        doc.add_paragraph("Запросы в выбранном диапазоне отсутствуют.", style="Data Missing")
        return False
    widths = [7.74, 1.50, 1.56, 8.25] if engine == "yandex" else [8.63, 1.17, 1.52, 7.42]
    cell_fills = []
    cluster_order = {}
    for row in rows:
        cluster = str(row[3] or "Без группы")
        cluster_order.setdefault(cluster, len(cluster_order))
        cluster_fill = "F7F8FA" if cluster_order[cluster] % 2 == 0 else None
        cell_fills.append((cluster_fill, cluster_fill, _position_fill(row[2]), cluster_fill))
    table = _table(
        doc,
        ("Запросы", "WS", engine_header, "Имя группы"),
        rows,
        widths,
        header_fill="EEF1F2",
        cell_fills=cell_fills,
    )
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for row in table.rows[1:]:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
        for column in (1, 2):
            row.cells[column].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def _render_daily_position_table(doc, payload, segment):
    sources = [
        source
        for source in payload.get("ranking_sources", [])
        if str(source.get("configuration_id")) == str(segment.get("configuration_id"))
    ]
    sources.sort(key=lambda item: str(item.get("date")))
    if not sources:
        return
    dates = [str(source.get("date"))[:10] for source in sources]
    by_date = {
        day: {
            row.get("normalized_query") or row.get("query"): row
            for row in source.get("positions", [])
        }
        for day, source in zip(dates, sources, strict=True)
    }
    latest = sources[-1].get("positions", [])
    query_rows = []
    current_group = None
    for row in sorted(latest, key=lambda item: (item.get("group") or "", item.get("query") or "")):
        group = row.get("group") or "Без группы"
        if group != current_group:
            query_rows.append((f"Группа: {group}", "", *([""] * len(dates))))
            current_group = group
        key = row.get("normalized_query") or row.get("query")
        query_rows.append(
            (
                row.get("query"),
                _number(row.get("frequency"), decimal_places=0),
                *[
                    _number((by_date[day].get(key) or {}).get("position"), decimal_places=0)
                    if (by_date[day].get(key) or {}).get("position") is not None
                    else "—"
                    for day in dates
                ],
            )
        )
    widths = [8.0, 2.2, *([max(1.6, 8.2 / len(dates))] * len(dates))]
    table = _table(
        doc,
        (
            "Ключевое слово",
            "Частота",
            *[date.fromisoformat(day).strftime("%d.%m") for day in dates],
        ),
        query_rows,
        widths,
        header_fill="F7F7F7",
    )
    for row_index, values in enumerate(query_rows, start=1):
        if str(values[0]).startswith("Группа: "):
            for cell in table.rows[row_index].cells:
                _shade_cell(cell, "F3F5F7")
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
            continue
        for column in range(2, len(table.rows[row_index].cells)):
            cell = table.rows[row_index].cells[column]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            value = values[column]
            if value != "—":
                _shade_cell(cell, _position_fill(int(str(value).replace(",", ".").split(".")[0])))
    _set_table_borders(table, "E2E4E8", size="3")


def _monthly_boundary_points(points, date_getter):
    points_by_month = {}
    for order, point in enumerate(points):
        point_date = date.fromisoformat(str(date_getter(point))[:10])
        points_by_month.setdefault((point_date.year, point_date.month), []).append(
            (point_date, order, point)
        )
    month_keys = sorted(points_by_month)
    monthly_points = []
    for index, month_key in enumerate(month_keys):
        dated_points = sorted(points_by_month[month_key], key=lambda item: (item[0], item[1]))
        # The first month is the comparison baseline; later months use their latest point.
        monthly_points.append(dated_points[0 if index == 0 else -1][2])
    return monthly_points


def _monthly_topvisor_rows(segment):
    rows = []
    depth = segment.get("ranking_depth") or 0
    monthly_points = _monthly_boundary_points(
        segment.get("three_month_series") or [], lambda point: point.get("month")
    )

    for point in monthly_points:
        distribution = point.get("distribution") or {}
        buckets = {item["label"]: item for item in _topvisor_buckets(distribution, depth)}

        def bucket(label, bucket_values=buckets):
            item = bucket_values.get(label) or {}
            return (
                f"{_number(item.get('share'), '%', decimal_places=0)} "
                f"({_number(item.get('count'))})"
            )

        final_label = "11-30" if depth >= 30 else "11-20"
        rows.append(
            (
                _month(point.get("month")).capitalize(),
                _number(point.get("visibility"), "%", decimal_places=0),
                bucket("1-3"),
                bucket("1-10"),
                bucket(final_label),
            )
        )
    return rows


def _render_monthly_topvisor_table(doc, segment, *, show_visibility=True):
    final_label = "11-30" if (segment.get("ranking_depth") or 0) >= 30 else "11-20"
    rows = _monthly_topvisor_rows(segment)
    if not rows:
        doc.add_paragraph("Месячные итоги отсутствуют.", style="Data Missing")
        return
    headers = ("Месяц", "в топ 3", "в топ 10", f"в топ {final_label}")
    rows = [(row[0], *row[2:]) for row in rows]
    widths = [4.6] * 4
    table = _table(doc, headers, rows, widths, header_fill="EEF1F2")
    for row in table.rows:
        for column in range(1, len(row.cells)):
            row.cells[column].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _comparison_phrase(current, previous):
    if current is None or previous is None:
        return "нет базы сравнения"
    delta = Decimal(str(current)) - Decimal(str(previous))
    if delta == 0:
        return "не изменилось"
    direction = "увеличилось" if delta > 0 else "уменьшилось"
    return f"{direction} на {_number(abs(delta), '%', decimal_places=1)}"


def _relative_comparison_phrase(current, previous):
    if current is None or previous is None:
        return "нет базы сравнения"
    current_number = Decimal(str(current))
    previous_number = Decimal(str(previous))
    if current_number == previous_number:
        return "не изменилось"
    if previous_number == 0:
        return "изменение не рассчитывается из-за нулевой базы"
    delta = (current_number - previous_number) / abs(previous_number) * Decimal(100)
    direction = "увеличилось" if delta > 0 else "уменьшилось"
    return f"{direction} на {_number(abs(delta), '%', decimal_places=0)}"


def _render_topvisor_comparison(doc, segment, *, show_visibility=True):
    history = segment.get("three_month_series") or []
    if not history:
        return
    current = history[-1]
    previous = history[-2] if len(history) >= 2 else None
    depth = segment.get("ranking_depth") or 0
    current_buckets = {
        item["label"]: item for item in _topvisor_buckets(current.get("distribution") or {}, depth)
    }
    previous_buckets = (
        {
            item["label"]: item
            for item in _topvisor_buckets(previous.get("distribution") or {}, depth)
        }
        if previous
        else {}
    )
    final_label = "11-30" if depth >= 30 else "11-20"
    doc.add_paragraph("В сравнении с прошлым месяцем доля запросов в топ:")
    for label in ("1-3", "1-10", final_label):
        bucket = current_buckets.get(label) or {}
        now = bucket.get("share")
        before = (previous_buckets.get(label) or {}).get("share")
        rendered_label = {"1-3": "3", "1-10": "10"}.get(label, label)
        paragraph = doc.add_paragraph(
            f"Запросов в топ {rendered_label} — {_number(now, '%', decimal_places=0)} "
            f"({_comparison_phrase(now, before)})."
        )
        paragraph.paragraph_format.space_after = Pt(0)
    if show_visibility:
        now_visibility = current.get("visibility")
        previous_visibility = previous.get("visibility") if previous else None
        engine = ENGINE_LABELS.get(segment.get("search_engine"), "Поиск")
        region = segment.get("region") or "регион не указан"
        paragraph = doc.add_paragraph(
            f"Общая видимость сайта в ПС {engine}.{region} — "
            f"{_number(now_visibility, '%', decimal_places=0)} "
            f"({_comparison_phrase(now_visibility, previous_visibility)})."
        )
        paragraph.paragraph_format.space_after = Pt(0)


def _manual_topvisor_segment(payload, segment):
    rows = payload.get("display_options", {}).get("topvisor_manual_rows") or []
    selected = [
        row
        for row in rows
        if (
            not row.get("configuration_id")
            or str(row.get("configuration_id")) == str(segment.get("configuration_id"))
        )
        and (not row.get("engine") or row.get("engine") == segment.get("search_engine"))
        and (not row.get("region") or row.get("region") == segment.get("region"))
        and row.get("month")
    ]
    if not selected:
        return segment
    depth = segment.get("ranking_depth") or 30
    history_by_month = {
        str(point.get("month"))[:7]: dict(point)
        for point in segment.get("three_month_series") or []
        if point.get("month")
    }
    for row in sorted(selected, key=lambda item: item["month"]):
        top3 = max(int(row.get("top3") or 0), 0)
        top10 = max(int(row.get("top10") or 0), top3)
        top11 = max(int(row.get("top11_30") or 0), 0)
        shares = {
            "1-3": max(0, min(100, float(row.get("top3_percent") or 0))),
            "1-10": max(0, min(100, float(row.get("top10_percent") or 0))),
            "11-30" if depth >= 30 else "11-20": max(
                0, min(100, float(row.get("top11_30_percent") or 0))
            ),
        }
        inferred_totals = [
            round(count * 100 / shares[label])
            for label, count in (("1-3", top3), ("1-10", top10))
            if shares[label] > 0
        ]
        total = max(int(row.get("total") or 0), *(inferred_totals or [0]), top10 + top11)
        ranges = {"1-3": top3, "4-10": max(top10 - top3, 0), "11-20": top11}
        if depth >= 30:
            ranges["21-30"] = 0
        final_label = "11-30" if depth >= 30 else "11-20"
        history_by_month[str(row["month"])[:7]] = {
            "month": row["month"],
            "visibility": row.get("visibility"),
            "distribution": {
                "total": total,
                "top_10": top10,
                "ranges": ranges,
                "manual_buckets": {
                    "1-3": {"count": top3, "share": shares["1-3"]},
                    "1-10": {"count": top10, "share": shares["1-10"]},
                    final_label: {"count": top11, "share": shares[final_label]},
                },
            },
            "ranking_depth": depth,
            "manual_override": True,
        }
    history = [history_by_month[key] for key in sorted(history_by_month)]
    return {
        **segment,
        "three_month_series": history,
        "chart_series": history,
        "distribution": history[-1]["distribution"],
    }


def _top_table_title(segment, start, end):
    range_label = f"TOP-{end}" if start == 1 else f"TOP-{start}–{end}"
    engine = ENGINE_LABELS.get(segment.get("search_engine"), "Поиск")
    region = segment.get("region") or "регион не указан"
    return f"Запросы в {range_label} по {engine}.{region}"


def _render_topvisor_segment(doc, payload, segment, blocks, *, report_url=""):
    segment = _manual_topvisor_segment(payload, segment)
    engine = ENGINE_LABELS.get(segment.get("search_engine"), "Поиск")
    region = segment.get("region") or "регион не указан"
    doc.add_heading(f"{engine}. {region}", level=2)
    history = segment.get("three_month_series") or []
    chart_history = segment.get("chart_series") or history
    depth = segment.get("ranking_depth") or 0
    if section_enabled(payload, "visibility"):
        doc.add_paragraph("График видимости сайта по основным ключевым словам за отчётный период.")
        _add_report_picture(
            doc,
            _visibility_chart(
                [(point.get("month"), point.get("visibility")) for point in chart_history]
            ),
        )
        doc.add_paragraph(
            "Видимость сайта — это доля показов сайта в поисковых системах, которая "
            "зависит от частот и позиций запросов."
        )
        if payload.get("display_options", {}).get("include_visibility_table"):
            doc.add_paragraph("Позиции по выбранным дням", style="Table Heading")
            _render_daily_position_table(doc, payload, segment)
    doc.add_paragraph(
        "График распределения по топам по основным ключевым словам за отчётный период."
    )
    _add_report_picture(doc, _distribution_chart(chart_history, depth), width=18.5)
    doc.add_paragraph(
        "Данная диаграмма не отражает зависимости запросов от частот и отражает только "
        "количество запросов в топ 3, топ 10, топ 30 и пр."
    )
    doc.add_paragraph("Данные по количеству запросов в топ:")
    _render_distribution_cards_table(
        doc,
        segment.get("distribution") or {},
        depth,
        engine=segment.get("search_engine"),
    )
    if section_enabled(payload, "position_dynamics") and payload.get("display_options", {}).get(
        "include_monthly_dynamics_table", True
    ):
        doc.add_paragraph("В динамике по месяцам")
        _render_monthly_topvisor_table(
            doc, segment, show_visibility=section_enabled(payload, "visibility")
        )
    _render_topvisor_comparison(
        doc, segment, show_visibility=section_enabled(payload, "visibility")
    )
    top_mode = payload.get("project", {}).get("top_11_20_mode", "auto")
    for code in ("top_5", "top_10", "top_20", "top_11_30", "top_30", "top_11_20"):
        if not section_enabled(payload, code):
            continue
        if code == "top_11_20" and top_mode == "disabled":
            continue
        start, end = TOP_SECTION_RANGES[code]
        if code == "top_11_20" and top_mode == "auto":
            source = _current_position_source(payload, segment)
            if not _position_rows(source, start, end, show_urls=False):
                continue
        doc.add_paragraph(_top_table_title(segment, start, end), style="Table Heading")
        rendered = _render_position_table(doc, payload, segment, start, end)
        comment = blocks.get(code)
        if rendered and comment and comment != "Данные раздела отсутствуют.":
            doc.add_paragraph(_clean(comment), style="Table Comment")
    if segment.get("search_engine") == "google" and depth:
        doc.add_paragraph(
            f"Глубина проверки Google — TOP-{depth}. Для запросов за пределами "
            "подтверждённой глубины точная позиция не определяется.",
            style="Depth Note",
        )
    if report_url:
        doc.add_paragraph(
            f"Подробный отчёт {engine} · {region}: {_clean(report_url)}",
            style="Compact",
        )


def _topvisor_segment_url(options, segment):
    urls = options.get("topvisor_report_urls") or {}
    configuration = str(
        segment.get("configuration_id") or segment.get("topvisor_configuration_id") or ""
    )
    return urls.get(configuration, "")


def _render_topvisor(doc, payload, segments, blocks):
    if not segments:
        return
    doc.add_heading(
        "1) Видимость сайта в поисковых системах Яндекс и Google по основным ключевым словам",
        level=1,
    )
    options = payload.get("display_options", {})
    report_urls = options.get("topvisor_report_urls") or {}
    legacy_url = options.get("topvisor_report_url", "")
    link_pending = bool(options.get("include_topvisor_report_link") and legacy_url)
    yandex_indexes = [i for i, item in enumerate(segments) if item.get("search_engine") == "yandex"]
    for index, segment in enumerate(segments):
        report_url = (
            _topvisor_segment_url(options, segment)
            if options.get("include_topvisor_report_link") and report_urls
            else legacy_url
            if link_pending
            and (index == yandex_indexes[-1] if yandex_indexes else index == len(segments) - 1)
            else ""
        )
        _render_topvisor_segment(doc, payload, segment, blocks, report_url=report_url)
        link_pending = link_pending and not bool(report_url)


def _metric_series(payload, source, codes):
    facts = _metric_source(payload, source)
    result = []
    for code in codes:
        points = [
            (row.get("month"), row.get("value"))
            for row in facts.get("three_month_series", {}).get(code, [])
        ]
        result.append((METRIC_LABELS.get(code, code), points))
    return result


def _metric_has_data(payload, source, code):
    facts = _metric_source(payload, source)
    change = (facts.get("normalized_changes") or {}).get(code) or {}
    if any(change.get(field) is not None for field in ("current", "previous")):
        return True
    return any(
        point.get("value") is not None
        for point in (facts.get("three_month_series") or {}).get(code, [])
    )


def _period_details(payload, source):
    return _metric_source(payload, source).get("period_details") or []


def _webmaster_chart_details(payload):
    details = _period_details(payload, "yandex_webmaster")
    if not details:
        return []
    if payload.get("display_options", {}).get("webmaster_chart_period") == "selected":
        return details
    return details[-1:]


def _detail_payload(detail):
    return detail.get("payload") or {}


def _daily_rows(details, key):
    rows = []
    for detail in details:
        rows.extend((_detail_payload(detail).get("daily") or {}).get(key, []))
    return sorted(rows, key=lambda item: item.get("date") or "")


def _date_label(value):
    parsed = date.fromisoformat(str(value)[:10])
    return f"{parsed.day} {MONTHS_GENITIVE[parsed.month - 1]}"


def _period_caption(doc, rows, *, detail="дням", provider="metrika", show_detail=True):
    dates = [row.get("date") for row in rows if row.get("date")]
    if not dates:
        return
    start, end = min(dates), max(dates)
    start_date, end_date = date.fromisoformat(start), date.fromisoformat(end)
    if provider == "webmaster":
        period = (
            f"{start_date:%d} {MONTHS[start_date.month - 1]} {start_date.year} — "
            f"{end_date:%d} {MONTHS[end_date.month - 1]} {end_date.year}"
        )
        image = _webmaster_period_image(period, detail=detail if show_detail else None)
        width = 11.5 if show_detail else 7.0
    else:
        period = (
            f"{start_date.day} {MONTHS_GENITIVE[start_date.month - 1]} — "
            f"{end_date.day} {MONTHS_GENITIVE[end_date.month - 1]}"
        )
        image = _metrika_period_image(period, detail=detail)
        width = 7.1
    _add_report_picture(doc, image, width=width)


def _nice_step(values, *, minimum=10):
    if not values:
        return minimum
    spread = max(values) - min(values)
    raw = max(minimum, spread / 4)
    magnitude = 10 ** math.floor(math.log10(raw)) if raw else 1
    normalized = raw / magnitude
    factor = 1 if normalized <= 1 else 2 if normalized <= 2 else 5 if normalized <= 5 else 10
    return max(minimum, factor * magnitude)


def _single_service_chart(
    points, *, title, color, fill=False, suffix="", minimum_step=None, step=False
):
    by_day = {}
    for day, value in points:
        if not day or value is None:
            continue
        by_day[date.fromisoformat(str(day)[:10])] = value
    useful = sorted(by_day.items())
    if not useful:
        return None
    dates = [day for day, _ in useful]
    values = [float(value) for _, value in useful]
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure, axis = plt.subplots(figsize=(7.2, 3.25), dpi=150, facecolor="white")
        if step:
            axis.step(dates, values, where="post", linewidth=2.0, color=color)
        else:
            axis.plot(
                dates,
                values,
                linewidth=2.0,
                color=color,
                solid_capstyle="round",
                solid_joinstyle="round",
            )
        if fill:
            axis.fill_between(
                dates,
                values,
                color=color,
                alpha=0.82,
                step="post" if step else None,
            )
        axis.set_title(title, loc="left", fontsize=13, color="#2F343B", pad=16)
        axis.yaxis.set_major_formatter(FuncFormatter(lambda value, _pos: f"{value:g}{suffix}"))
        if minimum_step:
            step = _nice_step(values, minimum=minimum_step)
            low = max(0, math.floor(min(values) / step) * step - step)
            high = math.ceil(max(values) / step) * step + step
            if high <= low:
                high = low + step * 2
            axis.set_ylim(low, high)
            axis.set_yticks([low + step * index for index in range(int((high - low) / step) + 1)])
        _style_axis(axis)
        axis.xaxis.set_major_locator(mdates.AutoDateLocator(minticks=3, maxticks=7))
        axis.xaxis.set_major_formatter(
            FuncFormatter(lambda value, _pos: _date_label(mdates.num2date(value).date()))
        )
        figure.tight_layout()
        return _save_figure(figure)


def _indexing_group_color(index, row):
    if row.get("path") == "Статус неизвестен":
        return "#F2B51D"
    return INDEXING_GROUP_COLORS[index % len(INDEXING_GROUP_COLORS)]


def _indexing_group_chart(points, distribution):
    rows = [
        row
        for row in ((distribution or {}).get("rows") or [])
        if _decimal_or_none(row.get("count")) not in (None, 0)
    ]
    if not rows:
        return _single_service_chart(
            points,
            title="Страницы в поиске",
            color="#00B945",
            fill=True,
            step=True,
        )
    by_day = {}
    for day, value in points:
        if not day or value is None:
            continue
        by_day[date.fromisoformat(str(day)[:10])] = float(value)
    useful = sorted(by_day.items())
    if not useful:
        return None
    dates = [day for day, _ in useful]
    totals = [value for _, value in useful]
    distribution_total = sum(float(row["count"]) for row in rows)
    if not distribution_total:
        return None
    series = [
        [total * float(row["count"]) / distribution_total for total in totals] for row in rows
    ]
    colors = [_indexing_group_color(index, row) for index, row in enumerate(rows)]
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure, axis = plt.subplots(figsize=(7.2, 3.25), dpi=150, facecolor="white")
        axis.stackplot(dates, *series, colors=colors, alpha=0.9, step="post")
        axis.set_title("Страницы в поиске", loc="left", fontsize=13, color="#2F343B", pad=16)
        axis.set_ylim(bottom=0)
        axis.yaxis.set_major_formatter(FuncFormatter(lambda value, _pos: f"{value:g}"))
        _style_axis(axis)
        axis.xaxis.set_major_locator(mdates.AutoDateLocator(minticks=3, maxticks=7))
        axis.xaxis.set_major_formatter(
            FuncFormatter(lambda value, _pos: _date_label(mdates.num2date(value).date()))
        )
        figure.tight_layout()
        return _save_figure(figure)


def _webmaster_search_chart(payload):
    configs = (
        ("Показы", "shows", "#8BCB55"),
        ("Клики", "clicks", "#F7C933"),
        ("CTR, %", "ctr", "#FF6657"),
        ("Ср. позиция", "average_position", "#84BFE0"),
    )
    rows = _daily_rows(_webmaster_chart_details(payload), "queries")
    if not rows:
        return None
    labels = [_date_label(row["date"]) for row in rows]
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure, axis = plt.subplots(figsize=(7.6, 4.05), dpi=150, facecolor="white")
        for index, row in enumerate(rows):
            if date.fromisoformat(row["date"]).weekday() >= 5:
                axis.axvspan(index - 0.5, index + 0.5, color="#F1F1F1", zorder=0)
        for label, code, color in configs:
            raw = [float(row[code]) if row.get(code) is not None else float("nan") for row in rows]
            finite = [value for value in raw if not math.isnan(value)]
            if not finite:
                continue
            low, high = min(finite), max(finite)
            values = [
                (value - low) / (high - low) if high > low and not math.isnan(value) else 0.5
                for value in raw
            ]
            _plot_smooth_line(
                axis,
                list(range(len(labels))),
                values,
                linewidth=1.4,
                color=color,
                label=label,
            )
        axis.set_xticks(range(len(labels)), labels)
        axis.set_title(
            "Показы, клики, CTR и средняя позиция",
            loc="left",
            fontsize=13,
            color="#2F343B",
            pad=16,
        )
        _style_axis(axis)
        axis.set_yticks([])
        axis.set_xlim(-0.8, len(labels) - 0.2)
        axis.set_ylim(-0.09, 1.09)
        axis.xaxis.set_major_locator(MaxNLocator(8))
        axis.legend(
            loc="upper center",
            bbox_to_anchor=(0.5, -0.13),
            ncol=4,
            frameon=False,
            fontsize=8,
        )
        figure.subplots_adjust(left=0.055, right=0.965, top=0.87, bottom=0.21)
        return _save_figure(figure)


def _change_table(doc, payload, source, codes, *, provider):
    table = _table(
        doc,
        (
            "Показатель",
            "Предыдущий период",
            "Текущий период",
            "Абсолютное изменение",
            "Относительное изменение",
        ),
        _change_rows(payload, source, codes),
        [5.4, 3.0, 3.0, 3.4, 3.7],
        header_fill="F5F7FA",
    )
    for row_index, row in enumerate(table.rows[1:], start=0):
        change = (_metric_source(payload, source).get("normalized_changes") or {}).get(
            codes[row_index]
        ) or {}
        delta = change.get("absolute")
        if delta is None:
            continue
        color = (
            "26A95B"
            if Decimal(str(delta)) > 0
            else "F04444"
            if Decimal(str(delta)) < 0
            else "7A8796"
        )
        for column in (3, 4):
            for run in row.cells[column].paragraphs[0].runs:
                run.font.color.rgb = RGBColor.from_string(color)
    _set_table_borders(table, "E1E7EE", size="3")
    return table


def _latest_webmaster_payload(payload):
    details = _period_details(payload, "yandex_webmaster")
    return _detail_payload(details[-1]) if details else {}


def _decimal_or_none(value):
    try:
        return Decimal(str(value)) if value is not None else None
    except (InvalidOperation, ValueError):
        return None


def _relative_delta(current, previous):
    current = _decimal_or_none(current)
    previous = _decimal_or_none(previous)
    if current is None or previous in (None, 0):
        return None
    return (current - previous) / abs(previous) * Decimal(100)


def _provider_value(code, value):
    if value is None:
        value = 0
    if code in {"ctr", "average_position", "bounce_rate", "conversion_rate"}:
        return _number(value, decimal_places=2)
    return _number(value, decimal_places=0)


def _webmaster_value(code, value):
    if value is None:
        value = 0
    if code in {"ctr", "average_position"}:
        number = Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_DOWN)
        return f"{number:.2f}".replace(".", ",").rstrip("0").rstrip(",")
    return _number(value, decimal_places=0)


def _change_color(current, previous, *, lower_is_better=False):
    current = _decimal_or_none(current)
    previous = _decimal_or_none(previous)
    if current is None or previous is None or current == previous:
        return "7A8796"
    improved = current < previous if lower_is_better else current > previous
    return "26A95B" if improved else "F04444"


def _style_run(run, *, size, color="000000", bold=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _style_provider_table(table):
    for index, cell in enumerate(table.rows[0].cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_margins(cell, top=35, bottom=35, left=55, right=55)
        for paragraph in cell.paragraphs:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                _style_run(run, size=11, bold=True)
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=30, bottom=30, left=55, right=55)
        for paragraph in row.cells[0].paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                _style_run(run, size=11)


def _paired_metric_cell(
    cell,
    code,
    current,
    previous,
    *,
    lower_is_better=False,
    show_change=True,
    color_previous=False,
    difference_only=False,
):
    cell.text = _webmaster_value(code, current)
    current_paragraph = cell.paragraphs[0]
    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    current_paragraph.paragraph_format.space_after = Pt(1)
    _style_run(current_paragraph.runs[0], size=11)
    if previous is None:
        return

    current_number = _decimal_or_none(current)
    previous_number = _decimal_or_none(previous)
    delta = abs((current_number or 0) - (previous_number or 0))
    previous_text = (
        _webmaster_value(code, delta) if difference_only else _webmaster_value(code, previous)
    )
    previous_paragraph = cell.add_paragraph(previous_text)
    previous_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    previous_paragraph.paragraph_format.space_after = Pt(0)
    previous_color = (
        _change_color(current, previous, lower_is_better=lower_is_better)
        if color_previous
        else "7A8796"
    )
    _style_run(previous_paragraph.runs[0], size=8, color=previous_color)

    if not show_change:
        return

    if code in {"shows", "clicks"}:
        relative = (
            delta / abs(previous_number) * Decimal(100)
            if previous_number not in (None, 0)
            else None
        )
        change_text = _number(relative, "%", decimal_places=0)
    else:
        change_text = _number(
            delta,
            "%" if code in {"ctr", "bounce_rate", "conversion_rate"} else "",
            decimal_places=2,
        )
    change = cell.add_paragraph(change_text)
    change.alignment = WD_ALIGN_PARAGRAPH.CENTER
    change.paragraph_format.space_after = Pt(0)
    _style_run(
        change.runs[0],
        size=8,
        color=_change_color(current, previous, lower_is_better=lower_is_better),
    )


def _webmaster_query_summary_table(doc, current, previous):
    headers = ("Группа запросов", "Показы", "Клики", "CTR, %", "Ср. позиция")
    codes = ("shows", "clicks", "ctr", "average_position")
    table = _table(
        doc,
        headers,
        [("Все запросы", "", "", "", "")],
        [6.4, 3.0, 3.0, 2.8, 3.3],
        header_fill="F7F7F7",
    )
    for index, code in enumerate(codes, start=1):
        _paired_metric_cell(
            table.rows[1].cells[index],
            code,
            current.get(code),
            previous.get(code),
            lower_is_better=code == "average_position",
            show_change=False,
            color_previous=True,
            difference_only=True,
        )
    _style_provider_table(table)
    for cell, color in zip(
        table.rows[0].cells[1:], ("8BCB55", "FFD54A", "FF6657", "84BFE0"), strict=True
    ):
        _set_cell_bottom_border(cell, color)
    _set_table_borders(table, "D7DADF", size="3")
    return table


def _webmaster_query_summary_from_changes(payload):
    changes = _metric_source(payload, "yandex_webmaster").get("normalized_changes", {})
    codes = {
        "shows": "search_impressions",
        "clicks": "search_clicks",
        "ctr": "search_ctr",
        "average_position": "average_position",
    }
    current = {
        target: changes[source].get("current")
        for target, source in codes.items()
        if source in changes and changes[source].get("current") is not None
    }
    previous = {
        target: changes[source].get("previous")
        for target, source in codes.items()
        if source in changes and changes[source].get("previous") is not None
    }
    return current, previous


def _webmaster_query_text(current, previous):
    shows = _relative_delta(current.get("shows"), previous.get("shows"))
    clicks = _relative_delta(current.get("clicks"), previous.get("clicks"))
    ctr_current = _decimal_or_none(current.get("ctr"))
    ctr_previous = _decimal_or_none(previous.get("ctr"))
    position_current = _decimal_or_none(current.get("average_position"))
    position_previous = _decimal_or_none(previous.get("average_position"))

    def movement(value):
        if value is None:
            return "не рассчитано"
        return (
            ("увеличилось" if value >= 0 else "снизилось")
            + " на "
            + _number(abs(value), "%", decimal_places=1)
        )

    ctr_delta = (
        ctr_current - ctr_previous if ctr_current is not None and ctr_previous is not None else None
    )
    position_delta = (
        position_current - position_previous
        if position_current is not None and position_previous is not None
        else None
    )
    ctr_text = (
        ("увеличилось" if ctr_delta >= 0 else "снизилось")
        + " на "
        + _number(abs(ctr_delta), "%", decimal_places=2)
        if ctr_delta is not None
        else "не рассчитано"
    )
    position_text = (
        ("улучшении" if position_delta < 0 else "снижении")
        + " средней позиции по всем запросам на "
        + _number(abs(position_delta), decimal_places=2)
        + " пункта"
        if position_delta is not None
        else "недоступной динамике средней позиции"
    )
    return (
        "По сравнению с предыдущим аналогичным периодом (сравнивается с периодом, "
        "равным по количеству дней с текущим) общее количество показов по всем запросам "
        f"в поисковой системе Яндекс {movement(shows)}. Количество кликов {movement(clicks)}. "
        f"Значение CTR {ctr_text} при {position_text}."
    )


def _webmaster_popular_table(doc, current_rows, previous_rows):
    previous = {row.get("query", "").casefold(): row for row in previous_rows}
    table = _table(
        doc,
        ("Запрос", "Показы", "Клики", "CTR, %", "Ср. позиция"),
        [(row.get("query"), "", "", "", "") for row in current_rows[:20]],
        [8.7, 2.7, 2.4, 2.4, 2.8],
        header_fill="F7F7F7",
    )
    for row_index, current in enumerate(current_rows[:20], start=1):
        before = previous.get(current.get("query", "").casefold(), {})
        for column, code in enumerate(("shows", "clicks", "ctr", "average_position"), start=1):
            _paired_metric_cell(
                table.rows[row_index].cells[column],
                code,
                current.get(code),
                before.get(code),
                lower_is_better=code == "average_position",
                show_change=False,
                color_previous=True,
                difference_only=True,
            )
    _style_provider_table(table)
    for cell, color in zip(
        table.rows[0].cells[1:], ("8BCB55", "FFD54A", "FF6657", "84BFE0"), strict=True
    ):
        _set_cell_bottom_border(cell, color)
    _set_table_borders(table, "D7DADF", size="3")
    return table


def _webmaster_popular_analysis(payload, current_rows, previous_rows):
    top = sorted(
        current_rows, key=lambda row: _decimal_or_none(row.get("clicks")) or 0, reverse=True
    )[:10]
    informational_words = ("как", "что", "почему", "симптом", "лечение", "причин", "можно ли")
    informational = sum(
        any(word in str(row.get("query") or "").casefold() for word in informational_words)
        for row in top
    )
    kind = "информационные" if informational >= max(1, len(top) / 2) else "коммерческие"
    domain = str(payload.get("project", {}).get("normalized_domain") or "").split(".")[0].casefold()
    brand_patterns = [
        str(rule.get("pattern") or "").casefold()
        for rule in payload.get("project", {}).get("brand_rules", [])
        if rule.get("active")
    ]
    needles = [value for value in [domain, *brand_patterns] if value]

    def brand_totals(rows):
        selected = [
            row
            for row in rows
            if any(needle in str(row.get("query") or "").casefold() for needle in needles)
        ]
        return (
            sum((_decimal_or_none(row.get("shows")) or 0 for row in selected), Decimal(0)),
            sum((_decimal_or_none(row.get("clicks")) or 0 for row in selected), Decimal(0)),
        )

    current_shows, current_clicks = brand_totals(current_rows)
    previous_shows, previous_clicks = brand_totals(previous_rows)
    parts = [f"Среди самых кликабельных запросов преобладают {kind}."]
    changes = []
    for label, current, previous in (
        ("Показы по брендовым запросам", current_shows, previous_shows),
        ("переходы", current_clicks, previous_clicks),
    ):
        delta = _relative_delta(current, previous)
        if delta in (None, 0):
            continue
        direction = "увеличилось" if delta > 0 else "уменьшилось"
        changes.append(f"{label}: {direction} на {_number(abs(delta), '%', decimal_places=1)}")
    if changes:
        parts.append("; ".join(changes) + ".")
    return " ".join(parts)


def _add_uploaded_picture(doc, payload):
    image = payload.get("display_options", {}).get("webmaster_queries_screenshot") or {}
    if not image.get("data"):
        return False
    try:
        data = io.BytesIO(base64.b64decode(image["data"], validate=True))
    except (ValueError, TypeError):
        return False
    _add_report_picture(doc, data)
    return True


def _render_indexing_legend(doc, distribution):
    rows = (distribution or {}).get("rows") or []
    if not rows:
        return
    paragraph = doc.add_paragraph()
    for index, row in enumerate(rows):
        color = _indexing_group_color(index, row).removeprefix("#")
        count = paragraph.add_run(_number(row.get("count"), decimal_places=0) + "  ")
        count.bold = True
        count.font.color.rgb = RGBColor.from_string(color)
        paragraph.add_run(str(row.get("path") or "") + "    ")


def _render_iks_explanation(doc):
    doc.add_paragraph(
        "Индекс качества сайта — это показатель того, насколько полезен сайт для пользователей "
        "с точки зрения Яндекса. Данный показатель напрямую не влияет на ранжирование сайтов "
        "в органической выдаче."
    )
    doc.add_paragraph(
        "При расчете индекса качества учитываются размер аудитории сайта, степень "
        "удовлетворенности пользователей, уровень доверия к сайту со стороны пользователей "
        "и Яндекса, а также множество других факторов, которые Яндекс не разглашает. Для "
        "расчета используются данные сервисов Яндекса. Значение ИКС регулярно пересчитывается "
        "с учетом множества факторов. За время между пересчетами, связанные с сайтом факторы "
        "могут измениться."
    )
    doc.add_paragraph("Для увеличения показателя ИКС необходимо:")
    for text in (
        "развивать популярность бренда как онлайн, так и офлайн;",
        "увеличивать трафик на сайт со всех возможных источников;",
        "расширять структуру сайта и семантику, задействовать максимальное количество URL, "
        "генерирующих трафик;",
        "работать над вовлеченностью пользователями ресурсом.",
    ):
        doc.add_paragraph(text, style="List Bullet")
    doc.add_paragraph(
        "Значение ИКС регулярно пересчитывается с учетом множества факторов, в том числе "
        "сезонности. За время между пересчетами, связанные с сайтом факторы могут измениться."
    )


def _render_webmaster(doc, payload, blocks):
    enabled = [
        code
        for code in ("iks", "indexing", "clicks_impressions", "ctr")
        if section_enabled(payload, code)
    ]
    if not enabled:
        return
    doc.add_heading("2) Индексация сайта (Яндекс.Вебмастер)", level=1)
    source = "yandex_webmaster"
    details = _webmaster_chart_details(payload)
    latest = _latest_webmaster_payload(payload)
    if "iks" in enabled:
        codes = tuple(
            code for code in ("iks", "quality_index") if _metric_has_data(payload, source, code)
        )[:1]
        if codes:
            code = codes[0]
            change = (_metric_source(payload, source).get("normalized_changes") or {}).get(
                code
            ) or {}
            iks_rows = _daily_rows(details, "iks")
            current = iks_rows[-1].get("value") if iks_rows else change.get("current")
            previous = iks_rows[0].get("value") if iks_rows else change.get("previous")
            delta = (
                abs(Decimal(str(current)) - Decimal(str(previous)))
                if current is not None and previous is not None
                else None
            )
            direction = None
            if current is not None and previous is not None:
                direction = (
                    "увеличилось"
                    if Decimal(str(current)) > Decimal(str(previous))
                    else "уменьшилось"
                    if Decimal(str(current)) < Decimal(str(previous))
                    else "не изменилось"
                )
            doc.add_paragraph(
                "В отчётном месяце значение ИКС сайта "
                + (
                    f"{direction} на {_number(delta, decimal_places=0)} - "
                    if delta not in (None, 0)
                    else "не изменилось - "
                    if delta == 0
                    else ""
                )
                + f"составляет {_number(current)} единиц."
            )
            _period_caption(
                doc,
                iks_rows,
                provider="webmaster",
                show_detail=False,
            )
            _add_report_picture(
                doc,
                _single_service_chart(
                    [(row.get("date"), row.get("value")) for row in iks_rows]
                    or _metric_series(payload, source, (code,))[0][1],
                    title=f"Индекс качества сайта (ИКС) — {_number(current)}",
                    color=WEBMASTER_COLORS[0],
                    minimum_step=10,
                    step=True,
                ),
            )
            _render_iks_explanation(doc)
    if "indexing" in enabled:
        doc.add_paragraph("Динамика количества страниц в поиске:", style="Table Heading")
        indexing_rows = _daily_rows(details, "indexed_pages")
        points = [(row.get("date"), row.get("value")) for row in indexing_rows]
        distribution = latest.get("path_distribution")
        _add_report_picture(
            doc,
            _indexing_group_chart(
                points or _metric_series(payload, source, ("indexed_pages",))[0][1],
                distribution,
            ),
        )
        _render_indexing_legend(doc, distribution)
        distribution_rows = (distribution or {}).get("rows") or []
        known_rows = [row for row in distribution_rows if row.get("path") != "Статус неизвестен"]
        first_pages = _decimal_or_none(points[0][1]) if points else None
        last_pages = _decimal_or_none(points[-1][1]) if points else None
        pages_delta = (
            last_pages - first_pages if first_pages is not None and last_pages is not None else None
        )
        if pages_delta is not None and abs(pages_delta) < 20:
            doc.add_paragraph("Количество страниц в индексе изменилось незначительно.")
        elif pages_delta is not None:
            doc.add_paragraph(
                f"Количество страниц в индексе {'увеличилось' if pages_delta > 0 else 'снизилось'} "
                f"на {_number(abs(pages_delta), decimal_places=0)}."
            )
        if known_rows:
            leader = max(known_rows, key=lambda row: row.get("count") or 0)
            doc.add_paragraph(
                f"Преимущественно в поиске находятся страницы раздела {leader.get('path')} — "
                f"{_number(leader.get('count'), decimal_places=0)} URL."
            )
            if (
                "blog" in str(leader.get("path") or "").casefold()
                or "стат" in str(leader.get("path") or "").casefold()
            ):
                doc.add_paragraph("Необходимо расширять структуру по коммерческим направлениям.")
    if "clicks_impressions" in enabled or "ctr" in enabled:
        doc.add_paragraph(
            "Данные по показам, кликам и CTR по всем запросам:", style="Table Heading"
        )
        query_rows = _daily_rows(details, "queries")
        _period_caption(doc, query_rows, provider="webmaster", detail="по дням")
        _add_report_picture(doc, _webmaster_search_chart(payload))
        current_summary = latest.get("query_summary") or {}
        previous_summary = latest.get("comparison_query_summary") or {}
        if not current_summary:
            current_summary, previous_summary = _webmaster_query_summary_from_changes(payload)
        if current_summary:
            _webmaster_query_summary_table(doc, current_summary, previous_summary)
        doc.add_paragraph(
            "Красным шрифтом представлены цифры, показывающие спад показателя по сравнению "
            "с предыдущим периодом, зелёным — рост. Запросы, по которым начинает показываться "
            "сайт, появляются неравномерно: сначала брендовые — наиболее кликабельные, а затем "
            "показываются небрендовые (они менее кликабельные), которые появляются на более "
            "низких позициях. Поэтому показатели CTR и позиции могут снижаться."
        )
        if current_summary and previous_summary:
            doc.add_paragraph(_webmaster_query_text(current_summary, previous_summary))
    if section_enabled(payload, "webmaster_popular_queries"):
        doc.add_paragraph("Самые кликабельные запросы:", style="Table Heading")
        current_queries = latest.get("popular_queries") or []
        previous_queries = latest.get("comparison_popular_queries") or []
        comment = blocks.get("webmaster_popular_queries") or payload.get("display_options", {}).get(
            "webmaster_queries_comment"
        )
        doc.add_paragraph(
            _clean(comment)
            if comment
            else (
                "Запросы в таблице отсортированы по кликабельности. Красным шрифтом "
                "представлены цифры, показывающие спад показателя по сравнению с предыдущим "
                "периодом, зелёным — рост, серым — оставшиеся без изменений."
            )
        )
        if current_queries:
            _webmaster_popular_table(doc, current_queries, previous_queries)
            doc.add_paragraph(
                _webmaster_popular_analysis(payload, current_queries, previous_queries)
            )
        elif not _add_uploaded_picture(doc, payload):
            doc.add_paragraph(
                "API не вернул список популярных запросов, и скриншот не загружен.",
                style="Data Missing",
            )


def _metrika_comparison_chart(payload, codes, *, title):
    changes = _metric_source(payload, "yandex_metrika").get("normalized_changes", {})
    rows = [(code, changes.get(code) or {}) for code in codes]
    rows = [item for item in rows if item[1].get("current") is not None]
    if not rows:
        return None
    labels = [METRIC_LABELS.get(code, code) for code, _ in rows]
    previous = [float(change.get("previous") or 0) for _, change in rows]
    current = [float(change.get("current") or 0) for _, change in rows]
    x = list(range(len(labels)))
    width = 0.26
    offset = 0.18
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure, axis = plt.subplots(figsize=(7.2, 3.55), dpi=150, facecolor="white")
        previous_bars = axis.bar(
            [value - offset for value in x],
            previous,
            width=width,
            color=[METRIKA_COLORS[i % len(METRIKA_COLORS)] for i in x],
            label="Предыдущий период",
        )
        current_bars = axis.bar(
            [value + offset for value in x],
            current,
            width=width,
            color="white",
            edgecolor=[METRIKA_COLORS[i % len(METRIKA_COLORS)] for i in x],
            hatch="////",
            linewidth=1.2,
            label="Текущий период",
        )
        for index, bar in enumerate(current_bars):
            bar.set_edgecolor(METRIKA_COLORS[index % len(METRIKA_COLORS)])
        axis.set_xticks(x, labels)
        axis.set_title(title, loc="left", fontsize=13, color="#30343B")
        _style_axis(axis, grid_axis="y")
        axis.legend(
            handles=(previous_bars[0], current_bars[0]),
            labels=("Предыдущий период", "Текущий период"),
            loc="upper center",
            bbox_to_anchor=(0.5, -0.14),
            ncol=2,
            frameon=False,
        )
        figure.subplots_adjust(left=0.09, right=0.98, top=0.9, bottom=0.27)
        return _save_figure(figure)


def _metrika_source_label(name):
    return {
        "search": "Переходы из поисковых систем",
        "direct": "Прямые заходы",
        "referral": "Переходы по ссылкам на сайтах",
        "ad": "Переходы по рекламе",
        "ads": "Переходы по рекламе",
        "advertising": "Переходы по рекламе",
        "social": "Переходы из социальных сетей",
        "internal": "Внутренние переходы",
        "recommend": "Переходы из рекомендательных систем",
        "messenger": "Переходы из мессенджеров",
        "saved": "Переходы с сохранённых страниц",
        "email": "Переходы из почтовых рассылок",
        "qrcode": "Переходы по QR-кодам",
        "other": "Остальные источники",
    }.get(name, name)


def _style_traffic_source_icons(table):
    markers = {
        "Переходы из поисковых систем": ("⌕", "F04444"),
        "Прямые заходы": ("➜", "8B98A7"),
        "Переходы по ссылкам на сайтах": ("↗", "3388FF"),
        "Переходы из социальных сетей": ("●", "9B51E0"),
        "Внутренние переходы": ("⌂", "0FBDA0"),
        "Переходы из рекомендательных систем": ("★", "FF3399"),
        "Переходы из мессенджеров": ("✉", "2D9CDB"),
        "Переходы с сохранённых страниц": ("▣", "8B98A7"),
        "Переходы по рекламе": ("AD", "FF8A00"),
        "Остальные источники": ("●", "8B98A7"),
    }
    for row in table.rows[1:]:
        cell = row.cells[0]
        label = cell.text.strip()
        if label == "Итого и среднее":
            continue
        marker, color = markers.get(label, ("●", "8B98A7"))
        cell.text = ""
        paragraph = cell.paragraphs[0]
        _style_run(
            paragraph.add_run(f"{marker}  "),
            size=8 if marker == "AD" else 10,
            color=color,
            bold=True,
        )
        _style_run(paragraph.add_run(label), size=11)


def _metrika_sources_chart(facts):
    order = (
        "search",
        "direct",
        "referral",
        "advertising",
        "social",
        "internal",
        "recommend",
        "messenger",
        "saved",
        "email",
        "qrcode",
        "other",
    )
    color_map = {
        "search": "#7A45E5",
        "direct": "#FF3399",
        "referral": "#0FBDA0",
        "ad": "#3388FF",
        "ads": "#3388FF",
        "advertising": "#3388FF",
        "social": "#FFB851",
        "internal": "#FFB851",
        "recommend": "#FF3399",
        "messenger": "#2D9CDB",
        "saved": "#8B98A7",
        "other": "#8B98A7",
    }
    names = [name for name in order if name in facts]
    names.extend(name for name in facts if name not in names)
    useful = [
        (
            name,
            [(point.get("month"), point.get("value")) for point in facts[name].get("series", [])],
        )
        for name in names
        if any(
            _decimal_or_none(point.get("value")) not in (None, 0)
            for point in facts[name].get("series", [])
        )
    ]
    if not useful:
        return None
    labels = [_month_short(month) for month, _ in useful[0][1]]
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure, axis = plt.subplots(figsize=(7.2, 4.35), dpi=150, facecolor="white")
        legend_handles = []
        for index, (name, points) in enumerate(useful):
            values = [float(value) if value is not None else float("nan") for _, value in points]
            change = facts[name].get("change") or {}
            current = change.get("current")
            legend_label = (
                f"{_metrika_source_label(name)} — {_compact_number(current)}"
                if current is not None
                else _metrika_source_label(name)
            )
            color = color_map.get(name, METRIKA_COLORS[index % len(METRIKA_COLORS)])
            axis.plot(
                labels,
                values,
                linewidth=1.8,
                color=color,
                label=legend_label,
            )
            legend_handles.append(Patch(facecolor=color, edgecolor=color, label=legend_label))
        axis.set_title("Источники, сводка", loc="left", fontsize=13, color="#30343B")
        axis.set_ylabel("Визиты", color="#677485")
        axis.yaxis.set_major_formatter(
            FuncFormatter(lambda value, _position: _axis_compact_number(value))
        )
        axis.yaxis.set_major_locator(MaxNLocator(nbins=5, integer=True))
        _style_axis(axis)
        axis.legend(
            handles=legend_handles,
            loc="upper center",
            bbox_to_anchor=(0.5, -0.15),
            ncol=2,
            frameon=False,
            fontsize=8,
        )
        legend_rows = max(1, math.ceil(len(legend_handles) / 2))
        figure.subplots_adjust(
            left=0.09, right=0.98, top=0.9, bottom=min(0.48, 0.2 + legend_rows * 0.065)
        )
        return _save_figure(figure)


def _metrika_period_rows(payload, key):
    options = payload.get("display_options", {})
    robotness = options.get("metrika_robotness", "humans")
    force_search = key in {"search_engines", "search_landing_pages"}
    detail_key = "landing_pages" if key == "search_landing_pages" else key
    segment = "search" if force_search or options.get("metrika_search_segment", True) else "all"
    periods = []
    for detail in _period_details(payload, "yandex_metrika"):
        source = _detail_payload(detail)
        variant = (
            ((source.get("detail_variants") or {}).get(segment) or {}).get(robotness)
            or (source.get("search_details") or {}).get(robotness)
            or source
        )
        rows = (
            source.get("traffic_source_details") or []
            if key == "traffic_source_details"
            else variant.get(detail_key) or []
        )
        periods.append(
            {
                "period_start": detail.get("period_start"),
                "period_end": detail.get("period_end"),
                "rows": rows,
                "total": (
                    source.get("traffic_source_total") or {}
                    if key == "traffic_source_details"
                    else variant.get(f"{detail_key}_total") or {}
                ),
                "payload": source,
            }
        )
    return periods


def _row_dimension(row, index):
    dimensions = row.get("dimensions") or []
    return dimensions[index] if index < len(dimensions) else {"id": "", "name": ""}


def _aggregate_detail_rows(rows, key):
    result = {}
    for row in rows:
        label = key(row)
        if not label:
            continue
        target = result.setdefault(
            label,
            {"visits": Decimal(0), "users": Decimal(0), "bounce_sum": Decimal(0)},
        )
        visits = _decimal_or_none(row.get("visits")) or Decimal(0)
        users = _decimal_or_none(row.get("users")) or Decimal(0)
        bounce = _decimal_or_none(row.get("bounce_rate")) or Decimal(0)
        target["visits"] += visits
        target["users"] += users
        target["bounce_sum"] += bounce * visits
    for values in result.values():
        values["bounce_rate"] = (
            values.pop("bounce_sum") / values["visits"] if values["visits"] else Decimal(0)
        )
    return result


def _aggregate_traffic_source_rows(periods):
    """Aggregate Metrika source rows using visit-weighted rate/average metrics."""
    result = {}
    for period in periods:
        for row in period.get("rows") or []:
            code = str(row.get("code") or _row_dimension(row, 0).get("id") or "other")
            target = result.setdefault(
                code,
                {
                    "visits": Decimal(0),
                    "users": Decimal(0),
                    "bounce_sum": Decimal(0),
                    "page_depth_sum": Decimal(0),
                    "duration_sum": Decimal(0),
                },
            )
            visits = _decimal_or_none(row.get("visits")) or Decimal(0)
            target["visits"] += visits
            target["users"] += _decimal_or_none(row.get("users")) or Decimal(0)
            target["bounce_sum"] += (_decimal_or_none(row.get("bounce_rate")) or 0) * visits
            target["page_depth_sum"] += (_decimal_or_none(row.get("page_depth")) or 0) * visits
            target["duration_sum"] += (
                _decimal_or_none(row.get("avg_visit_duration_seconds")) or 0
            ) * visits
    for values in result.values():
        visits = values["visits"]
        values["bounce_rate"] = values.pop("bounce_sum") / visits if visits else Decimal(0)
        values["page_depth"] = values.pop("page_depth_sum") / visits if visits else Decimal(0)
        values["avg_visit_duration_seconds"] = (
            values.pop("duration_sum") / visits if visits else Decimal(0)
        )
    return result


def _duration(value):
    seconds = max(0, int(round(float(_decimal_or_none(value) or 0))))
    minutes, seconds = divmod(seconds, 60)
    return f"{minutes} м {seconds} с" if minutes else f"{seconds} с"


def _metrika_sources_quarter_table(doc, periods):
    quarter_rows = (
        (periods[-1].get("payload") or {}).get("traffic_source_quarter_details") if periods else []
    )
    totals = _aggregate_traffic_source_rows(
        [{"rows": quarter_rows}] if quarter_rows else periods[-3:]
    )
    ordered = sorted(
        ((code, values) for code, values in totals.items() if values["visits"] > 0),
        key=lambda item: item[1]["visits"],
        reverse=True,
    )
    if not ordered:
        return None
    rows = [
        (
            _metrika_source_label(code),
            _number(values["visits"], decimal_places=0),
            _number(values["users"], decimal_places=0),
            _number(values["bounce_rate"], "%", decimal_places=2),
            _number(values["page_depth"], decimal_places=2),
            _duration(values["avg_visit_duration_seconds"]),
        )
        for code, values in ordered
    ]
    provider_total = (
        (periods[-1].get("payload") or {}).get("traffic_source_quarter_total") if periods else {}
    ) or {}
    if provider_total:
        rows.insert(
            0,
            (
                "Итого и среднее",
                _number(provider_total.get("visits"), decimal_places=0),
                _number(provider_total.get("users"), decimal_places=0),
                _number(provider_total.get("bounce_rate"), "%", decimal_places=2),
                _number(provider_total.get("page_depth"), decimal_places=2),
                _duration(provider_total.get("avg_visit_duration_seconds")),
            ),
        )
    table = _table(
        doc,
        ("Источник", "Визиты", "Посетители", "Отказы", "Глубина просмотра", "Время на сайте"),
        rows,
        [5.4, 2.3, 2.5, 2.3, 3.0, 2.7],
        header_fill="F7F7F7",
    )
    _style_provider_table(table)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
    _style_traffic_source_icons(table)
    if provider_total:
        for column_index, cell in enumerate(table.rows[1].cells):
            _shade_cell(cell, "F7F7F7")
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.bold = True
    _set_table_borders(table, "E2E4E8", size="3")
    return table


def _search_engine_name(row):
    dimension = _row_dimension(row, 0)
    raw = f"{dimension.get('id', '')} {dimension.get('name', '')}".casefold()
    if "yandex" in raw or "яндекс" in raw:
        return "Яндекс"
    if "google" in raw:
        return "Google"
    return str(dimension.get("name") or dimension.get("id") or "").strip()


def _search_engine_code(label):
    raw = str(label or "").casefold()
    if "yandex" in raw or "яндекс" in raw:
        return "yandex"
    for code in ("google", "bing", "yahoo"):
        if code in raw:
            return code
    return raw


def _compact_number(value):
    number = _decimal_or_none(value)
    if number is None:
        return "—"
    if abs(number) >= 1000:
        return f"{_number(number / Decimal(1000), decimal_places=1)} тыс."
    return _number(number, decimal_places=0)


def _axis_compact_number(value):
    number = _decimal_or_none(value)
    if number is None:
        return "0"
    if abs(number) >= 1000:
        return f"{_number(number / Decimal(1000), decimal_places=1)}к"
    return _number(number, decimal_places=0)


def _metrika_search_quarter_chart(periods):
    if not periods:
        return None
    aggregates = [
        _aggregate_detail_rows(period["rows"], _search_engine_name) for period in periods[-3:]
    ]
    latest = aggregates[-1]
    engines = sorted(
        latest,
        key=lambda label: latest[label].get("visits") or Decimal(0),
        reverse=True,
    )[:3]
    if not engines:
        return None
    labels = []
    for period in periods[-3:]:
        parsed = date.fromisoformat(str(period["period_start"])[:10])
        labels.append(MONTHS[parsed.month - 1][:3])
    series = []
    colors = {"Google": "#7A45E5", "Яндекс": "#FF3399"}
    for index, engine in enumerate(engines):
        series.append(
            (
                engine,
                [aggregate.get(engine, {}).get("visits") or Decimal(0) for aggregate in aggregates],
                colors.get(engine, METRIKA_COLORS[index % len(METRIKA_COLORS)]),
            )
        )
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure, axis = plt.subplots(figsize=(7.2, 3.45), dpi=150, facecolor="white")
        for label, values, color in series:
            axis.plot(
                labels,
                [float(value) for value in values],
                color=color,
                linewidth=1.6,
                label=label,
            )
        axis.set_title("Визиты", loc="left", fontsize=11, color="#30343B")
        axis.yaxis.set_major_formatter(
            FuncFormatter(lambda value, _position: _axis_compact_number(value))
        )
        axis.yaxis.set_major_locator(MaxNLocator(nbins=4, integer=True))
        axis.set_ylim(bottom=0)
        _style_axis(axis)
        legend_handles = [
            Line2D(
                [0],
                [0],
                marker="s",
                linestyle="none",
                markersize=5,
                markerfacecolor=color,
                markeredgecolor=color,
                label=f"{label}  {_compact_number(values[-1])}",
            )
            for label, values, color in series
        ]
        axis.legend(
            handles=legend_handles,
            loc="upper center",
            bbox_to_anchor=(0.5, -0.14),
            ncol=min(5, len(legend_handles)),
            frameon=False,
            fontsize=7.5,
            handletextpad=0.35,
            columnspacing=1.1,
        )
        figure.subplots_adjust(left=0.1, right=0.98, top=0.9, bottom=0.27)
        return _save_figure(figure)


def _metrika_comparison_bars(rows, *, title):
    useful = [row for row in rows if row.get("current") is not None]
    if not useful:
        return None
    x = list(range(len(useful)))
    width = 0.26
    offset = 0.18
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure, axis = plt.subplots(figsize=(7.2, 3.55), dpi=150, facecolor="white")
        for index, row in enumerate(useful):
            color = row.get("color") or METRIKA_COLORS[index % len(METRIKA_COLORS)]
            axis.bar(
                index - offset,
                float(row.get("previous") or 0),
                width=width,
                color=color,
                edgecolor=color,
                linewidth=0.5,
            )
            axis.bar(
                index + offset,
                float(row.get("current") or 0),
                width=width,
                color=color,
                alpha=0.72,
                hatch="////",
                edgecolor="#FFFFFF",
                linewidth=0.5,
            )
        axis.set_xticks(x, ["" for _row in useful])
        axis.set_title(title, loc="left", fontsize=13, color="#30343B")
        axis.yaxis.set_major_formatter(
            FuncFormatter(lambda value, _position: _axis_compact_number(value))
        )
        axis.yaxis.set_major_locator(MaxNLocator(nbins=5, integer=True))
        _style_axis(axis, grid_axis="y")
        handles = []
        labels = []
        for index, row in enumerate(useful):
            color = row.get("color") or METRIKA_COLORS[index % len(METRIKA_COLORS)]
            handles.append(Patch(facecolor=color, edgecolor=color))
            labels.append(
                f"{row['label']} — {_compact_number(row.get('previous'))} → "
                f"{_compact_number(row.get('current'))}"
            )
        axis.legend(
            handles,
            labels,
            loc="upper center",
            bbox_to_anchor=(0.5, -0.14),
            ncol=min(2, len(handles)),
            frameon=False,
            fontsize=8,
        )
        figure.subplots_adjust(left=0.09, right=0.98, top=0.9, bottom=0.27)
        return _save_figure(figure)


def _metrika_detail_table(
    doc,
    rows,
    *,
    first_header,
    metrics=("visits", "users", "bounce_rate"),
    include_total=True,
    total_values=None,
):
    metric_headers = {
        "visits": "Визиты",
        "users": "Посетители",
        "bounce_rate": "Отказы, %",
    }
    period_headers = []
    for code in metrics:
        period_headers.extend(
            (
                f"{metric_headers[code]}\nСегмент A",
                f"{metric_headers[code]}\nСегмент B",
            )
        )
    if len(metrics) == 3:
        widths = [5.3, *([2.2] * 6)]
    elif len(metrics) == 2:
        widths = [8.5, *([2.5] * 4)]
    else:
        widths = [10.5, 4.0, 4.0]
    display_rows = list(rows)
    if include_total and display_rows:

        def total(period_index):
            values = {"visits": Decimal(0), "users": Decimal(0), "bounce_sum": Decimal(0)}
            for _label, current, previous in display_rows:
                source = (current, previous)[period_index]
                visits = _decimal_or_none(source.get("visits")) or Decimal(0)
                values["visits"] += visits
                values["users"] += _decimal_or_none(source.get("users")) or Decimal(0)
                values["bounce_sum"] += (_decimal_or_none(source.get("bounce_rate")) or 0) * visits
            values["bounce_rate"] = (
                values.pop("bounce_sum") / values["visits"] if values["visits"] else Decimal(0)
            )
            return values

        current_total, previous_total = total_values or (total(0), total(1))
        display_rows.insert(0, ("Итого и среднее", current_total, previous_total))
    table_rows = []
    for label, current, previous in display_rows:
        values = [label]
        for code in metrics:
            values.extend(
                (
                    _provider_value(code, current.get(code)),
                    _provider_value(code, previous.get(code)),
                )
            )
        table_rows.append(tuple(values))
    table = _table(
        doc,
        (first_header, *period_headers),
        table_rows,
        widths,
        header_fill="F7F7F7",
    )
    for column_index, cell in enumerate(table.rows[0].cells):
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
        )
        for run in cell.paragraphs[0].runs:
            _style_run(
                run,
                size=11 if column_index == 0 else 9,
                color="000000" if column_index == 0 else "7A8796",
            )
    for row_index, (_label, current, previous) in enumerate(display_rows, start=1):
        label_cell = table.rows[row_index].cells[0]
        for run in label_cell.paragraphs[0].runs:
            _style_run(run, size=11)
        for metric_index, code in enumerate(metrics):
            current_cell = table.rows[row_index].cells[1 + metric_index * 2]
            previous_cell = table.rows[row_index].cells[2 + metric_index * 2]
            previous_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            previous_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            current_cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            for run in previous_cell.paragraphs[0].runs:
                _style_run(run, size=11)
            for run in current_cell.paragraphs[0].runs:
                _style_run(run, size=11)
            delta = _relative_delta(current.get(code), previous.get(code))
            if delta in (None, 0):
                continue
            change = current_cell.add_paragraph(_number(delta, "%", decimal_places=2))
            change.alignment = WD_ALIGN_PARAGRAPH.CENTER
            change.paragraph_format.space_after = Pt(0)
            _style_run(
                change.runs[0],
                size=8,
                color=_change_color(
                    current.get(code),
                    previous.get(code),
                    lower_is_better=code == "bounce_rate",
                ),
            )
    _set_table_borders(table, "E2E4E8", size="3")
    if include_total and len(table.rows) > 1:
        for cell in table.rows[1].cells:
            _shade_cell(cell, "F7F7F7")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
    return table


def _style_search_engine_icons(table):
    markers = {
        "google": ("G", "4285F4"),
        "яндекс": ("Я", "FC3F1D"),
        "yandex": ("Я", "FC3F1D"),
        "bing": ("⌕", "008373"),
        "duckduckgo": ("●", "DE5833"),
        "rambler": ("/", "315EFB"),
        "mail.ru": ("@", "168DE2"),
        "yahoo": ("Y!", "6001D2"),
        "yahoo!": ("Y!", "6001D2"),
    }
    for row in table.rows[1:]:
        cell = row.cells[0]
        label = cell.text.strip()
        if label == "Итого и среднее":
            for run in cell.paragraphs[0].runs:
                _style_run(run, size=11, bold=True)
            continue
        marker, color = markers.get(label.casefold(), ("●", "8B98A7"))
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        marker_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        marker_run.text = marker + "  "
        _style_run(marker_run, size=10, color=color, bold=True)
        _style_run(paragraph.add_run(label), size=11)


def _metrika_search_engine_text(current):
    total = sum((row.get("visits") or Decimal(0) for row in current.values()), Decimal(0))
    ordered = []
    for label in ("Яндекс", "Google"):
        values = current.get(label)
        if values:
            share = (
                _number(values["visits"] * 100 / total, "%", decimal_places=2) if total else None
            )
            ordered.append(
                f"{label} составляет {share} от всего поискового трафика"
                if share is not None
                else f"доля {label} не рассчитана"
            )
    other = total - sum(
        (current.get(label, {}).get("visits", Decimal(0)) for label in ("Яндекс", "Google")),
        Decimal(0),
    )
    suffix = (
        "На остальные поисковые системы приходится менее 1% всего поискового трафика, "
        "их данными можно пренебречь."
        if total and other * 100 / total < 1
        else (
            "На остальные поисковые системы приходится "
            f"{_number(other * 100 / total, '%', decimal_places=1)}."
        )
        if total
        else ""
    )
    return ". ".join(ordered) + (". " if ordered else "") + suffix


def _comparison_period_pills(doc, periods):
    if len(periods) < 2:
        return
    labels = []
    for period in reversed(periods[-2:]):
        start = date.fromisoformat(str(period["period_start"])[:10])
        end = date.fromisoformat(str(period.get("period_end") or period["period_start"])[:10])
        labels.append(
            f"{start.day}—{end.day} {MONTHS_GENITIVE[end.month - 1]} ⌄"
            if start.month == end.month and start.year == end.year
            else (
                f"{start.day} {MONTHS_GENITIVE[start.month - 1]} — "
                f"{end.day} {MONTHS_GENITIVE[end.month - 1]} ⌄"
            )
        )
    _add_report_picture(
        doc,
        _period_pills_image([(labels[0], "A"), ("⇄", "swap"), (labels[1], "B")]),
        width=11.2,
    )


def _region_key(row):
    area = str(_row_dimension(row, 0).get("name") or "").casefold()
    city = str(_row_dimension(row, 1).get("name") or "").casefold()
    if city in {"москва", "moscow"}:
        return "moscow"
    if city in {"санкт-петербург", "saint petersburg", "st. petersburg"}:
        return "saint_petersburg"
    if "область не определена" in area or area in {"area not defined"}:
        return "area_undefined"
    if city in {"", "не определено", "undefined", "not defined"}:
        return "undefined"
    return city or area


def _aggregate_regions(rows):
    result = _aggregate_detail_rows(rows, _region_key)
    for combined, needle, city_key in (
        ("moscow_region", "москов", "moscow"),
        ("saint_petersburg_region", "ленинград", "saint_petersburg"),
    ):
        selected = []
        for row in rows:
            area = str(_row_dimension(row, 0).get("name") or "").casefold()
            if needle in area or _region_key(row) == city_key:
                selected.append(row)
        aggregate = _aggregate_detail_rows(selected, lambda _row, key=combined: key)
        if combined in aggregate:
            result[combined] = aggregate[combined]
    return result


REGION_LABELS = {
    "moscow": "Москва",
    "moscow_region": "Москва и Московская область",
    "saint_petersburg": "Санкт-Петербург",
    "saint_petersburg_region": "Санкт-Петербург и Ленинградская область",
    "undefined": "Не определено",
    "area_undefined": "Область не определена",
}


def _landing_url(row):
    return str(_row_dimension(row, 1).get("name") or _row_dimension(row, 1).get("id") or "")


def _belongs_to_project(payload, row):
    url = _landing_url(row).strip()
    parsed = urlsplit(url)
    if not parsed.netloc:
        return True
    expected = str(
        payload.get("project", {}).get("normalized_domain")
        or payload.get("project", {}).get("domain")
        or ""
    ).strip()
    expected_host = urlsplit(expected if "://" in expected else f"//{expected}").hostname or ""
    actual_host = parsed.hostname or ""

    def normalize(value):
        return value.casefold().removeprefix("www.").rstrip(".")

    return bool(expected_host and normalize(actual_host) == normalize(expected_host))


def _url_group(payload, url):
    for group in sorted(
        (item for item in payload.get("project", {}).get("url_groups", []) if item.get("active")),
        key=lambda item: (-int(item.get("priority") or 0), item.get("name") or ""),
    ):
        for rule in sorted(
            (item for item in group.get("rules", []) if item.get("active")),
            key=lambda item: -int(item.get("priority") or 0),
        ):
            pattern = str(rule.get("pattern") or "")
            kind = rule.get("type")
            matched = (
                url.startswith(pattern)
                if kind == "starts_with"
                else pattern in url
                if kind == "contains"
                else re.search(pattern, url) is not None
                if kind == "regex"
                else False
            )
            if matched:
                return group.get("name") or group.get("slug")
    return "Остальные страницы"


def _configured_url_groups(payload, kind):
    configured = payload.get("display_options", {}).get("metrika_url_segments", {}).get(kind) or []
    if configured:
        return [
            {
                "name": str(group.get("name") or "Раздел"),
                "patterns": [str(value) for value in group.get("patterns") or [] if value],
            }
            for group in configured
            if group.get("patterns")
        ]
    if kind != "categories":
        return []
    return [
        {
            "name": str(group.get("name") or group.get("slug") or "Раздел"),
            "patterns": [
                {"value": str(rule.get("pattern")), "type": rule.get("type")}
                for rule in group.get("rules") or []
                if rule.get("active") and rule.get("pattern")
            ],
        }
        for group in payload.get("project", {}).get("url_groups", [])
        if group.get("active")
    ]


def _url_matches_pattern(url, pattern):
    url = str(url or "").strip()
    rule_type = pattern.get("type") if isinstance(pattern, dict) else ""
    pattern_value = pattern.get("value") if isinstance(pattern, dict) else pattern
    pattern_value = str(pattern_value or "").strip()
    if not url or not pattern_value:
        return False
    if rule_type == "starts_with":
        return url.casefold().startswith(pattern_value.casefold())
    if rule_type == "contains":
        return pattern_value.casefold() in url.casefold()
    if rule_type == "regex":
        try:
            return re.search(pattern_value, url) is not None
        except re.error:
            return False
    parsed = urlsplit(url)
    candidates = (url, parsed.path or "/", f"{parsed.netloc}{parsed.path}")
    wildcard = "*" in pattern_value or "?" in pattern_value
    if wildcard:
        return any(
            fnmatchcase(candidate.casefold(), pattern_value.casefold()) for candidate in candidates
        )
    return any(
        candidate.casefold().startswith(pattern_value.casefold()) for candidate in candidates
    )


def _url_pattern_label(pattern):
    return str(pattern.get("value") if isinstance(pattern, dict) else pattern)


def _aggregate_configured_groups(rows, groups):
    result = {
        group["name"]: {
            "visits": Decimal(0),
            "users": Decimal(0),
            "bounce_sum": Decimal(0),
        }
        for group in groups
    }
    for row in rows:
        url = _landing_url(row)
        visits = _decimal_or_none(row.get("visits")) or Decimal(0)
        users = _decimal_or_none(row.get("users")) or Decimal(0)
        bounce = _decimal_or_none(row.get("bounce_rate")) or Decimal(0)
        for group in groups:
            if not any(_url_matches_pattern(url, pattern) for pattern in group["patterns"]):
                continue
            target = result[group["name"]]
            target["visits"] += visits
            target["users"] += users
            target["bounce_sum"] += bounce * visits
    for values in result.values():
        values["bounce_rate"] = (
            values.pop("bounce_sum") / values["visits"] if values["visits"] else Decimal(0)
        )
    return result


def _group_period_values(periods, groups):
    return [
        {
            "period_start": period.get("period_start"),
            "values": _aggregate_configured_groups(period.get("rows") or [], groups),
        }
        for period in periods
    ]


def _metrika_groups_chart(periods, groups, *, title="Визиты", only_group=None):
    grouped = _group_period_values(periods, groups)
    if not grouped:
        return None
    selected = [group for group in groups if not only_group or group["name"] == only_group]
    if not selected:
        return None
    labels = [
        MONTHS[date.fromisoformat(str(row["period_start"])[:10]).month - 1][:3] for row in grouped
    ]
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 9}):
        figure, axis = plt.subplots(figsize=(7.2, 3.45), dpi=150, facecolor="white")
        legend_handles = []
        for index, group in enumerate(selected):
            values = [row["values"].get(group["name"], {}).get("visits", 0) for row in grouped]
            pattern_label = " › ".join(
                _url_pattern_label(pattern) for pattern in group.get("patterns") or []
            )
            legend = f"{pattern_label or group['name']}  {_compact_number(values[-1])}"
            color = METRIKA_COLORS[index % len(METRIKA_COLORS)]
            axis.plot(
                labels,
                [float(value) for value in values],
                color=color,
                linewidth=1.6,
                label=legend,
            )
            legend_handles.append(Patch(facecolor=color, edgecolor=color, label=legend))
        axis.set_title(title, loc="left", fontsize=11, color="#30343B")
        axis.yaxis.set_major_formatter(
            FuncFormatter(lambda value, _position: _axis_compact_number(value))
        )
        axis.yaxis.set_major_locator(MaxNLocator(nbins=4, integer=True))
        axis.set_ylim(bottom=0)
        _style_axis(axis)
        axis.legend(
            handles=legend_handles,
            loc="upper center",
            bbox_to_anchor=(0.5, -0.14),
            ncol=min(3, len(selected)),
            frameon=False,
            fontsize=7.5,
        )
        figure.subplots_adjust(left=0.1, right=0.98, top=0.9, bottom=0.29)
        return _save_figure(figure)


def _group_dynamics_text(periods, groups, *, prefix="Раздел", subsections=()):
    grouped = _group_period_values(periods, groups)
    if not grouped:
        return "Данные для сравнения отсутствуют."
    current = grouped[-1]["values"]
    previous = grouped[-2]["values"] if len(grouped) >= 2 else {}
    parts = []
    for group in groups:
        name = group["name"]
        delta = _relative_delta(
            current.get(name, {}).get("visits"), previous.get(name, {}).get("visits")
        )
        group_url = _configured_group_label(group)
        part = f"{prefix} «{name}» ({group_url}) — " + (
            f"{'+' if delta > 0 else '−' if delta < 0 else ''}"
            f"{_number(abs(delta), '%', decimal_places=1)}"
            if delta is not None
            else "нет базы сравнения"
        )
        latest_rows = periods[-1].get("rows") or []
        matching = [
            row
            for row in latest_rows
            if any(
                _url_matches_pattern(_landing_url(row), pattern)
                for pattern in group.get("patterns") or []
            )
        ]
        popular_subsections = []
        for subsection in subsections:
            relevant = [
                row
                for row in matching
                if any(
                    _url_matches_pattern(_landing_url(row), pattern)
                    for pattern in subsection.get("patterns") or []
                )
            ]
            if relevant:
                popular_subsections.append(
                    (
                        sum(
                            (_decimal_or_none(row.get("visits")) or 0 for row in relevant),
                            Decimal(0),
                        ),
                        subsection,
                    )
                )
        popular_subsections.sort(key=lambda item: (-item[0], item[1]["name"].casefold()))
        if popular_subsections:
            labels = [
                f"{item['name']} ({_configured_group_label(item)})"
                for _visits, item in popular_subsections[:2]
            ]
            part += ". Здесь самые популярные подразделы — " + ", ".join(labels)
        parts.append(part)
    return "\n".join(f"{part}." for part in parts)


def _group_overview_text(periods, groups, *, commercial=False):
    grouped = _group_period_values(periods, groups)
    if not grouped or not groups:
        return ""
    current = grouped[-1]["values"]
    previous = grouped[-2]["values"] if len(grouped) >= 2 else {}
    leader = max(groups, key=lambda group: current.get(group["name"], {}).get("visits", 0))
    current_total = sum(
        (current.get(group["name"], {}).get("visits", 0) for group in groups), Decimal(0)
    )
    previous_total = sum(
        (previous.get(group["name"], {}).get("visits", 0) for group in groups), Decimal(0)
    )
    delta = _relative_delta(current_total, previous_total)
    trend = (
        "увеличение"
        if delta is not None and delta > 0
        else "снижение"
        if delta is not None and delta < 0
        else "стабильный уровень"
    )
    if commercial:
        return f"По разделам услуг наблюдается {trend} поискового трафика."
    leader_label = f"«{leader['name']}» ({_configured_group_label(leader)})"
    return (
        f"На информационный раздел {leader_label} приходится основная доля трафика. "
        f"В отчётном месяце наблюдается {trend} трафика на информационные разделы."
    )


def _metrika_goal_icon(goal):
    goal_type = re.sub(r"[^a-z]", "", str(goal.get("type") or "").casefold())
    exact_types = {
        "conditionalcall": "call",
        "call": "call",
        "step": "step",
        "url": "url",
        "multi": "multi",
        "email": "email",
        "phone": "phone",
        "action": "action",
        "visitduration": "visit_duration",
    }
    if goal_type in exact_types:
        return exact_types[goal_type]
    text = " ".join(
        str(goal.get(key) or "") for key in ("label", "name", "identifier", "condition")
    ).casefold()
    if any(token in text for token in ("телефон", "phone", "звон")):
        return "call"
    if any(token in text for token in ("форма", "заяв", "заказ", "обратн", "отзыв")):
        return "action"
    if any(token in text for token in ("переход", "url", "http", "страниц")):
        return "url"
    if any(token in text for token in ("общая", "общий", "составная")):
        return "multi"
    return "action"


def _draw_metrika_goal_icon(figure, goal, *, x=0.042, y=0.878, color="#7A45E5"):
    """Draw the exact goal-type icons embedded in the supplied Metrika page."""
    kind = _metrika_goal_icon(goal)
    asset = Path(__file__).resolve().parent / "assets" / "metrika_goal_icons" / f"{kind}.png"
    if asset.exists():
        icon_axis = figure.add_axes((x - 0.011, y - 0.016, 0.024, 0.032))
        icon_axis.imshow(plt.imread(asset))
        icon_axis.axis("off")
        return
    transform = figure.transFigure
    linewidth = 1.15
    if kind == "target":
        figure.patches.extend(
            (
                Circle(
                    (x, y),
                    0.009,
                    transform=transform,
                    fill=False,
                    edgecolor=color,
                    linewidth=linewidth,
                ),
                Circle((x, y), 0.003, transform=transform, facecolor=color, edgecolor=color),
            )
        )
        return
    if kind == "page":
        figure.patches.append(
            Rectangle(
                (x - 0.007, y - 0.012),
                0.014,
                0.024,
                transform=transform,
                fill=False,
                edgecolor=color,
                linewidth=linewidth,
            )
        )
        figure.lines.extend(
            (
                Line2D(
                    (x + 0.001, x + 0.007),
                    (y + 0.012, y + 0.006),
                    transform=transform,
                    color=color,
                    linewidth=linewidth,
                ),
                Line2D(
                    (x + 0.001, x + 0.001),
                    (y + 0.012, y + 0.006),
                    transform=transform,
                    color=color,
                    linewidth=linewidth,
                ),
                Line2D(
                    (x + 0.001, x + 0.007),
                    (y + 0.006, y + 0.006),
                    transform=transform,
                    color=color,
                    linewidth=linewidth,
                ),
            )
        )
        return
    if kind == "sliders":
        offsets = ((0.006, -0.003), (0, 0.003), (-0.006, -0.002))
        for y_offset, knob_offset in offsets:
            line_y = y + y_offset
            figure.lines.append(
                Line2D(
                    (x - 0.009, x + 0.009),
                    (line_y, line_y),
                    transform=transform,
                    color=color,
                    linewidth=linewidth,
                )
            )
            figure.patches.append(
                Circle(
                    (x + knob_offset, line_y),
                    0.0026,
                    transform=transform,
                    facecolor="white",
                    edgecolor=color,
                    linewidth=linewidth,
                )
            )
        return
    figure.lines.extend(
        (
            Line2D(
                (x - 0.006, x + 0.006),
                (y + 0.008, y - 0.008),
                transform=transform,
                color=color,
                linewidth=2.0,
            ),
            Line2D(
                (x - 0.008, x - 0.004),
                (y + 0.01, y + 0.006),
                transform=transform,
                color=color,
                linewidth=linewidth,
            ),
            Line2D(
                (x + 0.004, x + 0.008),
                (y - 0.006, y - 0.01),
                transform=transform,
                color=color,
                linewidth=linewidth,
            ),
        )
    )


def _metrika_goal_image(goal, periods):
    labels = [_month_short(row["period_start"]) for row in periods]
    conversion = [float(row.get("conversion_rate") or 0) for row in periods]
    visits = [float(row.get("visits") or 0) for row in periods]
    reaches = [float(row.get("reaches") or 0) for row in periods]
    with plt.rc_context({"font.family": CHART_FONT, "font.size": 8}):
        figure = plt.figure(figsize=(7.2, 2.15), dpi=150, facecolor="white")
        figure.patches.append(
            Rectangle(
                (0.008, 0.025),
                0.984,
                0.95,
                transform=figure.transFigure,
                fill=False,
                edgecolor="#E3E8EF",
                linewidth=0.8,
            )
        )
        title = _clean(goal.get("label") or goal.get("name") or "Цель")
        title_lines = textwrap.wrap(title, width=30)[:2] or ["Цель"]
        _draw_metrika_goal_icon(figure, goal)
        figure.text(
            0.055, 0.9, "\n".join(title_lines), ha="left", va="top", fontsize=9, weight="bold"
        )
        figure.text(
            0.42,
            0.9,
            f"ID {_clean(goal.get('goal_id'))}",
            ha="left",
            va="top",
            fontsize=7,
            color="#526071",
        )
        identifier = goal.get("identifier") or goal.get("condition")
        figure.text(
            0.59,
            0.9,
            f"идентификатор: {_clean(identifier)}" if identifier else "",
            ha="left",
            va="top",
            fontsize=7,
            color="#526071",
        )
        metrics = (
            ("#7A45E5", "Конверсия", _number(goal.get("conversion_rate"), "%", decimal_places=2)),
            ("#FF3399", "Целевые визиты", _number(goal.get("visits"), decimal_places=0)),
            ("#0FBDA0", "Достижения цели", _number(goal.get("reaches"), decimal_places=0)),
        )
        for y, (color, label, value) in zip((0.59, 0.43, 0.27), metrics, strict=True):
            figure.patches.append(
                FancyBboxPatch(
                    (0.032, y - 0.018),
                    0.018,
                    0.036,
                    transform=figure.transFigure,
                    boxstyle="round,pad=0.001,rounding_size=0.006",
                    facecolor=color,
                    edgecolor=color,
                )
            )
            figure.text(
                0.041,
                y,
                "✓",
                color="#FFFFFF",
                fontsize=5.5,
                va="center",
                ha="center",
                weight="bold",
            )
            figure.text(0.057, y, label, color="#30343B", fontsize=7.5, va="center")
            figure.text(0.31, y, value, color="#30343B", fontsize=7.5, va="center", ha="right")

        left = figure.add_axes((0.42, 0.17, 0.53, 0.48))
        right = left.twinx()
        x = list(range(len(labels)))
        left.plot(x, conversion, color="#7A45E5", linewidth=1.25)
        right.plot(x, visits, color="#FF3399", linewidth=1.25)
        right.plot(x, reaches, color="#0FBDA0", linewidth=1.25)
        left.set_xticks(x, labels)
        _style_axis(left)
        left.set_ylim(0, max(max(conversion, default=0) * 1.2, 0.01))
        left.yaxis.set_major_locator(MaxNLocator(nbins=3))
        left.yaxis.set_major_formatter(FuncFormatter(lambda value, _position: f"{value:g}%"))
        right.grid(False)
        right.set_ylim(0, max(max([*visits, *reaches], default=0) * 1.2, 1))
        right.yaxis.set_major_locator(MaxNLocator(nbins=3, integer=True))
        right.yaxis.set_major_formatter(
            FuncFormatter(lambda value, _position: _axis_compact_number(value))
        )
        right.tick_params(colors="#8B98A7", labelsize=7, length=0)
        for spine in right.spines.values():
            spine.set_visible(False)
        left.tick_params(labelsize=6.5)
        return _save_figure(figure)


def _project_favicon_bytes(payload):
    encoded = (payload.get("project", {}).get("favicon") or {}).get("data")
    if not encoded:
        return None
    try:
        return base64.b64decode(encoded, validate=True)
    except (TypeError, ValueError):
        return None


def _style_metrika_url_column(table, payload):
    favicon = _project_favicon_bytes(payload)
    labels = [row.cells[0].text for row in table.rows[1:]]
    depths = [
        min(4, len([part for part in urlsplit(label).path.split("/") if part])) for label in labels
    ]
    for row_index, row in enumerate(table.rows[1:]):
        cell = row.cells[0]
        label = cell.text
        if label == "Итого и среднее":
            for run in cell.paragraphs[0].runs:
                _style_run(run, size=11, bold=True)
            continue
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        depth = depths[row_index]
        if depth == 1:
            expanded = row_index + 1 < len(depths) and depths[row_index + 1] == 2
            _style_run(
                paragraph.add_run("⌄  " if expanded else "›  "),
                size=9,
                color="9AA5B1",
            )
        if favicon:
            paragraph.add_run().add_picture(io.BytesIO(favicon), width=Cm(0.28))
            paragraph.add_run("  ")
        link = paragraph.add_run(label)
        link.font.color.rgb = RGBColor.from_string("5277D5")
        link.font.name = "Calibri"
        link._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
        link.font.size = Pt(11)
        paragraph.paragraph_format.left_indent = Cm(depth * 0.16)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _landing_pages_table(doc, payload, current, previous, *, limit=20):
    ordered = sorted(current.items(), key=lambda item: item[1]["visits"], reverse=True)[:limit]
    table = _metrika_detail_table(
        doc,
        [(url, values, previous.get(url, {})) for url, values in ordered],
        first_header="Страница входа",
        metrics=("visits", "users"),
    )
    _style_metrika_url_column(table, payload)
    return table


def _add_metrics(target, values):
    visits = _decimal_or_none(values.get("visits")) or Decimal(0)
    users = _decimal_or_none(values.get("users")) or Decimal(0)
    bounce = _decimal_or_none(values.get("bounce_rate")) or Decimal(0)
    target["visits"] += visits
    target["users"] += users
    target["bounce_sum"] += bounce * visits


def _landing_hierarchy(pages):
    result = {}
    for url, values in pages.items():
        parsed = urlsplit(url)
        if not parsed.netloc:
            continue
        parts = [part for part in parsed.path.split("/") if part]
        prefixes = [f"{parsed.scheme or 'https'}://{parsed.netloc}/"]
        for depth in (1, 2):
            if len(parts) >= depth:
                prefixes.append(
                    f"{parsed.scheme or 'https'}://{parsed.netloc}/" + "/".join(parts[:depth]) + "/"
                )
        for prefix in dict.fromkeys(prefixes):
            target = result.setdefault(
                prefix,
                {"visits": Decimal(0), "users": Decimal(0), "bounce_sum": Decimal(0)},
            )
            _add_metrics(target, values)
    for values in result.values():
        values["bounce_rate"] = (
            values.pop("bounce_sum") / values["visits"] if values["visits"] else Decimal(0)
        )
    return result


def _landing_hierarchy_order(hierarchy, expanded_patterns=None):
    by_domain = {}
    for url in hierarchy:
        parsed = urlsplit(url)
        by_domain.setdefault(parsed.netloc, []).append(url)
    ordered = []
    for domain in sorted(by_domain):
        urls = by_domain[domain]
        roots = [url for url in urls if not [p for p in urlsplit(url).path.split("/") if p]]
        ordered.extend(sorted(roots))
        sections = [
            url for url in urls if len([p for p in urlsplit(url).path.split("/") if p]) == 1
        ]
        sections.sort(key=lambda url: (-hierarchy[url]["visits"], url.casefold()))
        for section in sections:
            ordered.append(section)
            section_path = urlsplit(section).path.rstrip("/") + "/"
            should_expand = expanded_patterns is None or any(
                urlsplit(_url_pattern_label(pattern).replace("*", "").replace("?", ""))
                .path.rstrip("/")
                .startswith(section_path.rstrip("/"))
                for pattern in expanded_patterns
            )
            if not should_expand:
                continue
            section_part = [p for p in urlsplit(section).path.split("/") if p][0]
            children = [
                url
                for url in urls
                if len([p for p in urlsplit(url).path.split("/") if p]) == 2
                and [p for p in urlsplit(url).path.split("/") if p][0] == section_part
            ]
            children.sort(key=lambda url: (-hierarchy[url]["visits"], url.casefold()))
            ordered.extend(children)
    return ordered


def _landing_hierarchy_table(
    doc, payload, current, previous, *, metrics=("visits", "users"), expanded_groups=()
):
    current_hierarchy = _landing_hierarchy(current)
    previous_hierarchy = _landing_hierarchy(previous)
    expanded_patterns = [
        pattern for group in expanded_groups for pattern in group.get("patterns") or []
    ]
    ordered = _landing_hierarchy_order(current_hierarchy, expanded_patterns)
    current_total = next(
        (values for url, values in current_hierarchy.items() if urlsplit(url).path in {"", "/"}),
        {},
    )
    previous_total = next(
        (values for url, values in previous_hierarchy.items() if urlsplit(url).path in {"", "/"}),
        {},
    )
    table = _metrika_detail_table(
        doc,
        [(url, current_hierarchy[url], previous_hierarchy.get(url, {})) for url in ordered],
        first_header="Страница входа",
        metrics=metrics,
        total_values=(current_total, previous_total),
    )
    _style_metrika_url_column(table, payload)
    return table


def _landing_comparison_table(
    doc, payload, current_rows, previous_rows, engine, *, expanded_groups=()
):
    current = _aggregate_detail_rows(
        [row for row in current_rows if _search_engine_name(row) == engine], _landing_url
    )
    previous = _aggregate_detail_rows(
        [row for row in previous_rows if _search_engine_name(row) == engine], _landing_url
    )
    return _landing_hierarchy_table(
        doc, payload, current, previous, expanded_groups=expanded_groups
    )


def _configured_group_label(group):
    for pattern in group.get("patterns") or []:
        value = _url_pattern_label(pattern).strip()
        if value.startswith(("http://", "https://")):
            return value.rstrip("*?")
    return group.get("name") or "Раздел"


def _conclusion_page_label(url, groups):
    matching = [
        group
        for group in groups
        if any(_url_matches_pattern(url, pattern) for pattern in group.get("patterns") or [])
    ]
    if matching:
        matching.sort(
            key=lambda group: max(
                (len(_url_pattern_label(pattern)) for pattern in group.get("patterns") or []),
                default=0,
            ),
            reverse=True,
        )
        return f"«{matching[0]['name']}» ({url})"
    return f"страница ({url})"


def _configured_groups_table(
    doc, payload, current_rows, previous_rows, _groups, *, expanded_groups=()
):
    current_pages = _aggregate_detail_rows(current_rows, _landing_url)
    previous_pages = _aggregate_detail_rows(previous_rows, _landing_url)
    return _landing_hierarchy_table(
        doc,
        payload,
        current_pages,
        previous_pages,
        expanded_groups=expanded_groups,
    )


def _goal_card(doc, goal, periods):
    image = _metrika_goal_image(goal, periods)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run().add_picture(image, width=Cm(18.2))
    return paragraph


def _render_metrika(doc, payload, blocks):
    traffic_enabled = section_enabled(payload, "traffic")
    sources_enabled = section_enabled(payload, "traffic_sources")
    geography_enabled = section_enabled(payload, "geography")
    options = payload.get("display_options", {})
    if not options.get("include_metrika", True):
        return
    doc.add_heading("3) Сводная информация по переходам на сайт (Яндекс.Метрика)", level=1)
    doc.add_paragraph(
        "Трафик представлен за квартал по всем источникам в распределении по месяцам."
    )
    period_details = _period_details(payload, "yandex_metrika")
    if sources_enabled:
        doc.add_paragraph("Источники, сводка", style="Table Heading")
        if period_details:
            _period_caption(
                doc,
                [
                    {"date": str(period_details[0]["period_start"])[:10]},
                    {"date": str(period_details[-1]["period_end"])[:10]},
                ],
                detail="месяцам",
            )
        facts = _metric_source(payload, "yandex_metrika").get("traffic_source_dynamics", {})
        _add_report_picture(doc, _metrika_sources_chart(facts))
        if options.get("include_metrika_sources_table"):
            source_periods = _metrika_period_rows(payload, "traffic_source_details")
            if not any(period.get("rows") for period in source_periods):
                source_periods = [
                    {
                        "rows": [
                            {
                                "code": name,
                                "visits": (fact.get("change") or {}).get("current") or 0,
                                "users": 0,
                                "bounce_rate": 0,
                                "page_depth": 0,
                                "avg_visit_duration_seconds": 0,
                            }
                            for name, fact in facts.items()
                        ]
                    }
                ]
            if options.get("metrika_sources_compare_previous") and source_periods:
                current = _aggregate_traffic_source_rows(source_periods[-1:])
                previous = _aggregate_traffic_source_rows(source_periods[-2:-1])
                ordered = sorted(
                    (code for code, values in current.items() if values["visits"] > 0),
                    key=lambda code: current[code]["visits"],
                    reverse=True,
                )
                if ordered:
                    source_totals = None
                    if len(source_periods) >= 2:
                        current_total = source_periods[-1].get("total") or {}
                        previous_total = source_periods[-2].get("total") or {}
                        if current_total and previous_total:
                            source_totals = (current_total, previous_total)
                    table = _metrika_detail_table(
                        doc,
                        [
                            (_metrika_source_label(code), current[code], previous.get(code, {}))
                            for code in ordered
                        ],
                        first_header="Источник",
                        total_values=source_totals,
                    )
                    _style_traffic_source_icons(table)
            else:
                _metrika_sources_quarter_table(doc, source_periods)
    if traffic_enabled:
        search_periods = _metrika_period_rows(payload, "search_engines")
        if section_enabled(payload, "metrika_search_engines") and search_periods:
            doc.add_paragraph("Детально по поисковым системам", style="Table Heading")
            current = _aggregate_detail_rows(search_periods[-1]["rows"], _search_engine_name)
            previous = (
                _aggregate_detail_rows(search_periods[-2]["rows"], _search_engine_name)
                if len(search_periods) >= 2
                else {}
            )
            ordered = sorted(
                (
                    label
                    for label in current
                    if current[label].get("visits", 0) > 0
                    and previous.get(label, {}).get("visits", 0) > 0
                ),
                key=lambda key: current[key]["visits"],
                reverse=True,
            )
            colors = {
                "Google": "#7A45E5",
                "Яндекс": "#FF3399",
                "Bing": "#0FBDA0",
                "Yahoo": "#3388FF",
                "Yahoo!": "#3388FF",
            }
            selected_bar_engines = options.get("metrika_bar_search_engines")
            chart_labels = (
                [label for label in ordered if _search_engine_code(label) in selected_bar_engines]
                if selected_bar_engines is not None
                else ordered
            )
            chart_rows = [
                {
                    "label": label,
                    "previous": previous.get(label, {}).get("visits"),
                    "current": current[label].get("visits"),
                    "color": colors.get(label),
                }
                for label in chart_labels
            ]
            doc.add_paragraph(
                "Динамика по поисковым системам за квартал:",
                style="Chart Caption",
            )
            _period_caption(
                doc,
                [
                    {"date": str(search_periods[0]["period_start"])[:10]},
                    {"date": str(search_periods[-1]["period_end"])[:10]},
                ],
                detail="месяцам",
            )
            _add_report_picture(doc, _metrika_search_quarter_chart(search_periods))
            directions = []
            for label in ("Яндекс", "Google"):
                delta = _relative_delta(
                    current.get(label, {}).get("visits"), previous.get(label, {}).get("visits")
                )
                if delta is not None:
                    direction = (
                        "увеличилось"
                        if delta > 0
                        else "снизилось"
                        if delta < 0
                        else "не изменилось"
                    )
                    directions.append((label, direction))
            if directions:
                if len(directions) == 2 and directions[0][1] == directions[1][1]:
                    doc.add_paragraph(
                        "Число переходов из Яндекса и Google " + directions[0][1] + "."
                    )
                else:
                    doc.add_paragraph(
                        "Число переходов "
                        + ", ".join(
                            f"из {'Яндекса' if label == 'Яндекс' else label} {direction}"
                            for label, direction in directions
                        )
                        + "."
                    )
            doc.add_paragraph(
                "По поисковым системам в сравнении двух последних месяцев:",
                style="Table Heading",
            )
            _comparison_period_pills(doc, search_periods)
            _add_report_picture(
                doc,
                _metrika_comparison_bars(chart_rows, title="Визиты из поисковых систем"),
            )
            doc.add_paragraph(
                "Детально в цифрах по поисковым системам в сравнении двух месяцев:",
                style="Table Heading",
            )
            search_table = _metrika_detail_table(
                doc,
                [(label, current[label], previous.get(label, {})) for label in ordered],
                first_header="Поисковая система",
                total_values=(
                    (search_periods[-1].get("total") or {}, search_periods[-2].get("total") or {})
                    if len(search_periods) >= 2
                    and search_periods[-1].get("total")
                    and search_periods[-2].get("total")
                    else None
                ),
            )
            _style_search_engine_icons(search_table)
            movement = []
            for label in ("Яндекс", "Google"):
                if label not in current:
                    continue
                delta = _relative_delta(
                    current[label].get("visits"), previous.get(label, {}).get("visits")
                )
                if delta is not None:
                    movement.append(
                        f"Трафик из {label} {'увеличился' if delta >= 0 else 'снизился'} "
                        f"на {_number(abs(delta), '%', decimal_places=1)}"
                    )
            total_current = sum((row["visits"] for row in current.values()), Decimal(0))
            total_previous = sum((row["visits"] for row in previous.values()), Decimal(0))
            total_delta = _relative_delta(total_current, total_previous)
            if total_delta is not None:
                movement.append(
                    f"Общий поисковый трафик {'увеличился' if total_delta >= 0 else 'снизился'} "
                    f"на {_number(abs(total_delta), '%', decimal_places=1)}"
                )
            doc.add_paragraph(
                ". ".join(movement)
                + (". " if movement else "")
                + _metrika_search_engine_text(current)
            )
    if geography_enabled:
        geography_periods = _metrika_period_rows(payload, "search_geography")
        current = _aggregate_regions(geography_periods[-1]["rows"]) if geography_periods else {}
        previous = (
            _aggregate_regions(geography_periods[-2]["rows"]) if len(geography_periods) >= 2 else {}
        )
        flags = {
            "moscow": "geography_moscow",
            "moscow_region": "geography_moscow_region",
            "saint_petersburg": "geography_saint_petersburg",
            "saint_petersburg_region": "geography_saint_petersburg_region",
            "undefined": "geography_undefined",
            "area_undefined": "geography_area_undefined",
        }
        chart_selected = [
            key for key, flag in flags.items() if options.get(flag, True) and key in current
        ]
        table_selected = list(chart_selected)
        doc.add_paragraph(
            "Сравнение трафика за 2 последних месяца по основным регионам по поисковому трафику:",
            style="Table Heading",
        )
        _comparison_period_pills(doc, geography_periods)
        _add_report_picture(
            doc,
            _metrika_comparison_bars(
                [
                    {
                        "label": REGION_LABELS[key],
                        "previous": previous.get(key, {}).get("visits"),
                        "current": current[key].get("visits"),
                        "color": METRIKA_COLORS[index % len(METRIKA_COLORS)],
                    }
                    for index, key in enumerate(chart_selected)
                ],
                title="Поисковый трафик по регионам",
            ),
        )
        if table_selected:
            _metrika_detail_table(
                doc,
                [
                    (REGION_LABELS[key], current[key], previous.get(key, {}))
                    for key in table_selected
                ],
                first_header="Регион",
                total_values=(
                    (
                        geography_periods[-1].get("total") or {},
                        geography_periods[-2].get("total") or {},
                    )
                    if len(geography_periods) >= 2
                    and geography_periods[-1].get("total")
                    and geography_periods[-2].get("total")
                    else None
                ),
            )
            total_current = sum((row["visits"] for row in current.values()), Decimal(0))
            total_previous = sum((row["visits"] for row in previous.values()), Decimal(0))
            for key in table_selected:
                if key in {"undefined", "area_undefined"}:
                    continue
                delta = _relative_delta(
                    current[key].get("visits"), previous.get(key, {}).get("visits")
                )
                current_share = (
                    current[key]["visits"] * 100 / total_current if total_current else None
                )
                previous_share = (
                    previous.get(key, {}).get("visits", 0) * 100 / total_previous
                    if total_previous
                    else None
                )
                doc.add_paragraph(
                    f"Трафик из региона «{REGION_LABELS[key]}» "
                    f"{'увеличился' if delta is not None and delta >= 0 else 'снизился'} на "
                    f"{_number(abs(delta), '%', decimal_places=1) if delta is not None else '—'}. "
                    f"Изменение доли трафика: предыдущий месяц — "
                    f"{_number(previous_share, '%', decimal_places=1)}, отчётный месяц — "
                    f"{_number(current_share, '%', decimal_places=1)}."
                )

    landing_periods = [
        {
            **period,
            "rows": [row for row in period.get("rows") or [] if _belongs_to_project(payload, row)],
        }
        for period in _metrika_period_rows(payload, "landing_pages")
    ]
    if landing_periods:
        current_rows = landing_periods[-1]["rows"]
        previous_rows = landing_periods[-2]["rows"] if len(landing_periods) >= 2 else []
        category_groups = _configured_url_groups(payload, "categories")
        subsection_groups = _configured_url_groups(payload, "subsections")
        expanded_groups = [*category_groups, *subsection_groups]
        current_pages = _aggregate_detail_rows(current_rows, _landing_url)
        previous_pages = _aggregate_detail_rows(previous_rows, _landing_url)
        if section_enabled(payload, "metrika_landing_pages"):
            doc.add_paragraph("Популярные страницы входа", style="Table Heading")
            doc.add_paragraph(
                "Ниже приведены значения количества переходов в отчётном месяце по сравнению "
                "с предыдущим "
                + (
                    "только по поисковому трафику по страницам."
                    if options.get("metrika_search_segment", True)
                    else "по всему трафику по страницам."
                )
            )
            _landing_pages_table(doc, payload, current_pages, previous_pages)
            total = sum((row["visits"] for row in current_pages.values()), Decimal(0))
            root = next(
                (
                    (url, values)
                    for url, values in current_pages.items()
                    if urlsplit(url).path in {"", "/"}
                ),
                None,
            )
            internal = [
                (url, values)
                for url, values in current_pages.items()
                if urlsplit(url).path not in {"", "/"}
            ]
            popular = max(internal, key=lambda item: item[1]["visits"], default=None)
            text = []
            if root and total:
                text.append(
                    f"На главную страницу приходится "
                    f"{_number(root[1]['visits'] * 100 / total, '%', decimal_places=2)} всех "
                    f"визитов ({_number(root[1]['visits'], decimal_places=0)} визитов за месяц)"
                )
            if popular and total:
                text.append(
                    "Самой популярной внутренней страницей является "
                    f"{_conclusion_page_label(popular[0], expanded_groups)} "
                    f"({_number(popular[1]['visits'], decimal_places=0)} визитов — "
                    f"{_number(popular[1]['visits'] * 100 / total, '%', decimal_places=2)})"
                )
            if text:
                doc.add_paragraph(". ".join(text) + ".")
        if section_enabled(payload, "metrika_landing_page_comparison"):
            comparison_subsection_groups = _configured_url_groups(
                payload, "landing_comparison_subsections"
            )
            comparison_expanded_groups = [*category_groups, *comparison_subsection_groups]
            comparison_periods = [
                {
                    **period,
                    "rows": [
                        row for row in period.get("rows") or [] if _belongs_to_project(payload, row)
                    ],
                }
                for period in _metrika_period_rows(payload, "search_landing_pages")
            ]
            comparison_current = comparison_periods[-1]["rows"] if comparison_periods else []
            comparison_previous = (
                comparison_periods[-2]["rows"] if len(comparison_periods) >= 2 else []
            )
            commercial_groups = _configured_url_groups(payload, "commercial")
            named_conclusion_groups = [
                group
                for expected in ("лечение", "диагностика", "реабилитация")
                for group in commercial_groups
                if group["name"].casefold() == expected
            ] or commercial_groups
            for engine in ("Яндекс", "Google"):
                doc.add_paragraph(
                    "Страницы входа",
                    style="Table Heading",
                )
                doc.add_paragraph(
                    "Поисковые переходы по страницам входа только по поисковому трафику "
                    f"в сравнении двух последних месяцев по {engine}:"
                )
                _comparison_period_pills(doc, comparison_periods)
                _landing_comparison_table(
                    doc,
                    payload,
                    comparison_current,
                    comparison_previous,
                    engine,
                    expanded_groups=comparison_expanded_groups,
                )
                if named_conclusion_groups:
                    engine_periods = [
                        {
                            **period,
                            "rows": [
                                row
                                for row in period.get("rows") or []
                                if _search_engine_name(row) == engine
                            ],
                        }
                        for period in comparison_periods
                    ]
                    doc.add_paragraph(
                        _group_dynamics_text(
                            engine_periods,
                            named_conclusion_groups,
                            prefix=f"{engine}, раздел",
                            subsections=comparison_subsection_groups,
                        )
                    )
        information_groups = _configured_url_groups(payload, "information")
        commercial_groups = _configured_url_groups(payload, "commercial")
        if section_enabled(payload, "metrika_url_groups"):
            doc.add_paragraph(
                "Сравнение трафика, приходящегося на информационные и коммерческие страницы",
                style="Table Heading",
            )
            _comparison_period_pills(doc, landing_periods)
            _configured_groups_table(
                doc,
                payload,
                current_rows,
                previous_rows,
                [*information_groups, *commercial_groups],
                expanded_groups=(),
            )
            if information_groups:
                doc.add_paragraph(_group_overview_text(landing_periods, information_groups))
                doc.add_paragraph(
                    _group_dynamics_text(
                        landing_periods,
                        information_groups,
                        prefix="Информационный раздел",
                        subsections=subsection_groups,
                    )
                )
            if commercial_groups:
                doc.add_paragraph(
                    _group_overview_text(landing_periods, commercial_groups, commercial=True)
                )
                doc.add_paragraph(
                    _group_dynamics_text(
                        landing_periods,
                        commercial_groups,
                        prefix="Коммерческий раздел",
                        subsections=subsection_groups,
                    )
                )
        if section_enabled(payload, "metrika_sections"):
            doc.add_paragraph("Данные по разделам", style="Table Heading")
            for heading, groups in (
                ("Информационные разделы", information_groups),
                ("Коммерческие разделы", commercial_groups),
            ):
                if not groups:
                    continue
                doc.add_paragraph(heading, style="Chart Caption")
                _period_caption(
                    doc,
                    [
                        {"date": str(landing_periods[0]["period_start"])[:10]},
                        {"date": str(landing_periods[-1]["period_end"])[:10]},
                    ],
                    detail="месяцам",
                )
                _add_report_picture(doc, _metrika_groups_chart(landing_periods, groups))
                doc.add_paragraph(
                    _group_dynamics_text(landing_periods, groups, subsections=subsection_groups)
                )
            if not information_groups and not commercial_groups:
                doc.add_paragraph(
                    "URL-сегменты информационных и коммерческих разделов не настроены.",
                    style="Data Missing",
                )
        if section_enabled(payload, "metrika_categories"):
            doc.add_paragraph("Основные прорабатываемые категории", style="Table Heading")
            combined_categories = options.get("metrika_categories_combined", False)
            groups_to_render = (
                [category_groups]
                if combined_categories and category_groups
                else [[group] for group in category_groups]
            )
            for selected_groups in groups_to_render:
                if not combined_categories:
                    doc.add_paragraph(selected_groups[0]["name"], style="Table Heading")
                _period_caption(
                    doc,
                    [
                        {"date": str(landing_periods[0]["period_start"])[:10]},
                        {"date": str(landing_periods[-1]["period_end"])[:10]},
                    ],
                    detail="месяцам",
                )
                _add_report_picture(
                    doc,
                    _metrika_groups_chart(
                        landing_periods,
                        category_groups,
                        only_group=None if combined_categories else selected_groups[0]["name"],
                    ),
                    width=18.5,
                )
                doc.add_paragraph(
                    _group_dynamics_text(
                        landing_periods,
                        selected_groups,
                        prefix="Раздел",
                        subsections=subsection_groups,
                    )
                )
            if not category_groups:
                doc.add_paragraph(
                    "URL-сегменты прорабатываемых категорий не настроены.",
                    style="Data Missing",
                )

    if section_enabled(payload, "metrika_goals") and period_details:
        robotness = options.get("metrika_robotness", "humans")
        segment = "search" if options.get("metrika_search_segment", True) else "all"
        goal_periods = []
        for detail in period_details:
            source_payload = _detail_payload(detail)
            goal_periods.append(
                {
                    "period_start": detail.get("period_start"),
                    "rows": ((source_payload.get("goals_by_segment") or {}).get(segment) or {}).get(
                        robotness
                    )
                    or (source_payload.get("goals_by_robotness") or {}).get(robotness)
                    or source_payload.get("goals")
                    or [],
                }
            )
        current_goals = goal_periods[-1]["rows"]
        if current_goals:
            doc.add_heading("4) Сводная информация по конверсии (Яндекс.Метрика).", level=1)
            robotness_label = "только люди" if robotness == "humans" else "все визиты"
            segment_label = "переходы из поисковых систем" if segment == "search" else "весь трафик"
            segment_description = (
                "только по поисковому трафику" if segment == "search" else "по всему трафику"
            )
            doc.add_paragraph(
                "Ниже приведены диаграммы конверсии за квартал с детализацией "
                f"по месяцам {segment_description}."
            )
            _period_caption(
                doc,
                [
                    {"date": str(period_details[0]["period_start"])[:10]},
                    {"date": str(period_details[-1]["period_end"])[:10]},
                ],
                detail="месяцам",
            )
            doc.add_paragraph(
                f"Сегмент: {segment_label}. Роботность: {robotness_label}.", style="Compact"
            )
            for goal in current_goals:
                goal_id = str(goal.get("goal_id"))
                series = []
                for period in goal_periods:
                    row = next(
                        (item for item in period["rows"] if str(item.get("goal_id")) == goal_id),
                        None,
                    )
                    if row:
                        series.append({"period_start": period["period_start"], **row})
                summary_goal = goal
                if options.get("metrika_goals_quarter", True) and series:
                    visits = sum(
                        (_decimal_or_none(item.get("visits")) or 0 for item in series), Decimal(0)
                    )
                    reaches = sum(
                        (_decimal_or_none(item.get("reaches")) or 0 for item in series), Decimal(0)
                    )
                    bases = []
                    for item in series:
                        rate = _decimal_or_none(item.get("conversion_rate")) or 0
                        item_visits = _decimal_or_none(item.get("visits")) or 0
                        if rate:
                            bases.append(item_visits * 100 / rate)
                    conversion = (
                        visits * 100 / sum(bases, Decimal(0))
                        if bases and sum(bases, Decimal(0))
                        else 0
                    )
                    summary_goal = {
                        **goal,
                        "visits": visits,
                        "reaches": reaches,
                        "conversion_rate": conversion,
                    }
                _goal_card(doc, summary_goal, series)


def _change_rows(payload, source, codes):
    changes = _metric_source(payload, source).get("normalized_changes", {})
    rows = []
    for code in codes:
        change = changes.get(code)
        if not change:
            rows.append(
                (
                    METRIC_LABELS.get(code, code),
                    "Данные недоступны",
                    "Данные недоступны",
                    "Данные недоступны",
                    "Данные недоступны",
                )
            )
            continue
        unit = "%" if change.get("percentage_points") is not None else ""
        delta_label = "процентных пунктов" if unit else ""
        rows.append(
            (
                METRIC_LABELS.get(code, code),
                _number(change.get("previous"), unit),
                _number(change.get("current"), unit),
                _number(change.get("absolute"), f" {delta_label}"),
                _number(change.get("relative_percent"), "%"),
            )
        )
    return rows


def _add_hyperlink(paragraph, text, url, *, bold=False, italic=False):
    relationship = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "5277D5")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    if bold:
        properties.append(OxmlElement("w:b"))
    if italic:
        properties.append(OxmlElement("w:i"))
    run.append(properties)
    content = OxmlElement("w:t")
    content.text = text
    run.append(content)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class _RichTextDocxParser(HTMLParser):
    def __init__(self, doc):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.paragraph = None
        self.lists = []
        self.bold = False
        self.italic = False
        self.link = ""

    def _paragraph(self, style=None):
        self.paragraph = self.doc.add_paragraph(style=style)
        self.paragraph.paragraph_format.space_after = Pt(4)
        return self.paragraph

    def handle_starttag(self, tag, attrs):
        tag = tag.casefold()
        if tag == "p":
            self._paragraph()
        elif tag in {"ul", "ol"}:
            self.lists.append(tag)
        elif tag == "li":
            kind = "Number" if self.lists and self.lists[-1] == "ol" else "Bullet"
            level = min(len(self.lists), 3)
            style = f"List {kind}" + (f" {level}" if level > 1 else "")
            self._paragraph(style)
        elif tag == "br":
            (self.paragraph or self._paragraph()).add_run("\n")
        elif tag == "strong":
            self.bold = True
        elif tag == "em":
            self.italic = True
        elif tag == "a":
            self.link = dict(attrs).get("href", "")

    def handle_endtag(self, tag):
        tag = tag.casefold()
        if tag in {"p", "li"}:
            self.paragraph = None
        elif tag in {"ul", "ol"} and self.lists:
            self.lists.pop()
        elif tag == "strong":
            self.bold = False
        elif tag == "em":
            self.italic = False
        elif tag == "a":
            self.link = ""

    def handle_data(self, data):
        if not data:
            return
        paragraph = self.paragraph or self._paragraph()
        if self.link:
            _add_hyperlink(
                paragraph,
                data,
                self.link,
                bold=self.bold,
                italic=self.italic,
            )
            return
        run = paragraph.add_run(data)
        run.bold = self.bold
        run.italic = self.italic


def _render_rich_text(doc, value):
    parser = _RichTextDocxParser(doc)
    parser.feed(value)
    parser.close()


def _render_work(doc, payload, narrative):
    show_urls = payload.get("display_options", {}).get("show_urls", True)
    manual_text = payload.get("display_options", {}).get("completed_work_text", "").strip()
    if manual_text:
        _render_rich_text(doc, manual_text)
        return
    works = payload.get("completed_work", [])
    if works:
        for work in works:
            paragraph = doc.add_paragraph(style="List Number")
            title = work.get("title") or work.get("category") or "Работа"
            paragraph.add_run(_clean(title)).bold = True
            comment_lines = [
                re.sub(r"^[\s•●\-–—\d.)]+", "", line).strip()
                for line in str(work.get("comment") or "").splitlines()
                if line.strip()
            ]
            for line in comment_lines:
                doc.add_paragraph(_clean(line), style="List Bullet 2")
            if show_urls:
                for label, value in (
                    ("Страница", work.get("page_or_material_name") or work.get("url")),
                    ("Результат", work.get("result_url")),
                ):
                    if value:
                        doc.add_paragraph(f"{label}: {_clean(value)}", style="Compact")
    else:
        doc.add_paragraph(
            _clean(narrative) or "Выполненные работы отсутствуют.", style="Data Missing"
        )


def _configure_document(doc, domain, period):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = Cm(0.75)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.0)
    for name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.color.rgb = RGBColor.from_string("000000")
    doc.styles["Normal"].font.size = Pt(11)
    for name, size in (("Heading 1", 12), ("Heading 2", 11), ("Heading 3", 11)):
        doc.styles[name].font.size = Pt(size)
        doc.styles[name].font.bold = True
        doc.styles[name].paragraph_format.keep_with_next = True
        doc.styles[name].paragraph_format.space_after = Pt(6)
    for style_name, size, italic in (
        ("Chart Caption", 9, True),
        ("Data Missing", 10, True),
        ("KPI", 11, False),
        ("Depth Note", 9, True),
        ("Compact", 10, False),
        ("Table Heading", 11, False),
        ("Table Comment", 11, False),
    ):
        style = doc.styles.add_style(style_name, 1)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.italic = italic
    doc.styles["Compact"].paragraph_format.space_after = Pt(0)
    doc.styles["Table Heading"].font.bold = True
    doc.styles["Table Heading"].paragraph_format.keep_with_next = True
    doc.styles["Table Comment"].paragraph_format.space_before = Pt(4)
    table_style = doc.styles.add_style("Report Table", 3)
    table_style.base_style = doc.styles["Table Grid"]
    table_style.font.name = "Calibri"
    table_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    table_style.font.size = Pt(11)


def _docx(snapshot, narratives, issues, draft):
    payload = snapshot.payload
    project = payload.get("project", {})
    period = payload.get("periods", {}).get("report", {}).get("start")
    doc = Document()
    _configure_document(doc, project.get("domain"), period)
    doc.core_properties.author = "SEO"
    doc.core_properties.comments = "https://t.me/wmasterfl"
    if draft:
        paragraph = doc.add_paragraph("ЧЕРНОВИК")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(26)
        paragraph.runs[0].font.color.rgb = RGBColor(190, 35, 35)
    title = f"Отчёт по поисковому продвижению сайта {project.get('domain')} за {_month(period)}"
    doc.add_heading(_clean(title), 0)
    doc.add_paragraph(_clean(project.get("name")))
    doc.add_paragraph(f"Дата формирования: {timezone.localdate():%d.%m.%Y}")
    doc.add_page_break()
    narrative_rows = list(narratives)
    blocks = {block.section_code: block.effective_text for block in narrative_rows}
    table_comments = {
        block.section_code: block.edited_text.strip()
        for block in narrative_rows
        if block.edited_text.strip()
    }
    engine_order = {"yandex": 0, "google": 1}
    segments = sorted(
        payload.get("calculated", {}).get("positions", {}).get("segments", []),
        key=lambda item: (
            engine_order.get(item.get("search_engine"), 99),
            item.get("region") or "",
        ),
    )
    _render_topvisor(doc, payload, segments, table_comments)
    _render_webmaster(doc, payload, blocks)
    _render_metrika(doc, payload, blocks)
    warning_messages = {
        _clean(issue.message)
        for issue in issues
        if issue.severity == "warning"
        and section_enabled(payload, issue.section_code)
        and issue.section_code != "completed_work"
    }
    for message in sorted(warning_messages):
        doc.add_paragraph("Предупреждение: " + message, style="Depth Note")
    if section_enabled(payload, "completed_work"):
        doc.add_heading(TITLES["completed_work"], level=1)
        _render_work(
            doc,
            payload,
            blocks.get("completed_work") or "Выполненные работы отсутствуют.",
        )
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _xlsx_value(value):
    return _excel_safe(value) if isinstance(value, str) else value


def _xlsx(snapshot, draft):
    payload = snapshot.payload
    show_urls = payload.get("display_options", {}).get("show_urls", True)
    workbook = Workbook()
    metadata = workbook.active
    metadata.title = "Метаданные"
    project = payload.get("project", {})
    period = payload.get("periods", {}).get("report", {}).get("start")
    depths = "; ".join(
        f"{s.get('search_engine')} / {s.get('region')}: TOP-{s.get('ranking_depth')}"
        for s in payload.get("calculated", {}).get("positions", {}).get("segments", [])
    )
    for row in (
        ("проект", project.get("name")),
        ("домен", project.get("domain")),
        ("месяц", date.fromisoformat(str(period)[:10])),
        ("версия", snapshot.version.number),
        ("checksum snapshot", snapshot.checksum),
        ("дата создания", timezone.localtime(snapshot.created_at).replace(tzinfo=None)),
        ("formula_version", snapshot.formula_version),
        ("глубины", depths),
        ("черновик", draft),
    ):
        metadata.append(tuple(_xlsx_value(value) for value in row))
    positions = workbook.create_sheet("Позиции")
    position_headers = [
        "Поисковая система",
        "Регион",
        "Дата",
        "Запрос",
        "Частотность",
        "Позиция",
        "Статус",
        "Группа",
        "Фактическая глубина",
    ]
    if show_urls:
        position_headers.insert(8, "Релевантный URL")
    positions.append(position_headers)
    engine_order = {"yandex": 0, "google": 1}
    ranking_sources = sorted(
        payload.get("ranking_sources", []),
        key=lambda item: (
            engine_order.get(item.get("search_engine"), 99),
            item.get("region") or "",
            item.get("date") or "",
        ),
    )
    for source in ranking_sources:
        for row in source.get("positions", []):
            values = [
                source.get("search_engine"),
                source.get("region"),
                date.fromisoformat(source["date"]),
                row.get("query"),
                row.get("frequency"),
                row.get("position"),
                row.get("status"),
                row.get("group"),
                source.get("ranking_depth"),
            ]
            if show_urls:
                values.insert(8, row.get("target_url"))
            positions.append(tuple(_xlsx_value(value) for value in values))
    history = workbook.create_sheet("История")
    history.append(("Система", "Регион", "Месяц", "Видимость", "Глубина", "Распределение"))
    for raw_segment in payload.get("calculated", {}).get("positions", {}).get("segments", []):
        segment = _manual_topvisor_segment(payload, raw_segment)
        for row in segment.get("three_month_series", []):
            history.append(
                (
                    segment.get("search_engine"),
                    segment.get("region"),
                    date.fromisoformat(row["month"]),
                    row.get("visibility"),
                    row.get("ranking_depth"),
                    _excel_safe(str((row.get("distribution") or {}).get("ranges", {}))),
                )
            )
    metrics = workbook.create_sheet("Метрика и Вебмастер")
    metrics.append(("Источник", "Начало", "Конец", "Показатель", "Значение", "Единица"))
    for source in payload.get("source_snapshots", []):
        for metric in source.get("metrics", []):
            metrics.append(
                tuple(
                    _xlsx_value(value)
                    for value in (
                        source.get("source"),
                        date.fromisoformat(source["period_start"]),
                        date.fromisoformat(source["period_end"]),
                        metric.get("code"),
                        float(metric["value"]) if metric.get("value") is not None else None,
                        metric.get("unit"),
                    )
                )
            )
    traffic_facts = _metric_source(payload, "yandex_metrika").get("traffic_source_dynamics", {})
    for name, fact in traffic_facts.items():
        for point in fact.get("series", []):
            metrics.append(
                (
                    "yandex_metrika",
                    date.fromisoformat(point["month"]),
                    date.fromisoformat(point["month"]),
                    f"traffic_source_{_excel_safe(name)}_monthly_total",
                    float(point["value"]) if point.get("value") is not None else None,
                    "count",
                )
            )
    work = workbook.create_sheet("Выполненные работы")
    work_headers = ["Дата", "Категория", "Название", "Статус"]
    if show_urls:
        work_headers.extend(("Страница или материал", "URL", "Результат"))
    else:
        work_headers.append("Материал")
    work_headers.extend(("Объём", "Ответственный", "Комментарий"))
    work.append(work_headers)
    for item in payload.get("completed_work", []):
        values = [
            date.fromisoformat(item["date"]),
            item.get("category"),
            item.get("title"),
            item.get("status"),
        ]
        if show_urls:
            values.extend(
                (item.get("page_or_material_name"), item.get("url"), item.get("result_url"))
            )
        else:
            values.append(item.get("page_or_material_name"))
        values.extend((item.get("character_count"), item.get("responsible"), item.get("comment")))
        work.append(tuple(_xlsx_value(value) for value in values))
    for sheet in workbook.worksheets:
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="315B7D")
        for column in sheet.columns:
            sheet.column_dimensions[column[0].column_letter].width = min(
                55, max(12, max(len(str(cell.value or "")) for cell in column) + 2)
            )
            for cell in column:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _pdf(docx_bytes):
    with tempfile.TemporaryDirectory(prefix="seo-export-") as temporary:
        root = Path(temporary)
        source = root / "report.docx"
        source.write_bytes(docx_bytes)
        profile = root / "lo-profile"
        profile.mkdir()
        office_binary = shutil.which("libreoffice") or shutil.which("soffice")
        if office_binary is None:
            raise RuntimeError("LibreOffice is not installed")
        result = subprocess.run(
            [
                office_binary,
                "--headless",
                f"-env:UserInstallation={profile.as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(root),
                str(source),
            ],
            capture_output=True,
            text=True,
            timeout=settings.REPORT_PDF_TIMEOUT_SECONDS,
            check=False,
        )
        pdf = root / "report.pdf"
        if (
            result.returncode
            or not pdf.exists()
            or pdf.stat().st_size == 0
            or not pdf.read_bytes().startswith(b"%PDF-")
        ):
            raise RuntimeError(f"LibreOffice conversion failed (code {result.returncode})")
        return pdf.read_bytes()


def generate_artifact(*, version, artifact_type, is_draft=False, created_by=None):
    artifact = GeneratedArtifact.objects.create(
        report_version=version,
        artifact_type=artifact_type,
        is_draft=is_draft,
        created_by=created_by,
    )
    try:
        readiness = get_publication_readiness(version)
        if readiness.has_errors and not is_draft:
            raise ExportBlocked("Финальный экспорт заблокирован ошибками валидации.")
        snapshot = ReportDatasetSnapshot.objects.select_related("version").get(version=version)
        narratives = list(
            NarrativeBlock.objects.filter(report_version=version).order_by(
                "sort_order", "created_at"
            )
        )
        issues = list(ValidationIssue.objects.filter(version=version))
        log = "; ".join(_clean(issue.message) for issue in issues if issue.severity == "warning")
        if artifact_type == "docx":
            data = _docx(snapshot, narratives, issues, is_draft)
        elif artifact_type == "xlsx":
            data = _xlsx(snapshot, is_draft)
        elif artifact_type == "pdf":
            data = _pdf(_docx(snapshot, narratives, issues, is_draft))
        else:
            raise ValueError("Unsupported artifact type")
        domain = (
            re.sub(
                r"[^a-zA-Z0-9.-]+",
                "-",
                str(snapshot.payload.get("project", {}).get("normalized_domain") or "report"),
            ).strip(".-")
            or "report"
        )
        raw_month = str(snapshot.payload.get("periods", {}).get("report", {}).get("start"))[:10]
        parsed_month = date.fromisoformat(raw_month)
        month = MONTHS[parsed_month.month - 1]
        year = parsed_month.year
        suffix = "_draft" if is_draft else ""
        filename = f"{domain}_отчет_за_{month}_{year}_v{version.number}{suffix}.{artifact_type}"
        artifact.file.save(filename, ContentFile(data), save=False)
        artifact.filename = filename
        artifact.mime_type = MIMES[artifact_type]
        artifact.size = len(data)
        artifact.sha256 = hashlib.sha256(data).hexdigest()
        artifact.generation_log = log
        artifact.status = GeneratedArtifact.Status.READY
        artifact.generator_version = GENERATOR_VERSION
        artifact.save()
        return artifact
    except Exception as exc:
        artifact.status = GeneratedArtifact.Status.FAILED
        artifact.generation_log = "Экспорт не завершён (" + type(exc).__name__ + ")."
        artifact.save(update_fields=["status", "generation_log"])
        raise
