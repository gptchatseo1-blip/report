"""Compact Topvisor distribution summary table for DOCX/PDF exports."""

import math

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

_APPLIED = False


def _style_text(exp, cell, *, size=11, color="3D4655", bold=False, align=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            exp._style_run(run, size=size, color=color, bold=bold)


def _collapse_empty_cell_paragraphs(cell):
    """Nested tables require empty paragraphs; collapse them to one typographic point."""
    for paragraph in cell.paragraphs:
        if paragraph.text.strip():
            continue
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(1)
        run = paragraph.add_run("\u200b")
        run.font.size = Pt(1)


def _render_distribution_cards_table(exp, doc, distribution, depth, *, engine="yandex"):
    """Render a narrower 11 pt card table without visible empty spacer lines."""
    buckets = exp._topvisor_buckets(distribution, 20 if engine == "google" else depth)
    if engine == "google":
        buckets = [bucket for bucket in buckets if bucket["label"] in {"1-3", "1-10", "11-20"}]
    if not buckets:
        return None

    columns = 1 if engine == "google" else 2
    row_count = math.ceil(len(buckets) / columns)
    table = doc.add_table(rows=row_count, cols=columns)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    outer_width = 4.35 if columns == 2 else 4.5

    for row in table.rows:
        exp._prevent_row_split(row)
        for cell in row.cells:
            exp._set_cell_width(cell, outer_width)
            exp._set_cell_margins(cell, top=6, bottom=6, left=12, right=12)
            exp._shade_cell(cell, "F0F2F5")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""

    for index, bucket in enumerate(buckets):
        column = index // row_count
        row_index = index % row_count
        outer = table.rows[row_index].cells[column]
        nested = outer.add_table(rows=1, cols=3)
        nested.autofit = False
        widths = (1.35, 0.95, 1.55)
        for grid_column, width in zip(nested._tbl.tblGrid.gridCol_lst, widths, strict=True):
            grid_column.w = exp.Cm(width)
        label_cell, share_cell, count_cell = nested.rows[0].cells
        for cell, width in zip((label_cell, share_cell, count_cell), widths, strict=True):
            exp._set_cell_width(cell, width)
            exp._set_cell_margins(cell, top=4, bottom=4, left=10, right=10)
            exp._shade_cell(cell, "F0F2F5")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        exp._shade_cell(label_cell, exp.TOPVISOR_COLORS[bucket["label"]])
        label_cell.text = str(bucket["label"])
        share_cell.text = exp._number(bucket.get("share"), "%", decimal_places=0)
        count_cell.text = exp._number(bucket.get("count"), decimal_places=0)

        _style_text(
            exp,
            label_cell,
            size=11,
            color="FFFFFF",
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _style_text(
            exp,
            share_cell,
            size=11,
            color="8491A5",
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _style_text(
            exp,
            count_cell,
            size=11,
            color="3D4655",
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        exp._set_table_borders(nested, "F0F2F5", size="0")
        _collapse_empty_cell_paragraphs(outer)

    if len(buckets) % columns:
        empty_column = columns - 1
        empty_cell = table.rows[-1].cells[empty_column]
        exp._shade_cell(empty_cell, "FFFFFF")
        exp._set_cell_margins(empty_cell, top=0, bottom=0, left=0, right=0)
        _collapse_empty_cell_paragraphs(empty_cell)

    exp._set_table_borders(table, "FFFFFF", size="6")
    exp._keep_small_table_together(table)
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(1)
    gap.paragraph_format.line_spacing = Pt(1)
    gap.add_run("\u200b").font.size = Pt(1)
    return table


def apply():
    global _APPLIED
    if _APPLIED:
        return
    _APPLIED = True

    from . import exporting as exp

    exp._render_distribution_cards_table = lambda doc, distribution, depth, *, engine="yandex": (
        _render_distribution_cards_table(exp, doc, distribution, depth, engine=engine)
    )
    exp.GENERATOR_VERSION = "mvp1.12-2026-09-05"
